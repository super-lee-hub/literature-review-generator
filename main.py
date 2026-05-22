#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
文献综述自动生成器 - 工业级版本
支持身份基断点续传、双重工作模式、智能续写、项目命名空间、智能文件查找、双引擎PDF提取、适应性速率控制、并发处理、错误管理、自动重试机制和交互式安装向导。

# 作者: auto-generate 文献综述自动生成器开发团队
版本: 1.2
更新日期: 2025-10-15
"""

import sys
import os
import time
import argparse
import traceback
import concurrent.futures
import threading
import json
import hashlib
import logging
import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Set, Tuple, Iterator, Union, Mapping, Sequence
from datetime import datetime

ThreadPoolExecutor = concurrent.futures.ThreadPoolExecutor

# 添加项目根目录到Python路径
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# 导入项目模块
from models import (
    PaperInfo, ProcessingResult, FailedPaper, SummariesList,
    APIConfig, AISummary
)
from config_loader import load_config, ConfigDict
from zotero_parser import parse_zotero_report
from file_finder import create_file_index, FileIndex, find_pdf
from pdf_extractor import extract_text_from_pdf  # type: ignore
from ai_interface import (  # type: ignore
    get_summary_from_ai,
    get_summary_from_ai_with_fallback,
    get_concept_analysis,
    _call_ai_api,
    _call_ai_api_text_detailed,
    _smart_json_parser,
    _auto_correct_json,
)
from docx_writer import create_word_document, append_section_to_word_document, generate_word_table_of_contents, generate_apa_references_from_manifest
from report_generator import generate_excel_report, generate_failure_report, generate_retry_zotero_report  # type: ignore
from preprocess.service import PreprocessManager
from preprocess.visual_artifacts import Stage1VisualArtifactBuilder
from services.artifact_registry import ArtifactDependencyRef, ArtifactRegistry
from services.citation_metadata import sanitize_metadata_fields
from services.config_compat import CompatConfigView
from services.environment_service import (
    detect_runtime_environment,
    recommended_conda_activate_command,
    recommended_conda_create_command,
)
from services.job_workspace import JobWorkspace, atomic_write_json
from services.paper_artifact import build_paper_artifact_v1
from services.paper_identity import build_paper_key as build_canonical_paper_key, normalize_doi
from services.progress_state import (
    ResumeStateReport,
    Stage1ProgressSnapshot,
    load_stage1_progress_snapshot,
    write_stage1_progress_snapshot,
)
from services.queue_service import CancelToken, PersistentQueueService, QueueJobSpec, QueueState, create_queue_job_id
from services.review_draft import build_review_draft_v1, build_review_draft_v2
from services.stage1_input_builder import Stage1InputBuilder
from services.stage1_input_completeness import is_blocked_stage1_quality
from services.model_selection import (
    get_backup_reader_api_config,
    get_outline_api_config,
    get_reader_api_config,
    get_writer_api_config,
)
from services.source_normalizer import normalize_source_papers, project_descriptors_to_legacy_papers
from services.summary_reuse import (
    ResolvedSummarySet,
    SummaryCatalog,
    SummaryMatch,
    SummarySource,
    SummarySourceError,
    build_effective_summary_set,
    build_reused_summary,
    collect_summary_sources,
    describe_summary_candidate,
    load_summary_records,
)
from services.text_io import load_json_file_with_fallbacks
from utils import ensure_dir
from setup_wizard import run_setup_wizard
from summary_schema import get_core_analysis, get_paper_metadata
from free_mode.profile_manager import build_profile_context, load_profile


def get_paper_key(paper: 'Dict[str, Any] | PaperInfo') -> str:
    return build_canonical_paper_key(paper)

def normalize_checkpoint_paper_key(paper_key: str) -> str:
    normalized = normalize_doi(paper_key)
    if normalized:
        return normalized
    return str(paper_key or "").strip()

from context_manager import validate_summary_quality, optimize_context_for_synthesis, optimize_context_for_outline, estimate_tokens



# 优雅地处理可选依赖
try:
    from docx import Document  # type: ignore
    from docx.shared import Pt  # type: ignore
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    _docx_available = True
except ImportError:
    _docx_available = False
    Document = None  # type: ignore
    Pt = None  # type: ignore
    WD_PARAGRAPH_ALIGNMENT = None  # type: ignore
    qn = None  # type: ignore

# 定义常量，避免重定义问题
DOCX_AVAILABLE = _docx_available

try:
    from tqdm import tqdm  # type: ignore
    _tqdm_available = True
except ImportError:
    _tqdm_available = False
    # 创建一个假的tqdm类以避免在代码中进行大量的if检查
    class tqdm:
        def __init__(self, iterable: Optional[Any] = None, **kwargs: Any) -> None:
            self.iterable: List[Any] = iterable if iterable else []
        def __iter__(self) -> Iterator[Any]:
            return iter(self.iterable)
        def __enter__(self) -> 'tqdm':
            return self
        def __exit__(self, exc_type: Any, exc_val: Any, exc_tb: Any) -> None:
            pass
        def update(self, n: int = 1) -> None:
            pass
        def set_postfix_str(self, s: str) -> None:
            pass
        def close(self) -> None:
            pass

# 定义常量，避免重定义问题
TQDM_AVAILABLE = _tqdm_available

# 在文件开头打印警告信息（使用logging而不是print）
if not DOCX_AVAILABLE:
    logging.warning("未安装 'python-docx'。生成Word文档和第二阶段验证功能将不可用。请运行: pip install python-docx")
if not TQDM_AVAILABLE:
    logging.warning("未安装 'tqdm'。将无法显示进度条。请运行: pip install tqdm")

class CustomLogger(logging.Logger):
    def success(self, msg: str, *args: Any, **kwargs: Any) -> None:
        self.info(f"[SUCCESS] {msg}", *args, **kwargs)
    def warn(self, msg: str, *args: Any, **kwargs: Any) -> None:
        self.warning(f"[WARN] {msg}", *args, **kwargs)

logging.setLoggerClass(CustomLogger)

# ==========================================================

class Counter:
    """线程安全计数器 - 简化版本提高性能"""
    def __init__(self, initial_value: int = 0):
        self._value = initial_value
        self._lock = threading.Lock()

    def increment(self) -> int:
        with self._lock:
            self._value += 1
            return self._value

    def decrement(self) -> int:
        with self._lock:
            self._value -= 1
            return self._value

    @property
    def value(self) -> int:
        """获取当前值（属性方式访问）"""
        with self._lock:
            return self._value

    def set(self, new_value: int) -> None:
        """设置值，确保线程安全"""
        with self._lock:
            self._value = new_value

    def get_value(self) -> int:
        """获取当前值，避免属性装饰器的开销"""
        with self._lock:
            return self._value

    def set_value(self, new_value: int) -> None:
        """设置值，确保线程安全"""
        with self._lock:
            self._value = new_value

class ReportingService:
    """报告生成服务 - 专门负责生成各种分析报告"""

    def __init__(self, logger: CustomLogger):
        self.logger: CustomLogger = logger if logger is not None else logging.getLogger(__name__)  # type: ignore

    def generate_all_reports(self, generator: 'LiteratureReviewGenerator') -> None:
        """生成所有分析阶段的报告"""
        self.logger.info("正在生成所有分析报告...")

        # 生成Excel报告
        if not generate_excel_report(generator):
            self.logger.warning("Excel报告生成失败，但不影响整体处理结果")

        # 生成失败报告（如果有失败的论文）
        if generator.failed_papers:
            if not generate_failure_report(generator):
                self.logger.warning("失败报告生成失败，但不影响整体处理结果")

        # 只在Zotero模式下生成自动化重跑报告
        if generator.mode == "zotero" and generator.failed_papers:
            if not generate_retry_zotero_report(generator):
                self.logger.warning("重跑报告生成失败，但不影响整体处理结果")

        self.logger.success("所有分析报告生成完毕。")


class CheckpointManager:
    """检查点管理器 - 专门负责处理基于身份的断点续传"""

    def __init__(self, logger: CustomLogger):
        self.logger: CustomLogger = logger or logging.getLogger(__name__)  # type: ignore

    def save_checkpoint(self, generator: 'LiteratureReviewGenerator') -> bool:
        """保存基于身份的断点文件"""
        try:
            if not generator.output_dir or not generator.project_name:
                return False

            checkpoint_file = generator._get_stage1_checkpoint_file_path()

            # 创建已处理论文的身份集合
            processed_papers: Set[str] = set()
            for summary in generator.summaries:
                if summary.get('status') == 'success':
                    paper_info: PaperInfo = summary.get('paper_info', {})  # type: ignore
                    paper_key: str = LiteratureReviewGenerator.get_paper_key(paper_info)  # type: ignore
                    processed_papers.add(paper_key)

            # 创建失败论文的身份集合
            failed_papers: Set[str] = set()
            for failed_item in generator.failed_papers:
                paper_info: PaperInfo = failed_item.get('paper_info', {})  # type: ignore
                paper_key: str = LiteratureReviewGenerator.get_paper_key(paper_info)  # type: ignore
                failed_papers.add(paper_key)

            checkpoint_data: Dict[str, Any] = {
                'version': '2.0',  # 身份基断点版本
                'project_name': generator.project_name,
                'update_time': datetime.now().isoformat(),
                'total_papers': len(generator.papers),
                'processed_count': len(processed_papers),
                'failed_count': len(failed_papers),
                'processed_papers': list(processed_papers),  # 基于身份的已处理列表
                'failed_papers': list(failed_papers),        # 基于身份的失败列表
                'processing_stats': {
                    'processed_success': generator.processed_count.value,
                    'failed_attempts': generator.failed_count.value
                }
            }

            atomic_write_json(checkpoint_file, checkpoint_data)

            self.logger.info(f"[断点保存] 已保存处理进度: {len(processed_papers)}成功, {len(failed_papers)}失败")
            return True

        except Exception as e:
            self.logger.error(f"保存断点文件失败: {e}")
            return False

    def load_checkpoint(self, generator: 'LiteratureReviewGenerator') -> bool:
        """加载基于身份的断点文件"""
        try:
            if not generator.output_dir or not generator.project_name:
                return False

            checkpoint_file = generator._get_stage1_checkpoint_file_path()

            if not os.path.exists(checkpoint_file):
                self.logger.info("[断点加载] 未找到断点文件，将开始全新处理")
                return False

            with open(checkpoint_file, 'r', encoding='utf-8') as f:
                checkpoint_data: Dict[str, Any] = json.load(f)

            # 验证断点文件版本
            version = checkpoint_data.get('version', '1.0')
            if version != '2.0':
                self.logger.warning(f"[断点加载] 检测到旧版本断点文件(v{version})，将开始全新处理")
                return False

            # 验证项目名称匹配
            checkpoint_project = checkpoint_data.get('project_name')
            if checkpoint_project != generator.project_name:
                self.logger.warning(f"[断点加载] 项目名称不匹配({checkpoint_project} != {generator.project_name})，将开始全新处理")
                return False

            # 提取已处理和失败的论文身份
            processed_papers = {
                normalize_checkpoint_paper_key(item)
                for item in checkpoint_data.get('processed_papers', [])
                if isinstance(item, str) and normalize_checkpoint_paper_key(item)
            }
            failed_papers = {
                normalize_checkpoint_paper_key(item)
                for item in checkpoint_data.get('failed_papers', [])
                if isinstance(item, str) and normalize_checkpoint_paper_key(item)
            }
            update_time = checkpoint_data.get('update_time', '未知时间')

            self.logger.info(f"[断点加载] 成功加载断点文件 (更新时间: {update_time})")
            self.logger.info(f"[断点加载] 已处理论文: {len(processed_papers)}篇")
            self.logger.info(f"[断点加载] 失败论文: {len(failed_papers)}篇")

            # 将断点信息存储到实例变量中，供process_all_papers使用
            generator._checkpoint_processed_papers = processed_papers  # type: ignore
            generator._checkpoint_failed_papers = failed_papers  # type: ignore

            # 恢复计数器
            processing_stats: Dict[str, Any] = checkpoint_data.get('processing_stats') or {}
            generator.processed_count.set(processing_stats.get('processed_success', 0))  # type: ignore
            generator.failed_count.set(processing_stats.get('failed_attempts', 0))  # type: ignore

            return True

        except Exception as e:
            self.logger.error(f"加载断点文件失败: {e}")
            return False


class LiteratureReviewGenerator:
    """文献综述生成器主类"""
    
    logger: CustomLogger
    OUTLINE_ARTIFACT_ID = "literature_review_outline"
    OUTLINE_ARTIFACT_TYPE = "literature_review_outline"
    OUTLINE_ARTIFACT_ROLE = "outline"
    OUTLINE_ARTIFACT_VERSION = "v1"
    PAPER_ARTIFACT_TYPE = "paper_artifact"
    PAPER_ARTIFACT_ROLE = "paper"
    PAPER_ARTIFACT_VERSION = "v1"
    REVIEW_DRAFT_ARTIFACT_ID = "review_draft:full_review"
    REVIEW_DRAFT_ARTIFACT_TYPE = "review_draft"
    REVIEW_DRAFT_ARTIFACT_ROLE = "review_draft"
    REVIEW_DRAFT_ARTIFACT_VERSION = "v1"
    REVIEW_DRAFT_V2_ARTIFACT_ID = "review_draft_v2:full_review"
    REVIEW_DRAFT_V2_ARTIFACT_TYPE = "review_draft"
    REVIEW_DRAFT_V2_ARTIFACT_ROLE = "review_draft_v2"
    REVIEW_DRAFT_V2_ARTIFACT_VERSION = "v2"
    CITATION_MANIFEST_ARTIFACT_ID = "citation_manifest:v3"
    CITATION_MANIFEST_ARTIFACT_TYPE = "citation_manifest"
    CITATION_MANIFEST_ARTIFACT_ROLE = "citation_manifest"
    CITATION_MANIFEST_ARTIFACT_VERSION = "v3"
    
    def __init__(self, config_file: str = 'config.ini', project_name: Optional[str] = None, pdf_folder: Optional[str] = None, queue_file: str = 'output/_queue/queue.json', zotero_report: Optional[str] = None, library_path: Optional[str] = None):
        self.config_file: str = config_file
        self.project_name: Optional[str] = project_name
        self.pdf_folder: Optional[str] = pdf_folder
        self.queue_file: str = queue_file
        self.zotero_report: Optional[str] = zotero_report
        self.library_path: Optional[str] = library_path
        self.config: Optional['ConfigDict'] = None
        self.compat_config: Optional[CompatConfigView] = None
        self.output_dir: Optional[str] = None
        self.summary_file: Optional[str] = None
        self.summary_file_override: Optional[str] = None
        self.summary_source_overrides: List[str] = []
        self.reuse_stage1: bool = False
        self.reuse_summary_files: List[str] = []
        self.papers: List[PaperInfo] = []
        self.source_descriptors: List[Dict[str, Any]] = []
        self.summaries: SummariesList = []
        self.failed_papers: List[FailedPaper] = []
        self.preprocess_manager: Optional[PreprocessManager] = None
        self.processed_count: Counter = Counter(0)
        self.failed_count: Counter = Counter(0)
        self.save_lock: threading.Lock = threading.Lock()
        self.progress_tracker: Optional[Any] = None
        self.free_mode_profile_path: Optional[str] = None
        self.free_mode_profile: Optional[Dict[str, Any]] = None
        self.free_mode_idea: Optional[str] = None
        self.cancel_token: Optional[CancelToken] = None
        self.job_workspace: Optional[JobWorkspace] = None
        self.artifact_registry: Optional[ArtifactRegistry] = None
        self.job_fingerprint_bundle: Dict[str, Any] = {}
        self.resume_state_report: Optional[ResumeStateReport] = None
        self.queue_service: Optional[PersistentQueueService] = None
        self._stage1_reuse_report: Optional[Dict[str, Any]] = None
        self._stage1_reused_paper_keys: Set[str] = set()
        self.stage1_result_summary: Dict[str, Any] = {}
        self.root_log_path: str = ""
        self.workspace_log_path: str = ""
        self._workspace_log_handler: Optional[logging.Handler] = None
        self._stage1_reader_engine_lock: threading.Lock = threading.Lock()
        self._stage1_disabled_reader_engines: Set[str] = set()
        self._stage1_reader_disable_reasons: Dict[str, str] = {}

        # 身份基断点续传相关变量
        self._checkpoint_processed_papers: Set[str] = set()
        self._checkpoint_failed_papers: Set[str] = set()

        # 概念增强模式相关变量
        self.concept_mode: bool = False
        self.concept_profile: Optional[Dict[str, Any]] = None

        # 根据参数确定运行模式
        if pdf_folder:
            self.mode: str = "direct"  # 直接PDF模式
            self.pdf_folder = os.path.abspath(pdf_folder)
        else:
            self.mode: str = "zotero"  # Zotero模式（默认）

        # 初始化日志记录器
        self._init_logger()

        # 初始化服务组件
        self.reporting_service: ReportingService = ReportingService(self.logger)
        self.checkpoint_manager: CheckpointManager = CheckpointManager(self.logger)
        self.stage1_visual_builder = Stage1VisualArtifactBuilder(logger=self.logger)
        self.stage1_input_builder = Stage1InputBuilder(logger=self.logger)

    def bind_job_workspace(
        self,
        *,
        workspace: JobWorkspace,
        artifact_registry: ArtifactRegistry,
        compat_config: CompatConfigView,
        fingerprint_bundle: Dict[str, Any],
        resume_state_report: ResumeStateReport,
    ) -> None:
        self.job_workspace = workspace
        self.artifact_registry = artifact_registry
        self.compat_config = compat_config
        self.job_fingerprint_bundle = dict(fingerprint_bundle)
        self.resume_state_report = resume_state_report
        self.project_name = workspace.project_name
        self.output_dir = workspace.root_dir
        self.summary_file = workspace.artifact_path(f"{workspace.project_name}_summaries.json")
        self._attach_workspace_log_handler(workspace, artifact_registry)
        
        # 初始化队列服务
        if self.queue_file and os.path.exists(self.queue_file):
            # 如果提供了 queue_file 且文件存在，使用它
            queue_file_path = Path(self.queue_file)
        else:
            # 否则使用默认路径
            queue_file_path = Path(workspace.root_dir) / "_queue" / "queue.json"
        self.queue_service = PersistentQueueService(queue_file_path)

    def _attach_workspace_log_handler(self, workspace: JobWorkspace, artifact_registry: ArtifactRegistry) -> None:
        workspace_log_path = workspace.log_path("job.log")
        os.makedirs(os.path.dirname(workspace_log_path), exist_ok=True)
        logger_handlers = list(getattr(self.logger, "handlers", []) or [])

        if self.workspace_log_path == workspace_log_path and self._workspace_log_handler in logger_handlers:
            return

        if self._workspace_log_handler is not None:
            try:
                if hasattr(self.logger, "removeHandler"):
                    self.logger.removeHandler(self._workspace_log_handler)
                self._workspace_log_handler.close()
            except Exception:
                pass

        formatter = None
        if logger_handlers:
            formatter = logger_handlers[0].formatter
        if formatter is None:
            formatter = logging.Formatter('[%(asctime)s] [%(levelname)s] %(message)s', datefmt='%H:%M:%S')

        if hasattr(self.logger, "addHandler"):
            handler = logging.FileHandler(workspace_log_path, encoding='utf-8')
            handler.setLevel(logging.INFO)
            handler.setFormatter(formatter)
            self.logger.addHandler(handler)
            self._workspace_log_handler = handler
        else:
            Path(workspace_log_path).touch()
            self._workspace_log_handler = None
        self.workspace_log_path = workspace_log_path
        artifact_registry.register_file(
            artifact_role="log",
            artifact_type="job_log",
            artifact_version="v1",
            path=workspace_log_path,
            producer="main.LiteratureReviewGenerator._attach_workspace_log_handler",
            artifact_id="job_log",
        )
        self.logger.info(f"工作区日志已创建: {workspace_log_path}")

    def _init_queue_service(self) -> None:
        """初始化队列服务（向后兼容：在没有 job_workspace 时使用）"""
        if self.queue_service is None:
            # 始终使用传入的 queue_file 参数，确保所有队列操作使用同一个队列文件
            queue_file_path = Path(self.queue_file)
            # 确保队列文件目录存在
            queue_file_path.parent.mkdir(parents=True, exist_ok=True)
            self.queue_service = PersistentQueueService(queue_file_path)

    def submit_job_to_queue(self, job_type: str, parameters: Dict[str, Any]) -> str:
        """提交任务到队列"""
        self._init_queue_service()
        if self.queue_service is None or not self.project_name:
            raise RuntimeError("Queue service not initialized or project name not set")
        
        job_id = create_queue_job_id()
        job_spec = QueueJobSpec(
            job_id=job_id,
            job_type=job_type,
            project_name=self.project_name,
            parameters=parameters
        )
        return self.queue_service.add_job(job_spec)

    def list_queue_jobs(self) -> List[QueueJobSpec]:
        """列出所有队列任务"""
        self._init_queue_service()
        if self.queue_service is None:
            return []
        return self.queue_service.list_jobs()

    def cancel_queue_job(self, job_id: str) -> bool:
        """取消队列任务"""
        self._init_queue_service()
        if self.queue_service is None:
            return False
        return self.queue_service.update_job_state(job_id, QueueState.CANCELLED)

    def retry_queue_job(self, job_id: str) -> bool:
        """重试失败的队列任务"""
        self._init_queue_service()
        if self.queue_service is None:
            return False
        return self.queue_service.reset_job(job_id) and self.queue_service.increment_retry_count(job_id) > 0

    def clear_completed_queue_jobs(self) -> int:
        """清空已完成的队列任务"""
        self._init_queue_service()
        if self.queue_service is None:
            return 0
        count = 0
        for job in self.queue_service.list_jobs():
            runtime = self.queue_service.get_job_runtime(job.job_id)
            if runtime and runtime.state == QueueState.COMPLETED:
                self.queue_service.remove_job(job.job_id)
                count += 1
        return count

    def _get_stage1_checkpoint_file_path(self) -> str:
        if self.job_workspace and self.project_name:
            return self.job_workspace.checkpoint_path(f"{self.project_name}_checkpoint.json")
        if self.output_dir and self.project_name:
            return os.path.join(self.output_dir, f"{self.project_name}_checkpoint.json")
        return "checkpoint.json"

    def _get_summary_file_path(self) -> str:
        if self.summary_file:
            return self.summary_file
        if self.job_workspace and self.project_name:
            return self.job_workspace.artifact_path(f"{self.project_name}_summaries.json")
        if self.output_dir and self.project_name:
            return os.path.join(self.output_dir, f"{self.project_name}_summaries.json")
        return "summaries.json"

    def _get_summary_source_manifest_path(self) -> str:
        if self.job_workspace and self.project_name:
            return self.job_workspace.artifact_path(f"{self.project_name}_summary_source_manifest.json")
        if self.output_dir and self.project_name:
            return os.path.join(self.output_dir, f"{self.project_name}_summary_source_manifest.json")
        return "summary_source_manifest.json"

    def _get_summary_reuse_report_path(self) -> str:
        if self.job_workspace and self.project_name:
            return self.job_workspace.artifact_path(f"{self.project_name}_summary_reuse_report.json")
        if self.output_dir and self.project_name:
            return os.path.join(self.output_dir, f"{self.project_name}_summary_reuse_report.json")
        return "summary_reuse_report.json"

    def _get_report_file_path(self, suffix: str) -> str:
        if self.job_workspace and self.project_name:
            return self.job_workspace.report_path(f"{self.project_name}{suffix}")
        if self.output_dir and self.project_name:
            return os.path.join(self.output_dir, f"{self.project_name}{suffix}")
        return suffix.lstrip("_")

    def _get_concept_profile_file_path(self) -> str:
        if self.job_workspace and self.project_name:
            return self.job_workspace.artifact_path(f"{self.project_name}_concept_profile.json")
        if self.output_dir and self.project_name:
            return os.path.join(self.output_dir, f"{self.project_name}_concept_profile.json")
        return "concept_profile.json"

    def _keep_checkpoints_after_completion(self) -> bool:
        if self.compat_config:
            return self.compat_config.keep_checkpoints_after_completion()
        return False

    def _get_stage1_config(self) -> Dict[str, Any]:
        """Return Stage 1 strict/partial-success configuration.

        Defaults: allow_partial_success=false, min_success_ratio=1.0.
        """
        stage1_raw = (self.config or {}).get("Stage1", {}) or {}
        if isinstance(stage1_raw, dict):
            allow_partial = bool(stage1_raw.get("allow_partial_success", False))
            try:
                min_ratio = float(stage1_raw.get("min_success_ratio", 1.0))
            except (TypeError, ValueError):
                min_ratio = 1.0
            return {
                "allow_partial_success": allow_partial,
                "min_success_ratio": max(0.0, min(1.0, min_ratio)),
            }
        return {"allow_partial_success": False, "min_success_ratio": 1.0}

    def _stage1_validation_enabled(self) -> bool:
        if not self.compat_config and self.config is not None:
            self.compat_config = CompatConfigView.from_config(self.config)
        if self.compat_config:
            return self.compat_config.stage1_validation_enabled()
        return False

    def _stage2_validation_enabled(self) -> bool:
        if not self.compat_config and self.config is not None:
            self.compat_config = CompatConfigView.from_config(self.config)
        if self.compat_config:
            return self.compat_config.stage2_validation_enabled()
        return False

    def _ensure_compat_config(self) -> CompatConfigView:
        if not self.config:
            raise RuntimeError("Configuration is not loaded")
        if not self.compat_config:
            self.compat_config = CompatConfigView.from_config(self.config)
        return self.compat_config

    def _outline_v2_enabled(self) -> bool:
        if not self.config:
            return False
        return self._ensure_compat_config().outline_v2_enabled()

    def _outline_v2_model_call(self, route_name: str, prompt: str, metadata: Dict[str, Any]) -> Any:
        """Call the configured v2 model route and return provider JSON."""
        config = self.config or {}
        api_params = config.get("API_Parameters", {}) if isinstance(config.get("API_Parameters", {}), dict) else {}
        if route_name == "Outline_API":
            api_config = get_outline_api_config(config)
            max_token_key = "outline_max_tokens"
            default_max_tokens = 16000
            default_temperature = 0.4
        elif route_name == "Writer_API":
            api_config = get_writer_api_config(config)
            max_token_key = "writer_max_tokens"
            default_max_tokens = 32000
            default_temperature = 0.7
        elif route_name == "Primary_Reader_API":
            api_config = get_reader_api_config(config)
            max_token_key = "primary_max_tokens"
            default_max_tokens = 5000
            default_temperature = 0.3
        else:
            api_config = get_outline_api_config(config)
            max_token_key = "outline_max_tokens"
            default_max_tokens = 16000
            default_temperature = 0.4

        try:
            max_tokens = int(api_config.get("max_tokens") or api_params.get(max_token_key) or default_max_tokens)
        except (TypeError, ValueError):
            max_tokens = default_max_tokens
        try:
            temperature = float(api_params.get(max_token_key.replace("max_tokens", "temperature"), default_temperature))
        except (TypeError, ValueError):
            temperature = default_temperature
        stage = str(metadata.get("stage") or "outline_v2")
        default_timeout_seconds = 600 if stage == "outline_candidates" else 180
        timeout_key = f"outline_v2_{stage}_timeout_seconds"
        retries_key = f"outline_v2_{stage}_retry_attempts"
        try:
            timeout_seconds = int(api_params.get(timeout_key) or api_params.get("outline_v2_timeout_seconds") or default_timeout_seconds)
        except (TypeError, ValueError):
            timeout_seconds = default_timeout_seconds
        try:
            retry_attempts = int(api_params.get(retries_key) or api_params.get("outline_v2_retry_attempts") or (2 if stage == "outline_candidates" else 1))
        except (TypeError, ValueError):
            retry_attempts = 2 if stage == "outline_candidates" else 1

        system_prompt = (
            "You are an Outline Intelligence v2 structured JSON generator. "
            "Use only the provided controlled-corpus evidence and return strict JSON."
        )
        candidate_index = metadata.get("candidate_index")
        self.logger.info(
            f"Outline v2 model call start: stage={stage}, candidate_index={candidate_index}, "
            f"prompt_chars={len(prompt)}, max_tokens={max_tokens}, timeout_seconds={timeout_seconds}, "
            f"retry_attempts={retry_attempts}, route={route_name}"
        )
        response = _call_ai_api_text_detailed(
            prompt=prompt,
            api_config=api_config,
            system_prompt=system_prompt,
            max_tokens=max_tokens,
            temperature=temperature,
            logger=self.logger,
            retry_attempts=retry_attempts,
            timeout_seconds=timeout_seconds,
        )
        content = str(response.get("content") or "")
        parsed = _smart_json_parser(content) or _auto_correct_json(content)
        if parsed is not None:
            self.logger.info(f"Outline v2 model call parsed JSON: stage={stage}, candidate_index={candidate_index}")
            return parsed
        self.logger.error(
            f"Outline v2 model call returned unparseable JSON: stage={stage}, "
            f"candidate_index={candidate_index}, message={response.get('message', '')}"
        )
        return None

    def _load_paper_artifacts_for_outline_v2(self) -> List[Dict[str, Any]]:
        if not self.job_workspace:
            return []
        artifacts_dir = os.path.join(self.job_workspace.paths.artifacts_dir, "paper_artifacts")
        if not os.path.isdir(artifacts_dir):
            return []
        artifacts: List[Dict[str, Any]] = []
        for filename in sorted(os.listdir(artifacts_dir)):
            if not filename.endswith(".json"):
                continue
            path = os.path.join(artifacts_dir, filename)
            try:
                with open(path, "r", encoding="utf-8") as handle:
                    payload = json.load(handle)
                if isinstance(payload, dict):
                    artifacts.append(payload)
            except Exception as exc:
                self.logger.warning(f"读取论文工件失败，已跳过 {path}: {exc}")
        return artifacts

    def _load_literature_map_for_citation_resolution(self) -> Optional[Dict[str, Any]]:
        """Load Outline v2 literature map as an alias source for citation tokens."""
        if not self.job_workspace or not self.project_name:
            return None
        path = self.job_workspace.artifact_path(f"{self.project_name}_literature_map.json")
        if not os.path.exists(path):
            return None
        try:
            with open(path, "r", encoding="utf-8") as handle:
                payload = json.load(handle)
            return payload if isinstance(payload, dict) else None
        except Exception as exc:
            self.logger.warning(f"读取 literature_map 以解析引用别名失败，已跳过 {path}: {exc}")
            return None

    def adopt_outline(self) -> bool:
        """执行大纲采纳操作"""
        try:
            if self._outline_v2_enabled():
                self.logger.error("当前启用了 Outline Intelligence v2；请使用 --adopt-outline-v2。")
                return False
            self.logger.info("开始执行大纲采纳操作...")
            
            # 加载大纲文件（使用 JSON 格式的 OutlineDocument）
            if not self.job_workspace or not self.project_name:
                self.logger.error("工作空间或项目名称未设置")
                return False
            
            outline_doc_path = self.job_workspace.artifact_path(f"{self.project_name}_outline_document.json")
            if not os.path.exists(outline_doc_path):
                self.logger.error(f"大纲文档文件不存在: {outline_doc_path}")
                return False
            
            # 加载大纲数据
            with open(outline_doc_path, 'r', encoding='utf-8') as f:
                outline_data = json.load(f)
            
            # 转换为 OutlineDocument 对象
            from outline.models import OutlineDocument
            outline = OutlineDocument.from_dict(outline_data)
            
            # 检查大纲是否有 critiques
            if not outline.critiques:
                self.logger.warning("大纲没有 critiques，无法执行采纳操作")
                return False
            
            # 加载仲裁结果（如果存在）
            arbitration_path = self.job_workspace.artifact_path(f"{self.project_name}_outline_arbitration.json")
            if not os.path.exists(arbitration_path):
                self.logger.error(f"仲裁结果文件不存在: {arbitration_path}")
                return False
            
            with open(arbitration_path, 'r', encoding='utf-8') as f:
                arbitration_data = json.load(f)
            
            # 转换为 OutlineArbitrationResult 对象
            from outline.models import CritiqueArbitration, ArbitrationDecision
            arbitrations = [
                CritiqueArbitration(
                    critique_id=arb['critique_id'],
                    decision=ArbitrationDecision(arb['decision']),
                    reason=arb['reason'],
                    arbitrated_at=arb['arbitrated_at'],
                    arbitrated_by=arb['arbitrated_by']
                )
                for arb in arbitration_data.get('arbitrations', [])
            ]
            
            from outline.arbitration import run_arbitration
            arbitration_result = run_arbitration(
                outline=outline,
                critiques=outline.critiques,
                arbitrations=arbitrations,
                job_id=self.job_workspace.job_id,
                arbitrated_by="user"
            )
            
            # 执行采纳操作
            from outline.arbitration import adopt_outline
            reviewed_outline = adopt_outline(
                outline=outline,
                arbitration_result=arbitration_result,
                job_id=self.job_workspace.job_id,
                adopted_by="user"
            )
            
            # 保存采纳后的大纲（使用正确的文件名）
            reviewed_outline_path = self.job_workspace.artifact_path(f"{self.project_name}_reviewed_outline.json")
            from services.job_workspace import atomic_write_json
            atomic_write_json(reviewed_outline_path, reviewed_outline.to_dict())

            if self.artifact_registry:
                depends_on = []
                if os.path.exists(outline_doc_path):
                    depends_on.append(
                        ArtifactDependencyRef(artifact_type="outline_document", path=outline_doc_path)
                    )
                if os.path.exists(arbitration_path):
                    depends_on.append(
                        ArtifactDependencyRef(artifact_type="outline_arbitration", path=arbitration_path)
                    )
                self.artifact_registry.register_file(
                    artifact_role="reviewed_outline",
                    artifact_type="reviewed_outline_document",
                    artifact_version="v1",
                    path=reviewed_outline_path,
                    producer="main.LiteratureReviewGenerator.adopt_outline",
                    depends_on=depends_on,
                    artifact_id=f"reviewed_outline:{self.project_name}",
                )

            self.logger.success(f"大纲采纳成功，已保存到: {reviewed_outline_path}")
            return True
        except Exception as e:
            self.logger.error(f"执行大纲采纳操作时出错: {e}")
            import traceback
            self.logger.debug(f"详细错误信息: {traceback.format_exc()}")
            return False

    def adopt_outline_v2(self, adopted_by: str = "user") -> bool:
        """Explicitly adopt v2 final_outline.json after a passing, current audit."""
        try:
            self.logger.info("开始执行 Outline Intelligence v2 显式采纳...")
            if not self.job_workspace or not self.project_name:
                self.logger.error("工作空间或项目名称未设置")
                return False

            from outline.v2_models import CoverageAudit, FinalOutline
            from outline.adoption import adopt_final_outline, write_adopted_outline

            final_path = self.job_workspace.artifact_path(f"{self.project_name}_final_outline.json")
            audit_path = self.job_workspace.artifact_path(f"{self.project_name}_outline_coverage_audit.json")
            if not os.path.exists(final_path):
                self.logger.error(f"final_outline.json 不存在: {final_path}")
                return False
            if not os.path.exists(audit_path):
                self.logger.error(f"outline_coverage_audit.json 不存在: {audit_path}")
                return False

            with open(final_path, "r", encoding="utf-8") as handle:
                final_outline = FinalOutline.from_dict(json.load(handle))
            with open(audit_path, "r", encoding="utf-8") as handle:
                audit = CoverageAudit.from_dict(json.load(handle))

            adopted, message = adopt_final_outline(
                final_outline=final_outline,
                audit=audit,
                job_id=self.job_workspace.job_id,
                adopted_by=adopted_by,
            )
            if adopted is None:
                self.logger.error(f"Outline v2 采纳失败: {message}")
                return False

            adopted_path = self.job_workspace.artifact_path(f"{self.project_name}_adopted_final_outline.json")
            write_adopted_outline(adopted, adopted_path)

            if self.artifact_registry:
                from services.artifact_registry import file_sha256

                depends_on = [
                    ArtifactDependencyRef(
                        artifact_type="final_outline",
                        path=final_path,
                        content_hash=file_sha256(final_path),
                    ),
                    ArtifactDependencyRef(
                        artifact_type="outline_coverage_audit",
                        path=audit_path,
                        content_hash=file_sha256(audit_path),
                    ),
                ]
                self.artifact_registry.register_file(
                    artifact_role="adopted_final_outline",
                    artifact_type="adopted_final_outline",
                    artifact_version="v1",
                    path=adopted_path,
                    producer="main.LiteratureReviewGenerator.adopt_outline_v2",
                    depends_on=depends_on,
                    artifact_id="adopted_final_outline",
                )

            self.logger.success(f"Outline v2 采纳成功，已保存到: {adopted_path}")
            return True
        except Exception as e:
            self.logger.error(f"执行 Outline v2 采纳操作时出错: {e}")
            import traceback
            self.logger.debug(f"详细错误信息: {traceback.format_exc()}")
            return False

    def _check_cancelled(self) -> None:
        if self.cancel_token is not None:
            self.cancel_token.check_cancelled()

    def _register_workspace_artifact(
        self,
        *,
        artifact_role: str,
        artifact_type: str,
        artifact_version: str,
        path: str,
        producer: str = "main.LiteratureReviewGenerator",
        depends_on: Optional[List[ArtifactDependencyRef]] = None,
    ) -> None:
        if not self.artifact_registry:
            return
        self.artifact_registry.register_file(
            artifact_role=artifact_role,
            artifact_type=artifact_type,
            artifact_version=artifact_version,
            path=path,
            producer=producer,
            depends_on=depends_on or [],
        )

    def _write_stage1_progress_snapshot(self) -> bool:
        if not self.job_workspace or not self.project_name or not self.summary_file:
            return False

        snapshot_path = self.job_workspace.artifact_path("stage1_progress_snapshot.json")
        snapshot = Stage1ProgressSnapshot(
            artifact_type="stage1_progress_snapshot",
            artifact_version="v1",
            created_from_job_id=self.job_workspace.job_id,
            created_at=datetime.utcnow().replace(microsecond=0).isoformat() + "Z",
            project_name=self.project_name,
            job_id=self.job_workspace.job_id,
            summary_file=self.summary_file,
            summary_count=len(self.summaries),
            processed_papers=sorted(self._checkpoint_processed_papers),
            failed_papers=sorted(self._checkpoint_failed_papers),
            fingerprint_bundle=dict(self.job_fingerprint_bundle),
            checkpoint_file=self._get_stage1_checkpoint_file_path(),
        )
        write_stage1_progress_snapshot(snapshot_path, snapshot)
        self._register_workspace_artifact(
            artifact_role="progress",
            artifact_type="stage1_progress_snapshot",
            artifact_version="v1",
            path=snapshot_path,
            depends_on=[
                ArtifactDependencyRef(
                    artifact_type="summary_file",
                    path=self.summary_file,
                )
            ],
        )
        return True

    @staticmethod
    def _paper_artifact_hash(paper_key: str) -> str:
        return hashlib.sha256(paper_key.encode("utf-8")).hexdigest()[:16]

    @staticmethod
    def _paper_artifact_key(paper: Mapping[str, Any]) -> str:
        explicit_key = str(paper.get("canonical_paper_key") or paper.get("source_paper_id") or "").strip()
        if explicit_key:
            return explicit_key
        legacy_paper: Dict[str, Any] = dict(paper)
        return LiteratureReviewGenerator.get_paper_key(legacy_paper)

    def _paper_artifact_id(self, paper: Mapping[str, Any]) -> str:
        paper_key = self._paper_artifact_key(paper)
        return f"paper_artifact:{self._paper_artifact_hash(paper_key)}"

    def _paper_artifact_path(self, paper: Mapping[str, Any]) -> str:
        if not self.job_workspace:
            raise ValueError("job workspace is not configured")
        artifact_hash = self._paper_artifact_hash(self._paper_artifact_key(paper))
        return self.job_workspace.artifact_path(f"paper_artifacts/{artifact_hash}.json")

    def _persist_paper_artifact(self, result: Mapping[str, Any]) -> bool:
        if result.get("status") != "success":
            return True
        if not self.job_workspace or not self.artifact_registry:
            return True

        paper = result.get("paper_info")
        if not isinstance(paper, Mapping):
            return True

        try:
            paper_key = self._paper_artifact_key(paper)
            paper_artifact = build_paper_artifact_v1(
                job_id=self.job_workspace.job_id,
                paper=paper,
                result=result,
                paper_key=paper_key,
            )
            artifact_path = self._paper_artifact_path(paper)
            atomic_write_json(artifact_path, paper_artifact.to_dict())

            depends_on: List[ArtifactDependencyRef] = []
            source_pdf = str(paper_artifact.source.get("source_pdf") or "")
            if source_pdf:
                depends_on.append(
                    ArtifactDependencyRef(
                        artifact_type="source_pdf",
                        path=source_pdf,
                        content_hash=str(paper_artifact.source.get("source_pdf_fingerprint") or ""),
                    )
                )
            visual_manifest_path = str(paper_artifact.stage1_inputs.get("visual_artifact_manifest_path") or "")
            if visual_manifest_path:
                depends_on.append(
                    ArtifactDependencyRef(
                        artifact_type="visual_manifest",
                        path=visual_manifest_path,
                    )
                )

            self.artifact_registry.register_file(
                artifact_role=self.PAPER_ARTIFACT_ROLE,
                artifact_type=self.PAPER_ARTIFACT_TYPE,
                artifact_version=self.PAPER_ARTIFACT_VERSION,
                path=artifact_path,
                producer="main.LiteratureReviewGenerator.process_paper",
                depends_on=depends_on,
                artifact_id=self._paper_artifact_id(paper),
            )
            return True
        except Exception as exc:
            self.logger.error(f"Failed to persist paper_artifact_v1: {exc}")
            return False

    def _review_draft_path(self) -> str:
        if not self.job_workspace:
            raise ValueError("job workspace is not configured")
        project_name = self.project_name or "review"
        return self.job_workspace.artifact_path(f"review_drafts/{project_name}_review_draft_v1.json")

    def _review_draft_v2_path(self) -> str:
        if not self.job_workspace:
            raise ValueError("job workspace is not configured")
        project_name = self.project_name or "review"
        return self.job_workspace.artifact_path(f"review_drafts/{project_name}_review_draft_v2.json")

    def _citation_manifest_path(self) -> str:
        if not self.job_workspace:
            raise ValueError("job workspace is not configured")
        project_name = self.project_name or "review"
        return self.job_workspace.artifact_path(f"citation_manifests/{project_name}_citation_manifest_v3.json")

    def _citation_manifest_v2_path(self) -> str:
        if not self.job_workspace:
            raise ValueError("job workspace is not configured")
        project_name = self.project_name or "review"
        return self.job_workspace.artifact_path(f"citation_manifests/{project_name}_citation_manifest_v2.json")

    def _citation_manifest_v1_path(self) -> str:
        """Legacy v1 path for compatibility projection only."""
        if not self.job_workspace:
            raise ValueError("job workspace is not configured")
        project_name = self.project_name or "review"
        return self.job_workspace.artifact_path(f"citation_manifests/{project_name}_citation_manifest_v1.json")

    def _citation_migration_report_path(self) -> str:
        if not self.job_workspace:
            raise ValueError("job workspace is not configured")
        project_name = self.project_name or "review"
        return self.job_workspace.artifact_path(f"citation_manifests/{project_name}_citation_migration_report.json")

    def _extract_review_sections_from_word_document(
        self,
        word_file: str,
        *,
        section_titles_by_number: Mapping[int, str],
    ) -> Dict[int, Dict[str, Any]]:
        if not os.path.exists(word_file) or not DOCX_AVAILABLE or Document is None:
            return {}

        try:
            document = Document(word_file)  # type: ignore[operator]
            section_map: Dict[int, Dict[str, Any]] = {}
            current_section_number: Optional[int] = None
            current_paragraphs: List[str] = []

            def _flush_current_section() -> None:
                nonlocal current_section_number, current_paragraphs
                if current_section_number is None:
                    return
                section_map[current_section_number] = {
                    "section_number": current_section_number,
                    "section_title": section_titles_by_number.get(current_section_number, ""),
                    "content": "\n\n".join(current_paragraphs).strip(),
                }
                current_section_number = None
                current_paragraphs = []

            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "")

                if style_name.startswith("Heading 1"):
                    _flush_current_section()
                    heading_match = re.search(r"(\d+)", text)
                    if heading_match:
                        candidate = int(heading_match.group(1))
                        if candidate in section_titles_by_number:
                            current_section_number = candidate
                    continue

                if current_section_number is not None and text:
                    current_paragraphs.append(text)

            _flush_current_section()
            return section_map
        except Exception as exc:
            self.logger.warning(f"Failed to extract existing review sections from Word document: {exc}")
            return {}

    def _persist_review_draft(
        self,
        *,
        outline_file: str,
        review_sections: List[Dict[str, Any]],
        references: List[str],
        word_file: str,
        generation_mode: str = "full_review",
    ) -> bool:
        if not self.job_workspace or not self.artifact_registry:
            return True

        try:
            review_draft = build_review_draft_v1(
                job_id=self.job_workspace.job_id,
                project_name=self.project_name or "review",
                draft_id=self.REVIEW_DRAFT_ARTIFACT_ID,
                outline_artifact_id=self.OUTLINE_ARTIFACT_ID,
                outline_source_path=outline_file,
                summary_file=self.summary_file or "",
                review_word_path=word_file,
                sections=review_sections,
                references=references,
                generation_mode=generation_mode,
            )
            artifact_path = self._review_draft_path()
            atomic_write_json(artifact_path, review_draft.to_dict())

            depends_on: List[ArtifactDependencyRef] = []
            if outline_file:
                depends_on.append(
                    ArtifactDependencyRef(
                        artifact_type=self.OUTLINE_ARTIFACT_TYPE,
                        path=outline_file,
                    )
                )
            if self.summary_file:
                depends_on.append(
                    ArtifactDependencyRef(
                        artifact_type="summary_file",
                        path=self.summary_file,
                    )
                )

            self.artifact_registry.register_file(
                artifact_role=self.REVIEW_DRAFT_ARTIFACT_ROLE,
                artifact_type=self.REVIEW_DRAFT_ARTIFACT_TYPE,
                artifact_version=self.REVIEW_DRAFT_ARTIFACT_VERSION,
                path=artifact_path,
                producer="main.LiteratureReviewGenerator.generate_full_review_from_outline",
                depends_on=depends_on,
                artifact_id=self.REVIEW_DRAFT_ARTIFACT_ID,
            )
            return True
        except Exception as exc:
            self.logger.error(f"Failed to persist review_draft_v1: {exc}")
            return False

    def _persist_review_draft_v2(
        self,
        *,
        outline_file: str,
        review_sections: List[Dict[str, Any]],
        references: List[str],
        word_file: str,
        generation_mode: str = "full_review",
        paper_summaries: Optional[List[Dict[str, Any]]] = None,
    ) -> bool:
        if not self.job_workspace or not self.artifact_registry:
            return True

        try:
            # Use self.summaries if paper_summaries is not explicitly provided
            if paper_summaries is None:
                paper_summaries = [dict(summary) for summary in self.summaries]
            
            review_draft = build_review_draft_v2(
                job_id=self.job_workspace.job_id,
                project_name=self.project_name or "review",
                draft_id=self.REVIEW_DRAFT_V2_ARTIFACT_ID,
                outline_artifact_id=self.OUTLINE_ARTIFACT_ID,
                outline_source_path=outline_file,
                summary_file=self.summary_file or "",
                review_word_path=word_file,
                sections=review_sections,
                references=references,
                generation_mode=generation_mode,
                paper_summaries=paper_summaries,
                allow_legacy_regex_citations=False,
            )
            artifact_path = self._review_draft_v2_path()
            atomic_write_json(artifact_path, review_draft.to_dict())

            depends_on: List[ArtifactDependencyRef] = []
            if outline_file:
                depends_on.append(
                    ArtifactDependencyRef(
                        artifact_type=self.OUTLINE_ARTIFACT_TYPE,
                        path=outline_file,
                    )
                )
            if self.summary_file:
                depends_on.append(
                    ArtifactDependencyRef(
                        artifact_type="summary_file",
                        path=self.summary_file,
                    )
                )

            self.artifact_registry.register_file(
                artifact_role=self.REVIEW_DRAFT_V2_ARTIFACT_ROLE,
                artifact_type=self.REVIEW_DRAFT_V2_ARTIFACT_TYPE,
                artifact_version=self.REVIEW_DRAFT_V2_ARTIFACT_VERSION,
                path=artifact_path,
                producer="main.LiteratureReviewGenerator.generate_full_review_from_outline",
                depends_on=depends_on,
                artifact_id=self.REVIEW_DRAFT_V2_ARTIFACT_ID,
            )
            return True
        except Exception as exc:
            self.logger.error(f"Failed to persist review_draft_v2: {exc}")
            return False

    def _build_citation_manifest_v3_from_review_draft(
        self,
        review_draft_path: str,
        review_word_path: str,
    ) -> Any:
        """Build canonical CitationManifestV3 from review_draft_v2 block structure."""
        from services.citation_manifest import build_citation_manifest_v3_from_review_draft
        
        # Load review_draft_v2 to get block structure
        try:
            with open(review_draft_path, 'r', encoding='utf-8') as f:
                review_draft_v2 = json.load(f)
        except Exception as e:
            self.logger.error(f"Failed to load review_draft_v2: {e}")
            raise
        
        return build_citation_manifest_v3_from_review_draft(
            job_id=self.job_workspace.job_id if self.job_workspace else "unknown",
            project_name=self.project_name or "review",
            manifest_id=self.CITATION_MANIFEST_ARTIFACT_ID,
            review_draft_path=review_draft_path,
            review_word_path=review_word_path,
            review_draft_v2=review_draft_v2,
            paper_summaries=[dict(summary) for summary in self.summaries],
            literature_map=self._load_literature_map_for_citation_resolution(),
        )

    def _persist_citation_manifest(
        self,
        *, 
        review_draft_path: str,
        review_word_path: str,
        citations: Optional[list[dict[str, Any]]] = None,
    ) -> bool:
        if not self.job_workspace or not self.artifact_registry:
            return True

        try:
            citation_manifest_v3 = self._build_citation_manifest_v3_from_review_draft(
                review_draft_path=review_draft_path,
                review_word_path=review_word_path,
            )
            
            artifact_path_v3 = self._citation_manifest_path()
            atomic_write_json(artifact_path_v3, citation_manifest_v3.to_dict())
            atomic_write_json(self._citation_migration_report_path(), citation_manifest_v3.migration_report.to_dict())

            depends_on: List[ArtifactDependencyRef] = []
            if review_draft_path:
                depends_on.append(
                    ArtifactDependencyRef(
                        artifact_type=self.REVIEW_DRAFT_V2_ARTIFACT_TYPE,
                        path=review_draft_path,
                    )
                )

            self.artifact_registry.register_file(
                artifact_role=self.CITATION_MANIFEST_ARTIFACT_ROLE,
                artifact_type=self.CITATION_MANIFEST_ARTIFACT_TYPE,
                artifact_version="v3",
                path=artifact_path_v3,
                producer="main.LiteratureReviewGenerator.generate_full_review_from_outline",
                depends_on=depends_on,
                artifact_id=self.CITATION_MANIFEST_ARTIFACT_ID,
            )

            artifact_path_v1 = self._citation_manifest_v1_path()
            v1_compatible = {
                "artifact_type": "citation_manifest",
                "artifact_version": "v1",
                "created_from_job_id": citation_manifest_v3.created_from_job_id,
                "created_at": citation_manifest_v3.created_at,
                "manifest_identity": {
                    **citation_manifest_v3.manifest_identity,
                    "projection_from": "v3",
                },
                "review_reference": citation_manifest_v3.review_reference,
                "citations": [
                    {
                        "citation_id": occ.occurrence_id,
                        "paper_id": occ.paper_id,
                        "text": occ.citation_token,
                        "context": occ.context_before,
                        "section_number": occ.section_number,
                        "section_title": occ.section_title,
                        "block_id": occ.block_id,
                        "block_order": occ.block_order,
                        "review_draft_version": "v2",
                    }
                    for occ in citation_manifest_v3.occurrences
                ],
            }
            atomic_write_json(artifact_path_v1, v1_compatible)
            
            return True
        except Exception as exc:
            self.logger.error(f"Failed to persist citation_manifest: {exc}")
            return False

    def _init_logger(self):
        """初始化日志记录器"""
        import logging
        import os
        from datetime import datetime
        
        # 创建日志记录器
        self.logger = logging.getLogger(f"auto-generate_{datetime.now().strftime('%Y%m%d_%H%M%S')}")  # type: ignore
        self.logger.setLevel(logging.INFO)
        
        # 如果记录器已经有处理器，先清除
        if self.logger.handlers:
            self.logger.handlers.clear()
        
        # 创建控制台处理器
        console_handler = logging.StreamHandler(sys.stdout)
        console_handler.setLevel(logging.INFO)
        
        # 创建格式器
        formatter = logging.Formatter('[%(asctime)s] [%(levelname)s] %(message)s', 
                                    datefmt='%H:%M:%S')
        console_handler.setFormatter(formatter)
        
        # 创建文件处理器
        try:
            # 创建logs目录（如果不存在）
            logs_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logs')
            os.makedirs(logs_dir, exist_ok=True)
            
            # 生成日志文件名：使用时间戳确保唯一性
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            log_file = os.path.join(logs_dir, f'llm_reviewer_{timestamp}.log')
            self.root_log_path = log_file

            file_handler = logging.FileHandler(log_file, encoding='utf-8')
            file_handler.setLevel(logging.INFO)
            file_handler.setFormatter(formatter)
            
            # 添加处理器到记录器
            self.logger.addHandler(console_handler)
            self.logger.addHandler(file_handler)
            
            # 记录日志文件位置
            self.logger.info(f"日志文件已创建: {log_file}")
            
        except Exception as e:
            # 如果文件日志失败，只使用控制台日志
            self.logger.warning(f"无法创建文件日志，仅使用控制台日志: {e}")
            self.logger.addHandler(console_handler)
    
    def load_configuration(self) -> bool:
        """加载配置文件"""
        try:
            self.config = load_config(self.config_file)
            if not self.config:
                self.logger.error("配置文件加载失败或为空")
                return False
            self.compat_config = CompatConfigView.from_config(self.config)
            self.preprocess_manager = PreprocessManager(config=self.config, logger=self.logger)
            self.logger.success("配置文件加载成功")
            return True
        except Exception as e:
            self.logger.error(f"配置文件加载异常: {e}")
            return False
    
    def setup_output_directory(self) -> bool:
        """设置输出目录"""
        try:
            # 检查配置是否已加载
            if not self.config:
                self.logger.error("配置未加载，无法设置输出目录")
                return False
            
            # 确定项目名称
            if not self.project_name:
                if self.mode == "zotero":
                    # Zotero模式使用默认项目名
                    self.project_name = "literature_review"
                else:
                    # 直接PDF模式使用文件夹名作为项目名
                    self.project_name = os.path.basename((self.pdf_folder or '').rstrip('/\\'))
            
            # 清理项目名称，移除非法字符
            self.project_name = sanitize_path_component(self.project_name)
            
            # 确定输出路径
            paths_config: Dict[str, str] = self.config.get('Paths', {}) if self.config else {}
            output_base_path: str = paths_config.get('output_path', './output')
            if self.job_workspace is None:
                pointer_path = os.path.join(os.path.abspath(output_base_path), self.project_name, "_latest_job.json")
                if os.path.exists(pointer_path):
                    try:
                        with open(pointer_path, 'r', encoding='utf-8') as handle:
                            pointer_payload = json.load(handle)
                        workspace_path = str(pointer_payload.get("workspace_path", "") or "")
                        job_id = str(pointer_payload.get("job_id", "") or "")
                        if workspace_path and os.path.exists(workspace_path):
                            self.job_workspace = JobWorkspace.from_workspace_path(
                                workspace_path=workspace_path,
                                project_name=self.project_name,
                                job_id=job_id or None,
                            )
                    except Exception as exc:
                        self.logger.warning(f"读取 latest job pointer 失败，将创建新 workspace: {exc}")
                if self.job_workspace is None:
                    self.job_workspace = JobWorkspace.create(
                        base_output_dir=output_base_path,
                        project_name=self.project_name,
                    )

            self.job_workspace.ensure_exists()
            if self.artifact_registry is None:
                self.artifact_registry = ArtifactRegistry(
                    self.job_workspace.paths.registry_path,
                    self.job_workspace.job_id,
                )
            self.output_dir = self.job_workspace.root_dir
            self.summary_file = self._get_summary_file_path()
            
            # 确保输出目录存在
            if ensure_dir(self.output_dir):
                self.logger.success(f"输出目录已创建: {self.output_dir}")
            else:
                self.logger.error(f"无法创建输出目录: {self.output_dir}")
                return False
            
            return True
        except Exception as e:
            self.logger.error(f"设置输出目录失败: {e}")
            return False
    
    def scan_pdf_folder(self) -> bool:
        """扫描PDF文件夹（直接模式专用）"""
        try:
            if self.mode != "direct":
                self.logger.error("scan_pdf_folder只能在直接PDF模式下调用")
                return False
            
            if not self.pdf_folder or not os.path.exists(self.pdf_folder):
                self.logger.error(f"PDF文件夹不存在: {self.pdf_folder}")
                return False
            
            self.logger.info(f"正在扫描PDF文件夹: {self.pdf_folder}")
            
            # 查找所有PDF文件
            pdf_files: List[str] = []
            for root, _dirs, files in os.walk(self.pdf_folder):
                for file in files:
                    if file.lower().endswith('.pdf'):
                        pdf_files.append(os.path.join(root, file))
            
            self.logger.info(f"找到 {len(pdf_files)} 个PDF文件")
            
            # 为每个PDF文件创建论文信息
            self.papers: List[PaperInfo] = []
            for i, pdf_path in enumerate(pdf_files):
                # 从文件名提取标题（移除.pdf扩展名）
                title = os.path.splitext(os.path.basename(pdf_path))[0]
                
                # 尝试从PDF文件中提取额外信息
                pdf_info: Optional[Dict[str, str]] = None  # 明确指定字典值的类型
                try:
                    from pdf_extractor import get_pdf_info  # type: ignore
                    pdf_info = get_pdf_info(pdf_path)
                except Exception as e:
                    self.logger.warning(f"无法从PDF文件提取元数据: {pdf_path}, 错误: {e}")
                
                # 创建论文信息字典
                paper_info: PaperInfo = {
                    'title': title,
                    'authors': [],  # 初始化为空列表
                    'year': '未知年份',  # 年份通常需要OCR才能从PDF中提取，暂时设为默认值
                    'journal': '未知期刊',  # 期刊信息通常需要OCR才能从PDF中提取，暂时设为默认值
                    'doi': '',  # 直接模式下DOI为空
                    'pdf_path': pdf_path,  # PDF文件路径
                    'file_index': i  # 文件索引
                }
                
                # 从PDF信息中提取作者
                if pdf_info:
                    author_str = pdf_info.get('author', '')
                    if author_str and author_str.strip():
                        # 将作者字符串转换为列表格式
                        paper_info['authors'] = [author_str.strip()]
                
                # 如果PDF信息为空或无作者，设为空数组
                authors = paper_info.get('authors', [])
                if not authors:
                    paper_info['authors'] = []
                elif any(author.strip() in ['Unknown', '未知'] for author in authors):
                    paper_info['authors'] = []
                
                # 尝试从文件名中提取年份（简单模式匹配）
                import re
                year_match = re.search(r'(20\d{2})', title)  # 搜索2020-2099年份
                if year_match:
                    paper_info['year'] = year_match.group(1)
                
                # 尝试从文件名中提取作者（如果文件名格式包含下划线分隔的作者名）
                if '_' in title:
                    # 假设文件名格式为: "标题_作者.pdf" 或 "标题_作者_其他信息.pdf"
                    parts = title.split('_')
                    if len(parts) >= 2:
                        potential_author = parts[-1].strip()
                        if potential_author and potential_author != '侯甜甜' and potential_author != '贺爱忠' and potential_author != '周冲' and potential_author != '盘城' and potential_author != '张赛楠' and potential_author != '彭丽徽' and potential_author != '康超' and potential_author != '刘伟华' and potential_author != '朱华东':
                            paper_info['authors'] = [potential_author]
                
                self.papers.append(paper_info)
            
            descriptors = normalize_source_papers("direct", self.papers)
            self.source_descriptors = [descriptor.to_dict() for descriptor in descriptors]
            self.papers = project_descriptors_to_legacy_papers(self.papers, descriptors)
            self.logger.success(f"PDF文件夹扫描完成，共 {len(self.papers)} 篇论文")
            return True
            
        except Exception as e:
            self.logger.error(f"扫描PDF文件夹失败: {e}")
            return False
    
    def parse_zotero_report(self, override_path: Optional[str] = None) -> bool:
        """解析Zotero报告（Zotero模式专用）"""
        try:
            if self.mode != "zotero":
                self.logger.error("parse_zotero_report只能在Zotero模式下调用")
                return False
            
            # 确定Zotero报告路径
            if override_path:
                zotero_report_path = override_path
            elif self.zotero_report:
                zotero_report_path = self.zotero_report
            else:
                paths_config: Dict[str, str] = self.config.get('Paths', {}) if self.config else {}
                zotero_report_path: str = paths_config.get('zotero_report', '')
            
            if not zotero_report_path or not os.path.exists(zotero_report_path):
                self.logger.error(f"Zotero报告文件不存在: {zotero_report_path}")
                return False
            
            self.logger.info(f"正在解析Zotero报告: {zotero_report_path}")
            
            # 解析报告
            self.papers = parse_zotero_report(zotero_report_path)
            
            if not self.papers:
                self.logger.error("Zotero报告解析失败或报告为空")
                return False
            
            descriptors = normalize_source_papers("zotero", self.papers)
            self.source_descriptors = [descriptor.to_dict() for descriptor in descriptors]
            self.papers = project_descriptors_to_legacy_papers(self.papers, descriptors)
            self.logger.success(f"Zotero报告解析完成，共 {len(self.papers)} 篇论文")
            return True
            
        except Exception as e:
            self.logger.error(f"解析Zotero报告失败: {e}")
            return False
    
    @staticmethod
    def get_paper_key(paper: 'Dict[str, Any] | PaperInfo') -> str:
        """为论文生成唯一身份标识（调用模块级别函数）"""
        return get_paper_key(paper)
    
    def _load_summary_records_from_path(self, path: str) -> List[Dict[str, Any]]:
        return load_summary_records(path, logger=self.logger)

    def _explicit_summary_source_paths(self) -> List[str]:
        paths: List[str] = []
        seen: Set[str] = set()
        for raw_path in [self.summary_file_override, *self.summary_source_overrides]:
            value = str(raw_path or "").strip()
            if not value:
                continue
            normalized = os.path.abspath(value)
            if normalized in seen:
                continue
            seen.add(normalized)
            paths.append(normalized)
        return paths

    def _materialize_effective_summaries(
        self,
        summaries: List[Dict[str, Any]],
        *,
        source_path: str = "",
        source_kind: str,
        producer: str,
        source_items: Optional[List[Dict[str, Any]]] = None,
        rejected_candidates: Optional[List[Dict[str, Any]]] = None,
    ) -> bool:
        if not self.summary_file:
            return False

        Path(self.summary_file).parent.mkdir(parents=True, exist_ok=True)
        atomic_write_json(self.summary_file, summaries)
        effective_source_items = list(source_items or [])
        if not effective_source_items and source_path:
            effective_source_items = [
                {
                    "path": os.path.abspath(source_path),
                    "source_type": "explicit",
                    "label": source_kind,
                    "priority": 0,
                }
            ]
        self._register_workspace_artifact(
            artifact_role="summary",
            artifact_type="summary_file",
            artifact_version="v1",
            path=self.summary_file,
            producer=producer,
            depends_on=[
                ArtifactDependencyRef(
                    artifact_type="summary_source",
                    path=str(item.get("path") or ""),
                )
                for item in effective_source_items
                if str(item.get("path") or "").strip()
            ],
        )

        manifest_path = self._get_summary_source_manifest_path()
        primary_source_path = ""
        if source_path:
            primary_source_path = os.path.abspath(source_path)
        elif effective_source_items:
            primary_source_path = str(effective_source_items[0].get("path") or "")
        manifest_payload = {
            "artifact_type": "summary_source_manifest",
            "artifact_version": "v2",
            "created_at": datetime.utcnow().replace(microsecond=0).isoformat() + "Z",
            "project_name": self.project_name or "",
            "source_kind": source_kind,
            "source_path": primary_source_path,
            "source_items": effective_source_items,
            "rejected_candidates": list(rejected_candidates or []),
            "materialized_summary_file": self.summary_file,
            "summary_count": len(summaries),
        }
        atomic_write_json(manifest_path, manifest_payload)
        self._register_workspace_artifact(
            artifact_role="summary_source",
            artifact_type="summary_source_manifest",
            artifact_version="v1",
            path=manifest_path,
            producer=producer,
            depends_on=[
                ArtifactDependencyRef(
                    artifact_type="summary_file",
                    path=self.summary_file,
                )
            ],
        )
        return True

    def _load_summaries_from_sources(
        self,
        paths: Sequence[str],
        *,
        source_kind: str = "explicit_summary_sources",
        producer: str = "main.LiteratureReviewGenerator.load_existing_summaries",
    ) -> bool:
        sources = [
            SummarySource(
                path=os.path.abspath(path),
                source_type="explicit",
                priority=index,
                label=f"explicit:{index + 1}",
            )
            for index, path in enumerate(paths)
        ]
        resolved: ResolvedSummarySet = build_effective_summary_set(sources, logger=self.logger)
        self.summaries = [dict(item) for item in resolved.summaries]
        success_count = len(self.summaries)
        self.logger.success(f"Loaded summary sources: {success_count} success, 0 failed")
        if resolved.rejected_candidates:
            self.logger.info(f"Summary source merge skipped {len(resolved.rejected_candidates)} non-reusable records")
        self._materialize_effective_summaries(
            self.summaries,
            source_path=str(paths[0]) if len(paths) == 1 else "",
            source_kind=source_kind,
            producer=producer,
            source_items=resolved.source_items,
            rejected_candidates=resolved.rejected_candidates,
        )
        return True

    def _load_summaries_from_path(
        self,
        path: str,
        *,
        materialize: bool = False,
        source_kind: str = "existing_summary_file",
        producer: str = "main.LiteratureReviewGenerator.load_existing_summaries",
    ) -> bool:
        loaded_summaries = self._load_summary_records_from_path(path)
        self.summaries = [dict(item) for item in loaded_summaries]

        success_count = len([s for s in self.summaries if s.get("status") == "success"])
        failed_count = len([s for s in self.summaries if s.get("status") == "failed"])
        self.logger.success(f"Loaded summary file: {success_count} success, {failed_count} failed")

        if materialize:
            self._materialize_effective_summaries(
                self.summaries,
                source_path=path,
                source_kind=source_kind,
                producer=producer,
                source_items=[
                    {
                        "path": os.path.abspath(path),
                        "source_type": "explicit",
                        "label": source_kind,
                        "priority": 0,
                    }
                ],
            )
        return True

    def load_existing_summaries(self) -> bool:
        """Load an existing summaries file for resume or downstream generation."""
        try:
            explicit_summary_sources = self._explicit_summary_source_paths()
            if explicit_summary_sources:
                source_kind = "explicit_summary_file" if len(explicit_summary_sources) == 1 else "explicit_summary_sources"
                return self._load_summaries_from_sources(
                    explicit_summary_sources,
                    source_kind=source_kind,
                )

            if self.summary_file and os.path.exists(self.summary_file):
                try:
                    return self._load_summaries_from_path(self.summary_file)
                except SummarySourceError as exc:
                    self.logger.warning(f"Current summary file is invalid; starting fresh instead: {exc}")
                    self.summaries = []

            if self.project_name:
                paths_config: Dict[str, str] = self.config.get('Paths', {}) if self.config else {}
                output_base_path: str = paths_config.get('output_path', './output')
                output_base_path_abs = os.path.abspath(output_base_path)

                import glob
                workspace_pattern = os.path.join(output_base_path_abs, f"{self.project_name}__*")
                workspaces = glob.glob(workspace_pattern)
                workspaces.sort(key=os.path.getmtime, reverse=True)

                for workspace_path in workspaces:
                    summary_file = os.path.join(workspace_path, "artifacts", f"{self.project_name}_summaries.json")
                    if not os.path.exists(summary_file):
                        continue
                    try:
                        self.logger.info(f"Found historical summary file: {summary_file}")
                        return self._load_summaries_from_path(summary_file)
                    except SummarySourceError as exc:
                        self.logger.warning(f"Historical summary file is invalid; checking the next candidate: {exc}")
                        continue

            self.logger.info("No existing summary file was found; starting a fresh run")
            self.summaries = []
            return True

        except SummarySourceError as exc:
            self.logger.error(f"Failed to load summary file: {exc}")
            self.summaries = []
            return False
        except Exception as e:
            self.logger.warning(f"Failed to load existing summary file; starting a fresh run: {e}")
            self.summaries = []
            return True

    def _restore_stage1_progress_from_snapshot(self, snapshot_path: str) -> bool:
        """从 progress snapshot 中恢复阶段一跳过集合。"""
        try:
            snapshot = load_stage1_progress_snapshot(snapshot_path)
            if snapshot is None:
                return False
            if snapshot.project_name != self.project_name:
                self.logger.warning(
                    f"[进度恢复] snapshot 项目名称不匹配({snapshot.project_name} != {self.project_name})，忽略该快照"
                )
                return False

            processed_papers = {
                normalize_checkpoint_paper_key(paper_key)
                for paper_key in snapshot.processed_papers
                if isinstance(paper_key, str) and normalize_checkpoint_paper_key(paper_key)
            }
            failed_papers = {
                normalize_checkpoint_paper_key(paper_key)
                for paper_key in snapshot.failed_papers
                if isinstance(paper_key, str) and normalize_checkpoint_paper_key(paper_key)
            }

            self._checkpoint_processed_papers = processed_papers
            self._checkpoint_failed_papers = failed_papers
            self.processed_count.set(len(processed_papers))
            self.failed_count.set(len(failed_papers))

            self.logger.info(
                f"[进度恢复] 已从 progress snapshot 恢复: {len(processed_papers)}成功, {len(failed_papers)}失败"
            )
            return True
        except Exception as exc:
            self.logger.warning(f"[进度恢复] 读取 progress snapshot 失败，将继续尝试其他恢复路径: {exc}")
            return False

    def _rebuild_stage1_progress_from_loaded_summaries(self) -> bool:
        """在缺少 checkpoint 时，根据已加载摘要重建阶段一跳过集合。"""
        if not self.summaries:
            return False

        processed_papers, failed_papers = self._stage1_progress_sets_from_loaded_summaries()
        if not processed_papers and not failed_papers:
            return False

        self._checkpoint_processed_papers = processed_papers
        self._checkpoint_failed_papers = failed_papers
        self.processed_count.set(len(processed_papers))
        self.failed_count.set(len(failed_papers))

        self.logger.info(
            f"[进度恢复] 已根据现有摘要重建处理进度: {len(processed_papers)}成功, {len(failed_papers)}失败"
        )
        return True

    def _stage1_progress_sets_from_loaded_summaries(self) -> Tuple[Set[str], Set[str]]:
        """从已加载 summaries 提取阶段一成功/失败身份集合。"""
        processed_papers: Set[str] = set()
        failed_papers: Set[str] = set()

        for summary in self.summaries:
            if not isinstance(summary, Mapping):
                continue
            paper_info = summary.get("paper_info")
            if not isinstance(paper_info, Mapping):
                continue

            status = str(summary.get("status") or "").strip().lower()
            paper_key = LiteratureReviewGenerator.get_paper_key(dict(paper_info))
            if not paper_key:
                continue

            if status == "success":
                processed_papers.add(paper_key)
                failed_papers.discard(paper_key)
            elif status == "failed":
                if paper_key not in processed_papers:
                    failed_papers.add(paper_key)

        failed_papers.difference_update(processed_papers)
        return processed_papers, failed_papers

    def _merge_stage1_progress_from_loaded_summaries(self) -> bool:
        """用 summaries 中的 durable 结果校正旧 checkpoint，避免重试已成功论文。"""
        if not self.summaries:
            return False

        processed_papers, failed_papers = self._stage1_progress_sets_from_loaded_summaries()
        if not processed_papers and not failed_papers:
            return False

        old_processed = set(self._checkpoint_processed_papers)
        old_failed = set(self._checkpoint_failed_papers)

        self._checkpoint_processed_papers.update(processed_papers)
        self._checkpoint_failed_papers.difference_update(processed_papers)
        self._checkpoint_failed_papers.update(
            key for key in failed_papers if key not in self._checkpoint_processed_papers
        )
        self.processed_count.set(len(self._checkpoint_processed_papers))
        self.failed_count.set(len(self._checkpoint_failed_papers))

        if (
            old_processed != self._checkpoint_processed_papers
            or old_failed != self._checkpoint_failed_papers
        ):
            self.logger.info(
                "[进度恢复] 已用现有摘要校正 checkpoint: "
                f"{len(self._checkpoint_processed_papers)}成功, "
                f"{len(self._checkpoint_failed_papers)}失败"
            )

        return True
    
    def reset_counters(self):
        """重置计数器"""
        self.processed_count.set(0)
        self.failed_count.set(0)

    @staticmethod
    def _paper_progress_label(paper: Mapping[str, Any] | PaperInfo | None) -> str:
        title = str((paper or {}).get('title') or '').strip()
        return title or '未知标题'

    @staticmethod
    def _short_progress_message(message: str, max_length: int = 180) -> str:
        normalized = " ".join(str(message or "").split())
        if len(normalized) <= max_length:
            return normalized
        return normalized[: max_length - 1] + "…"

    def _emit_progress(self, **kwargs: Any) -> None:
        tracker = self.progress_tracker
        if tracker is not None:
            tracker.emit(**kwargs)

    def _emit_stage1_progress(
        self,
        *,
        total: int,
        current: int,
        success_count: int,
        failure_count: int,
        message: str,
        item_label: str = "",
        retry_round: int = 0,
        retry_total_rounds: int = 0,
    ) -> None:
        self._emit_progress(
            stage="analyze",
            total=total,
            current=current,
            success_count=success_count,
            failure_count=failure_count,
            remaining_count=max(total - current, 0),
            indeterminate=False,
            message=message,
            item_label=item_label,
            retry_round=retry_round,
            retry_total_rounds=retry_total_rounds,
        )

    def reset_stage1_reader_engine_round_state(self) -> None:
        with self._stage1_reader_engine_lock:
            self._stage1_disabled_reader_engines.clear()
            self._stage1_reader_disable_reasons.clear()

    def _is_stage1_reader_engine_disabled(self, engine_type: str) -> bool:
        with self._stage1_reader_engine_lock:
            return engine_type in self._stage1_disabled_reader_engines

    def _disable_stage1_reader_engine_for_round(self, engine_type: str, result: Mapping[str, Any]) -> None:
        label = "主引擎" if engine_type == "primary" else "备用引擎"
        message = f"{label}余额/额度不足，本轮自动跳过。"
        reason = str(result.get("message") or result.get("provider_code") or result.get("error_kind") or "")
        first_disable = False
        with self._stage1_reader_engine_lock:
            if engine_type not in self._stage1_disabled_reader_engines:
                self._stage1_disabled_reader_engines.add(engine_type)
                self._stage1_reader_disable_reasons[engine_type] = reason
                first_disable = True
        if first_disable:
            detail = f" 原因: {reason}" if reason else ""
            self.logger.warning(f"{message}{detail}")
            self._emit_progress(stage="analyze", message=message)

    def _call_stage1_reader_with_scheduler(
        self,
        analysis_prompt: str,
        reader_api_config: APIConfig,
        backup_api_config: APIConfig,
        *,
        user_content: Any = None,
        skip_engines: Optional[Set[str]] = None,
    ) -> Dict[str, Any]:
        result = get_summary_from_ai_with_fallback(
            analysis_prompt,
            reader_api_config,
            backup_api_config,
            logger=self.logger,
            config=self.config,
            user_content=user_content,
            return_detailed=True,
            disable_engine_callback=self._disable_stage1_reader_engine_for_round,
            is_engine_disabled_callback=self._is_stage1_reader_engine_disabled,
            skip_engines=skip_engines,
        )
        if isinstance(result, dict) and "status" in result and "content" in result:
            return result
        if result:
            legacy_engine = "backup" if skip_engines and "primary" in skip_engines else "primary"
            return {
                "status": "success",
                "error_kind": None,
                "message": "",
                "content": result,
                "engine_type": legacy_engine,
            }
        return {
            "status": "failed",
            "error_kind": "invalid_response",
            "message": "AI summary generation failed",
            "content": None,
            "engine_type": None,
        }

    def _collect_preprocess_result_metadata(self, preprocess_result: Any, input_kind: str) -> Dict[str, Any]:
        metadata = {
            'analysis_input_kind': input_kind,
            'extractor_used': getattr(preprocess_result, 'extractor_used', ''),
            'layout_fidelity': getattr(preprocess_result, 'layout_fidelity', ''),
            'conversion_used': getattr(preprocess_result, 'conversion_used', ''),
            'used_ocr': bool(getattr(preprocess_result, 'used_ocr', False)),
            'low_quality': bool(getattr(preprocess_result, 'low_quality', False)),
            'scanned_like': bool(getattr(preprocess_result, 'scanned_like', False)),
            'mineru_attempted': bool(getattr(preprocess_result, 'mineru_attempted', False)),
            'mineru_succeeded': bool(getattr(preprocess_result, 'mineru_succeeded', False)),
            'mineru_token_present': bool(getattr(preprocess_result, 'mineru_token_present', False)),
            'mineru_remote_requested': bool(getattr(preprocess_result, 'mineru_remote_requested', False)),
            'mineru_remote_enabled': bool(getattr(preprocess_result, 'mineru_remote_enabled', False)),
            'mineru_base_url': getattr(preprocess_result, 'mineru_base_url', ''),
            'markdown_path': getattr(preprocess_result, 'markdown_path', ''),
            'plain_text_path': getattr(preprocess_result, 'plain_text_path', ''),
            'page_index_path': getattr(preprocess_result, 'page_index_path', ''),
            'structured_json_path': getattr(preprocess_result, 'structured_json_path', ''),
            'diagnostics_path': getattr(preprocess_result, 'diagnostics_path', ''),
            'manifest_path': getattr(preprocess_result, 'manifest_path', ''),
            'stage1_input_path': getattr(preprocess_result, 'stage1_input_path', ''),
            'stage1_input_manifest_path': getattr(preprocess_result, 'stage1_input_manifest_path', ''),
            'stage1_quality_report_path': getattr(preprocess_result, 'stage1_quality_report_path', ''),
            'selected_text_source': getattr(preprocess_result, 'selected_text_source', ''),
            'stage1_quality_level': getattr(preprocess_result, 'stage1_quality_level', ''),
            'chunk_count': getattr(preprocess_result, 'chunk_count', 0),
            'cache_dir': getattr(preprocess_result, 'cache_dir', ''),
        }
        quality_report_path = str(metadata.get('stage1_quality_report_path') or '')
        if quality_report_path:
            try:
                with open(quality_report_path, "r", encoding="utf-8") as handle:
                    quality_report = json.load(handle)
                metadata['stage1_quality_reasons'] = quality_report.get('stage1_quality_reasons', [])
                metadata['stage1_completeness_metrics'] = quality_report.get('completeness_metrics', {})
            except Exception:
                metadata['stage1_quality_reasons'] = []
                metadata['stage1_completeness_metrics'] = {}
        else:
            metadata['stage1_quality_reasons'] = []
            metadata['stage1_completeness_metrics'] = {}
        manifest_path = str(metadata.get('stage1_input_manifest_path') or '')
        if manifest_path:
            try:
                with open(manifest_path, "r", encoding="utf-8") as handle:
                    stage1_manifest = json.load(handle)
                metadata['selected_text_length'] = int(stage1_manifest.get('selected_text_length') or 0)
                metadata['stage1_page_count'] = int(stage1_manifest.get('page_count') or 0)
                if not metadata.get('stage1_completeness_metrics'):
                    metadata['stage1_completeness_metrics'] = stage1_manifest.get('completeness_metrics', {})
            except Exception:
                metadata.setdefault('selected_text_length', 0)
                metadata.setdefault('stage1_page_count', 0)
        return metadata

    @staticmethod
    def _stage1_route_snapshot(strategy: str, metadata: Mapping[str, Any]) -> Dict[str, Any]:
        def text(key: str, default: str = "") -> str:
            value = metadata.get(key, default)
            return str(value if value is not None else default)

        parser_mode = text('parser_mode')
        primary_parser = text('primary_parser')
        remote_candidate = (
            strategy == 'mineru'
            or parser_mode in {'remote', 'remote_first'}
            or (parser_mode == 'hybrid' and primary_parser == 'mineru_remote')
        )
        requested = bool(metadata.get('mineru_remote_requested')) if 'mineru_remote_requested' in metadata else remote_candidate
        enabled = bool(metadata.get('mineru_remote_enabled')) if 'mineru_remote_enabled' in metadata else (
            strategy == 'mineru' or parser_mode in {'remote', 'remote_first'}
        )
        attempted = bool(metadata.get('mineru_attempted'))
        succeeded = bool(metadata.get('mineru_succeeded'))
        token_present = bool(metadata.get('mineru_token_present'))

        if succeeded:
            mineru_route = 'mineru_remote_succeeded'
        elif attempted:
            mineru_route = 'mineru_remote_failed_local_fallback'
        elif requested and not enabled:
            mineru_route = 'hybrid_local_baseline_met'
        elif requested and enabled and not token_present:
            mineru_route = 'mineru_token_missing'
        elif requested and enabled:
            mineru_route = 'mineru_remote_not_attempted'
        elif not requested:
            mineru_route = 'mineru_not_requested'
        else:
            mineru_route = 'mineru_not_used'

        return {
            'strategy': strategy,
            'preprocess_strategy': strategy,
            'preprocess_profile': text('preprocess_profile', strategy),
            'parser_mode': parser_mode,
            'primary_parser': primary_parser,
            'fallback_parser': text('fallback_parser'),
            'allow_local_parse_fallback': bool(metadata.get('allow_local_parse_fallback')),
            'extractor_used': text('extractor_used', 'unknown'),
            'selected_text_source': text('selected_text_source') or text('analysis_input_kind'),
            'stage1_quality_level': text('stage1_quality_level'),
            'stage1_quality_reasons': list(metadata.get('stage1_quality_reasons') or []),
            'selected_text_length': int(metadata.get('selected_text_length') or 0),
            'stage1_page_count': int(metadata.get('stage1_page_count') or 0),
            'mineru_token_present': token_present,
            'mineru_remote_requested': requested,
            'mineru_remote_enabled': enabled,
            'mineru_attempted': attempted,
            'mineru_succeeded': succeeded,
            'mineru_route': mineru_route,
        }

    @staticmethod
    def _format_stage1_route_snapshot(snapshot: Mapping[str, Any]) -> str:
        return (
            "阶段一输入路由: "
            f"strategy={snapshot.get('strategy')}, "
            f"parser_mode={snapshot.get('parser_mode')}, "
            f"extractor_used={snapshot.get('extractor_used')}, "
            f"selected_text_source={snapshot.get('selected_text_source')}, "
            f"stage1_quality_level={snapshot.get('stage1_quality_level')}, "
            f"mineru_remote_requested={snapshot.get('mineru_remote_requested')}, "
            f"mineru_remote_enabled={snapshot.get('mineru_remote_enabled')}, "
            f"mineru_attempted={snapshot.get('mineru_attempted')}, "
            f"mineru_succeeded={snapshot.get('mineru_succeeded')}, "
            f"mineru_route={snapshot.get('mineru_route')}"
        )

    @staticmethod
    def _stage1_route_human_message(snapshot: Mapping[str, Any]) -> str:
        route = str(snapshot.get('mineru_route') or '')
        if route == 'hybrid_local_baseline_met':
            return "配置把 MinerU 作为候选，但本地基线达标，所以本篇未发起 MinerU 请求。"
        if route == 'mineru_token_missing':
            return "配置要求 MinerU 远程解析，但 MINERU_API_TOKEN 缺失，所以本篇未发起 MinerU 请求。"
        if route == 'mineru_remote_failed_local_fallback':
            return "MinerU 远程解析已尝试但未成功，本篇已按配置回退到本地解析结果。"
        if route == 'mineru_remote_not_attempted':
            return "配置要求 MinerU 远程解析，但远程请求未实际发起；请检查配置和预处理诊断。"
        if route == 'mineru_not_requested':
            return "本篇阶段一输入没有请求 MinerU 远程解析，使用当前本地/兼容解析路径。"
        if route == 'mineru_remote_succeeded':
            return "本篇阶段一输入已使用 MinerU 远程解析结果。"
        return ""

    def _load_stage1_prompt_template(self) -> str:
        """加载阶段一结构化分析提示词模板。"""
        with open('prompts/optimized_prompt_analyze_router.txt', 'r', encoding='utf-8') as handle:
            return handle.read()

    def _prepare_stage1_input(self, pdf_path: str, preprocess_strategy: str = 'hybrid') -> Tuple[str, Dict[str, Any]]:
        """优先使用预处理工件，必要时回退到旧版文本提取。"""

        preprocess_metadata: Dict[str, Any] = {}
        preprocess_metadata['preprocess_strategy'] = preprocess_strategy
        preprocess_metadata['preprocess_profile'] = {
            'hybrid': 'hybrid',
            'docling': 'forced_docling',
            'mineru': 'forced_mineru_remote',
            'legacy': 'legacy_extractor',
        }.get(preprocess_strategy, preprocess_strategy)

        if self.preprocess_manager:
            original_parser_mode = self.preprocess_manager.parser_mode
            original_primary_parser = self.preprocess_manager.primary_parser
            original_force_rebuild = self.preprocess_manager.force_rebuild
            original_allow_local_parse_fallback = self.preprocess_manager.allow_local_parse_fallback
            original_force_docling_strategy = self.preprocess_manager.force_docling_strategy
            try:
                # 根据策略调整配置
                if preprocess_strategy == 'mineru':
                    # 强制使用 MinerU 远程解析
                    self.preprocess_manager.parser_mode = 'remote'
                    self.preprocess_manager.force_rebuild = True
                    self.preprocess_manager.allow_local_parse_fallback = False
                    self.preprocess_manager.force_docling_strategy = False
                    self.logger.info(f"强制使用 MinerU 远程解析策略: {os.path.basename(pdf_path)}")
                elif preprocess_strategy == 'docling':
                    # 强制使用 Docling 解析
                    self.preprocess_manager.parser_mode = 'local'
                    self.preprocess_manager.force_rebuild = True
                    self.preprocess_manager.allow_local_parse_fallback = False
                    self.preprocess_manager.force_docling_strategy = True
                    self.logger.info(f"强制使用 Docling 解析策略: {os.path.basename(pdf_path)}")
                else:
                    # 其他策略保持默认配置
                    self.preprocess_manager.force_rebuild = True
                    self.preprocess_manager.allow_local_parse_fallback = original_allow_local_parse_fallback
                    self.preprocess_manager.force_docling_strategy = False

                preprocess_metadata.update({
                    'parser_mode': self.preprocess_manager.parser_mode,
                    'primary_parser': self.preprocess_manager.primary_parser,
                    'fallback_parser': self.preprocess_manager.fallback_parser,
                    'allow_local_parse_fallback': self.preprocess_manager.allow_local_parse_fallback,
                })
                
                # 准备PDF
                preprocess_result = self.preprocess_manager.prepare_pdf(
                    pdf_path
                )

                if preprocess_result:
                    stage1_text = preprocess_result.stage1_input_text
                    input_kind = preprocess_result.selected_text_source or "stage1_input"
                    preprocess_metadata.update(self._collect_preprocess_result_metadata(preprocess_result, input_kind))
                    if is_blocked_stage1_quality(
                        preprocess_result.stage1_quality_level,
                        preprocess_metadata.get('stage1_quality_reasons'),
                    ):
                        preprocess_metadata['analysis_input_kind'] = preprocess_result.selected_text_source or 'blocked_stage1_input'
                        self.logger.warning(
                            f"阶段一输入质量闸阻止当前预处理结果: {os.path.basename(pdf_path)} "
                            f"({preprocess_result.stage1_quality_level}: "
                            f"{', '.join(preprocess_metadata.get('stage1_quality_reasons') or [])})"
                        )
                        return "", preprocess_metadata
                    if stage1_text and len(stage1_text.strip()) >= 500:
                        self.logger.info(
                            f"阶段一输入使用预处理结果: {os.path.basename(pdf_path)} -> "
                            f"{preprocess_result.extractor_used} / {input_kind} (策略: {preprocess_strategy})"
                        )
                        return stage1_text, preprocess_metadata
                    if preprocess_result.stage1_quality_level in {"REPROCESS", "BLOCK"}:
                        preprocess_metadata['analysis_input_kind'] = preprocess_result.selected_text_source or 'blocked_stage1_input'
                        self.logger.warning(
                            f"阶段一输入质量闸阻止当前预处理结果: {os.path.basename(pdf_path)} "
                            f"({preprocess_result.stage1_quality_level})"
                        )

                    self.logger.warning(
                        f"预处理结果文本过短，尝试其他策略: {os.path.basename(pdf_path)} "
                        f"({len(stage1_text) if stage1_text else 0} 字符)"
                    )
            except Exception as exc:
                self.logger.warning(f"预处理阶段失败，尝试其他策略: {exc}")
            finally:
                self.preprocess_manager.parser_mode = original_parser_mode
                self.preprocess_manager.primary_parser = original_primary_parser
                self.preprocess_manager.force_rebuild = original_force_rebuild
                self.preprocess_manager.allow_local_parse_fallback = original_allow_local_parse_fallback
                self.preprocess_manager.force_docling_strategy = original_force_docling_strategy

        # 如果是legacy策略或其他策略失败，使用旧版文本提取
        if preprocess_strategy == 'legacy':
            self.logger.info(f"阶段一输入使用旧版 PDF 文本提取: {os.path.basename(pdf_path)}")
            legacy_text = str(extract_text_from_pdf(pdf_path) or "")  # type: ignore
            # 直接从环境变量读取 MinerU 配置
            mineru_api_token = str(os.getenv("MINERU_API_TOKEN", "")).strip()
            mineru_base_url = str(os.getenv("MINERU_BASE_URL", "https://mineru.net/api/v4")).strip().rstrip("/")
            preprocess_metadata.update({
                'analysis_input_kind': 'legacy_text',
                'extractor_used': 'legacy_pdf_extractor',
                'parser_mode': 'legacy',
                'primary_parser': 'legacy_pdf_extractor',
                'fallback_parser': '',
                'allow_local_parse_fallback': False,
                'layout_fidelity': 'plain_text_only',
                'conversion_used': 'native_pdf',
                'used_ocr': False,
                'low_quality': False,
                'scanned_like': False,
                'mineru_attempted': False,
                'mineru_succeeded': False,
                'mineru_token_present': bool(mineru_api_token),
                'mineru_remote_requested': False,
                'mineru_remote_enabled': False,
                'mineru_base_url': mineru_base_url,
                'selected_text_source': 'legacy_text',
            })
            return legacy_text, preprocess_metadata

        # 策略失败，返回空文本和元数据
        # 确保即使在策略失败的情况下，也能正确设置 mineru_token_present
        if not preprocess_metadata.get('mineru_token_present'):
            mineru_api_token = str(os.getenv("MINERU_API_TOKEN", "")).strip()
            preprocess_metadata['mineru_token_present'] = bool(mineru_api_token)
        return "", preprocess_metadata

    @staticmethod
    def _normalize_metadata_scan_text(value: Any) -> str:
        text = str(value or "")
        text = text.replace("\u00a0", " ")
        text = text.replace("–", "-").replace("—", "-").replace("−", "-")
        text = re.sub(r"[*#_`]+", " ", text)
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    @classmethod
    def _metadata_match_key(cls, value: Any) -> str:
        text = cls._normalize_metadata_scan_text(value).casefold()
        return re.sub(r"[\W_]+", "", text, flags=re.UNICODE)

    @staticmethod
    def _extract_doi_from_text(value: Any) -> str:
        text = str(value or "").strip()
        if not text:
            return ""

        url_match = re.search(r"https?://(?:dx\.)?doi\.org/(10\.\d{4,9}/\S+)", text, flags=re.IGNORECASE)
        if url_match:
            text = url_match.group(1)
        else:
            doi_match = re.search(r"(10\.\d{4,9}/[-._;()/:A-Z0-9]+)", text, flags=re.IGNORECASE)
            if doi_match:
                text = doi_match.group(1)

        return text.rstrip(".,;:)]}>\"'")

    @staticmethod
    def _is_metadata_noise_line(line: str) -> bool:
        lowered = line.casefold()
        if not lowered:
            return True

        noise_tokens = (
            "issn:",
            "journal homepage",
            "to cite this article",
            "to link to this article",
            "published online",
            "submit your article",
            "article views",
            "view related articles",
            "view crossmark data",
            "full terms & conditions",
            "abstract",
            "article history",
            "keywords",
            "subjects",
            "received ",
            "accepted ",
            "revised ",
        )
        if any(token in lowered for token in noise_tokens):
            return True

        if " | " in line and "article" in lowered:
            return True

        if re.search(r"\b(?:vol|issue|no\.)\b", lowered) and re.search(r"\b(?:19|20)\d{2}\b", line):
            return True

        return False

    @staticmethod
    def _is_affiliation_line(line: str) -> bool:
        lowered = line.casefold()
        affiliation_tokens = (
            "university",
            "universitas",
            "faculty",
            "department",
            "school",
            "college",
            "hospital",
            "institute",
            "centre",
            "center",
            "laboratory",
            "contact ",
            "@",
            "http://",
            "https://",
        )
        return any(token in lowered for token in affiliation_tokens)

    @classmethod
    def _looks_like_inverted_author_parts(cls, parts: List[str]) -> bool:
        if len(parts) < 2 or len(parts) % 2 != 0:
            return False

        for surname_part, given_part in zip(parts[::2], parts[1::2]):
            surname_tokens = surname_part.split()
            given_tokens = given_part.split()
            if len(surname_tokens) != 1 or not 1 <= len(given_tokens) <= 3:
                return False

            combined_tokens = surname_tokens + given_tokens
            if any(any(ch.isdigit() for ch in token) for token in combined_tokens):
                return False
            if not all(any(ch.isalpha() for ch in token) for token in combined_tokens):
                return False

        return True

    @classmethod
    def _split_author_candidates(cls, value: Any) -> List[str]:
        text = cls._normalize_metadata_scan_text(value)
        if not text or cls._is_metadata_noise_line(text) or cls._is_affiliation_line(text):
            return []

        text = re.sub(r"\[[^\]]+\]", " ", text)
        text = re.sub(r"\b(?:orcid|id)\b", " ", text, flags=re.IGNORECASE)
        text = re.sub(r"[©®†‡]", " ", text)
        text = re.sub(r"\S+@\S+", " ", text)
        text = text.replace(" & ", "; ")
        text = re.sub(r"\band\b", ";", text, flags=re.IGNORECASE)
        text = re.sub(r"\s+", " ", text).strip(" ,;:-")

        authors: List[str] = []
        for raw_segment in re.split(r"\s*;\s*", text):
            segment = cls._normalize_metadata_scan_text(raw_segment).strip(" ,;:-")
            if not segment:
                continue

            parts: List[str] = []
            for raw_part in re.split(r"\s*,\s*", segment):
                part = cls._normalize_metadata_scan_text(raw_part)
                if not part:
                    continue
                part = re.sub(r"\s+[a-z](?:,[a-z])*$", "", part)
                part = part.strip(" ,;:-")
                if not part or any(ch.isdigit() for ch in part):
                    continue
                if cls._is_affiliation_line(part) or cls._is_metadata_noise_line(part):
                    continue
                parts.append(part)

            if not parts:
                continue

            if cls._looks_like_inverted_author_parts(parts):
                authors.extend(
                    f"{given_part} {surname_part}"
                    for surname_part, given_part in zip(parts[::2], parts[1::2])
                )
                continue

            authors.extend(parts)

        unique_authors: List[str] = []
        for author in authors:
            if author not in unique_authors:
                unique_authors.append(author)
        return unique_authors

    @classmethod
    def _extract_citation_metadata(cls, lines: List[str], title_hint: str) -> Dict[str, Any]:
        citation_line = next((line for line in lines[:80] if "to cite this article" in line.casefold()), "")
        if not citation_line:
            return {}

        match = re.search(
            r"to cite this article:\s*(?P<authors>.+?)\s*\((?P<year>(?:19|20)\d{2})\)\s*(?P<rest>.+)$",
            citation_line,
            flags=re.IGNORECASE,
        )
        if not match:
            return {"doi": cls._extract_doi_from_text(citation_line)}

        rest = cls._normalize_metadata_scan_text(match.group("rest"))
        doi = cls._extract_doi_from_text(rest or citation_line)
        if doi:
            rest = re.sub(rf",?\s*(?:DOI|doi)\s*:?\s*{re.escape(doi)}.*$", "", rest, flags=re.IGNORECASE)

        segments = [cls._normalize_metadata_scan_text(segment) for segment in rest.split(",")]
        segments = [segment for segment in segments if segment]

        metadata: Dict[str, Any] = {
            "authors": cls._split_author_candidates(match.group("authors")),
            "year": match.group("year"),
            "doi": doi,
        }

        title_key = cls._metadata_match_key(title_hint)
        if title_key and segments:
            cumulative = ""
            for index, segment in enumerate(segments):
                cumulative = f"{cumulative}, {segment}" if cumulative else segment
                if title_key and title_key in cls._metadata_match_key(cumulative):
                    metadata["title"] = cumulative
                    for candidate in segments[index + 1:]:
                        if cls._is_metadata_noise_line(candidate):
                            continue
                        if re.search(r"\b(?:19|20)\d{2}\b", candidate):
                            continue
                        if re.match(r"^\d", candidate):
                            continue
                        metadata["journal"] = candidate
                        break
                    break

        return metadata

    @classmethod
    def _extract_stage1_metadata(cls, stage1_text: str, paper: Mapping[str, Any]) -> Dict[str, Any]:
        if not stage1_text:
            return {}

        scan_text = stage1_text[:20000]
        lines = [
            cls._normalize_metadata_scan_text(line)
            for line in scan_text.splitlines()
        ]
        lines = [line for line in lines if line]
        if not lines:
            return {}

        metadata = cls._extract_citation_metadata(lines, str(paper.get("title") or ""))
        metadata.setdefault("doi", cls._extract_doi_from_text(scan_text))

        title_hint_key = cls._metadata_match_key(paper.get("title"))
        title_index: Optional[int] = None
        if title_hint_key:
            for index, line in enumerate(lines[:80]):
                if cls._is_metadata_noise_line(line):
                    continue
                line_key = cls._metadata_match_key(line)
                if not line_key:
                    continue
                if line_key == title_hint_key or title_hint_key in line_key or line_key in title_hint_key:
                    title_index = index
                    break

        if title_index is not None:
            title_line = lines[title_index]
            if title_line:
                metadata.setdefault("title", title_line)

            if not metadata.get("journal"):
                for index in range(title_index - 1, max(-1, title_index - 6), -1):
                    candidate = lines[index]
                    if cls._is_metadata_noise_line(candidate) or cls._is_affiliation_line(candidate):
                        continue
                    if re.search(r"\b(?:19|20)\d{2}\b", candidate):
                        continue
                    if re.search(r"\b(?:vol|issue|no\.)\b", candidate.casefold()):
                        continue
                    metadata["journal"] = candidate
                    break

            if not metadata.get("authors"):
                author_lines: List[str] = []
                for index in range(title_index + 1, min(len(lines), title_index + 4)):
                    candidate = lines[index]
                    if cls._is_metadata_noise_line(candidate) or cls._is_affiliation_line(candidate):
                        break
                    if re.search(r"\b(?:19|20)\d{2}\b", candidate):
                        break
                    author_lines.append(candidate)
                authors = cls._split_author_candidates(" ".join(author_lines))
                if authors:
                    metadata["authors"] = authors

            if not metadata.get("year"):
                for index in range(title_index - 1, max(-1, title_index - 6), -1):
                    candidate = lines[index]
                    year_match = re.search(r"\b(?:19|20)\d{2}\b", candidate)
                    if year_match:
                        metadata["year"] = year_match.group(0)
                        break

        return {
            "title": cls._normalize_metadata_scan_text(metadata.get("title")),
            "authors": list(metadata.get("authors") or []),
            "year": cls._normalize_metadata_scan_text(metadata.get("year")),
            "journal": cls._normalize_metadata_scan_text(metadata.get("journal")),
            "doi": cls._extract_doi_from_text(metadata.get("doi")),
        }

    def _apply_stage1_text_metadata_backfill(self, paper: PaperInfo, stage1_text: str) -> List[str]:
        needs_backfill = (
            not paper.get("authors")
            or self._is_placeholder_metadata_value(paper.get("year"))
            or self._is_placeholder_metadata_value(paper.get("journal"))
            or not str(paper.get("doi") or "").strip()
        )
        if not needs_backfill:
            return []

        metadata = sanitize_metadata_fields(self._extract_stage1_metadata(stage1_text, paper))
        updated_fields: List[str] = []

        extracted_authors = list(metadata.get("authors") or [])
        extracted_year = str(metadata.get("year") or "").strip()
        extracted_journal = str(metadata.get("journal") or "").strip()
        extracted_doi = str(metadata.get("doi") or "").strip()

        if extracted_authors and not paper.get("authors"):
            paper["authors"] = extracted_authors
            updated_fields.append("作者")

        if extracted_year and (
            self._is_placeholder_metadata_value(paper.get("year"))
            or not str(paper.get("year") or "").strip()
        ):
            paper["year"] = extracted_year
            updated_fields.append("年份")

        if extracted_journal and (
            self._is_placeholder_metadata_value(paper.get("journal"))
            or not str(paper.get("journal") or "").strip()
        ):
            paper["journal"] = extracted_journal
            updated_fields.append("期刊")

        if extracted_doi and not str(paper.get("doi") or "").strip():
            paper["doi"] = extracted_doi
            updated_fields.append("DOI")

        return updated_fields


    def _resolve_free_mode_context(self) -> str:
        """Build prompt context from a saved free-mode profile or an ad-hoc idea."""

        profile = self.free_mode_profile
        profile_path = (self.free_mode_profile_path or "").strip()

        if profile is None and profile_path:
            try:
                with open(profile_path, "r", encoding="utf-8") as handle:
                    loaded = json.load(handle)
                if isinstance(loaded, dict):
                    self.free_mode_profile = loaded
                    profile = loaded
                    self.logger.info(f"已加载自由模式 profile: {profile_path}")
            except Exception as exc:
                self.logger.warning(f"加载自由模式 profile 失败，将回退到临时意图文本: {exc}")

        if profile is None and self.output_dir and self.project_name:
            try:
                profile = load_profile(self.output_dir, self.project_name)
                if profile:
                    self.free_mode_profile = profile
                    self.logger.info("已自动加载项目自由模式 profile。")
            except Exception as exc:
                self.logger.warning(f"自动加载项目自由模式 profile 失败: {exc}")

        if profile:
            return build_profile_context(profile)

        free_mode_idea = str(self.free_mode_idea or "").strip()
        if free_mode_idea:
            return f"\n[FREE MODE IDEA]\n{free_mode_idea}\n"

        return ""

    def _inject_free_mode_context(self, prompt_template: str) -> str:
        free_mode_context = self._resolve_free_mode_context()
        prompt_with_placeholder = prompt_template.replace("{{FREE_MODE_CONTEXT}}", free_mode_context)
        if free_mode_context and "{{FREE_MODE_CONTEXT}}" not in prompt_template:
            return f"{free_mode_context}\n{prompt_with_placeholder}"
        return prompt_with_placeholder

    def _build_stage1_analysis_prompt(self, pdf_text: str) -> str:
        prompt_template = self._load_stage1_prompt_template()
        prompt_template = self._inject_free_mode_context(prompt_template)
        return prompt_template.replace("{{PAPER_FULL_TEXT}}", pdf_text)

    def _build_stage1_visual_bundle(
        self,
        *,
        paper: Mapping[str, Any],
        pdf_path: str,
        preprocess_metadata: Mapping[str, Any],
    ) -> Optional[Dict[str, Any]]:
        if not self.job_workspace or not self.artifact_registry:
            return None

        paper_key = self._paper_artifact_key(paper)
        artifact_hash = self._paper_artifact_hash(paper_key)
        output_dir = self.job_workspace.artifact_path(f"stage1_visuals/{artifact_hash}")
        self.stage1_visual_builder.logger = self.logger

        try:
            bundle = self.stage1_visual_builder.build_bundle(
                job_id=self.job_workspace.job_id,
                paper_key=paper_key,
                paper_info=paper,
                source_pdf=pdf_path,
                output_dir=output_dir,
                artifact_registry=self.artifact_registry,
                preprocess_metadata=preprocess_metadata,
            )
            return bundle.to_dict() if bundle else None
        except Exception as exc:
            self.logger.warning(f"Failed to build stage-1 visual bundle, continuing with text-first fallback: {exc}")
            return None

    def _build_stage1_model_input(
        self,
        *,
        pdf_text: str,
        reader_api_config: Mapping[str, Any],
        visual_bundle: Optional[Mapping[str, Any]] = None,
        paper: Optional[Mapping[str, Any]] = None,
    ) -> Dict[str, Any]:
        self.stage1_input_builder.logger = self.logger
        prompt_template = self._load_stage1_prompt_template()
        prompt_template = self._inject_free_mode_context(prompt_template)
        
        # Use registry-first visual artifact resolution if available
        resolved_visual_bundle = dict(visual_bundle or {})
        if paper and self.artifact_registry and self.job_workspace:
            from services.visual_artifact_resolver import VisualArtifactResolver
            resolver = VisualArtifactResolver(self.artifact_registry, self.logger)
            
            # Try to resolve paper artifact
            paper_key = self._paper_artifact_key(paper)
            artifact_hash = self._paper_artifact_hash(paper_key)
            paper_artifact_path = self.job_workspace.artifact_path(f"paper_artifacts/{artifact_hash}.json")
            
            # Resolve selected visual refs from paper artifact or registry
            selected_visual_refs = resolver.resolve_selected_visual_refs(paper_artifact_path)
            if selected_visual_refs:
                resolved_visual_bundle["selected_visual_refs"] = selected_visual_refs
                resolved_visual_bundle.setdefault("visual_manifest_path", "")
                resolved_visual_bundle.setdefault("bundle_path", "")
                resolved_visual_bundle.setdefault("selection_policy_snapshot", {})

                manifest_path = resolver.resolve_visual_manifest_path(paper_artifact_path)
                if manifest_path:
                    resolved_visual_bundle["visual_manifest_path"] = manifest_path
        
        built_input = self.stage1_input_builder.build(
            prompt_template=prompt_template,
            paper_text=pdf_text,
            reader_api_config=reader_api_config,
            visual_bundle=resolved_visual_bundle,
        )
        return built_input.to_metadata_dict()

    def _get_outline_file_path(self) -> str:
        if not self.output_dir:
            raise ValueError("输出目录未设置")
        if self.job_workspace is not None and self.project_name:
            return self.job_workspace.artifact_path(f"{self.project_name}_literature_review_outline.md")
        if self.project_name:
            return os.path.join(self.output_dir, f"{self.project_name}_literature_review_outline.md")
        return os.path.join(self.output_dir, "literature_review_outline.md")

    def _get_legacy_outline_file_path(self) -> str:
        if self.project_name:
            if self.job_workspace is not None:
                return os.path.join(
                    self.job_workspace.project_pointer_dir(),
                    f"{self.project_name}_literature_review_outline.md",
                )

            if self.config:
                output_base_path = self.config.get("Paths", {}).get("output_path", "./output")  # type: ignore[union-attr]
                return os.path.join(
                    os.path.abspath(output_base_path),
                    self.project_name,
                    f"{self.project_name}_literature_review_outline.md",
                )

        if not self.output_dir:
            raise ValueError("output directory is not configured")
        return os.path.join(self.output_dir, "literature_review_outline.md")

    def _write_outline_artifact(self, outline_text: str, *, producer: str) -> str:
        outline_file = self._get_outline_file_path()
        with open(outline_file, "w", encoding="utf-8") as handle:
            handle.write(outline_text)

        depends_on: List[ArtifactDependencyRef] = []
        if self.summary_file:
            depends_on.append(
                ArtifactDependencyRef(
                    artifact_type="summary_file",
                    path=self.summary_file,
                )
            )

        if self.artifact_registry:
            self.artifact_registry.register_file(
                artifact_role=self.OUTLINE_ARTIFACT_ROLE,
                artifact_type=self.OUTLINE_ARTIFACT_TYPE,
                artifact_version=self.OUTLINE_ARTIFACT_VERSION,
                path=outline_file,
                producer=producer,
                depends_on=depends_on,
                artifact_id=self.OUTLINE_ARTIFACT_ID,
            )

        return outline_file

    def _resolve_outline_file_path(self) -> Optional[str]:
        if self.artifact_registry:
            record = self.artifact_registry.get(self.OUTLINE_ARTIFACT_ID)
            if record and record.status == "ready" and os.path.exists(record.path):
                self.logger.info(f"Using registered outline artifact: {record.path}")
                return record.path
            if record and not os.path.exists(record.path):
                self.logger.warning(f"Registered outline artifact is missing on disk: {record.path}")

        workspace_outline_file = self._get_outline_file_path()
        if os.path.exists(workspace_outline_file):
            self.logger.info(f"Using workspace outline file without registry lookup fallback: {workspace_outline_file}")
            return workspace_outline_file

        legacy_outline_file = self._get_legacy_outline_file_path()
        if legacy_outline_file != workspace_outline_file and os.path.exists(legacy_outline_file):
            self.logger.warning(f"Using legacy outline compatibility fallback: {legacy_outline_file}")
            return legacy_outline_file

        return None

    def _load_outline_artifact(self) -> Optional[Tuple[str, str]]:
        """Load the primary markdown outline artifact used by downstream review generation."""
        if self._outline_v2_enabled():
            from outline.runtime_resolver import OutlineRuntimeResolver

            resolver = OutlineRuntimeResolver(
                config=self.config or {},
                artifact_registry=self.artifact_registry,
                workspace_path=self.job_workspace.paths.root_dir if self.job_workspace else (self.output_dir or ""),
                project_name=self.project_name or "",
                legacy_outline_path=self._get_legacy_outline_file_path(),
            )
            resolved = resolver.resolve_for_review()
            if resolved is None:
                self.logger.error(
                    "Outline Intelligence v2 已启用，但未找到有效 adopted_final_outline.json；"
                    "请先运行 --adopt-outline-v2，且不会回退到 legacy Markdown。"
                )
                return None
            self.logger.info(f"Using v2 adopted outline artifact: {resolved.source_path}")
            return resolved.source_path, resolved.markdown

        outline_file = self._resolve_outline_file_path()
        if outline_file:
            try:
                with open(outline_file, "r", encoding="utf-8") as handle:
                    return outline_file, handle.read()
            except Exception as exc:
                self.logger.error(f"Failed to read outline artifact: {exc}")

        # 如果当前工作空间找不到大纲文件，检查历史工作空间
        if self.project_name:
            paths_config: Dict[str, str] = self.config.get('Paths', {}) if self.config else {}
            output_base_path: str = paths_config.get('output_path', './output')
            output_base_path_abs = os.path.abspath(output_base_path)
            
            # 查找所有历史工作空间
            import glob
            workspace_pattern = os.path.join(output_base_path_abs, f"{self.project_name}__*")
            workspaces = glob.glob(workspace_pattern)
            
            # 按修改时间排序，找到最新的工作空间
            workspaces.sort(key=os.path.getmtime, reverse=True)
            
            for workspace_path in workspaces:
                # 检查该工作空间是否有大纲文件
                outline_file = os.path.join(workspace_path, "artifacts", f"{self.project_name}_literature_review_outline.md")
                if os.path.exists(outline_file):
                    self.logger.info(f"在历史工作空间中找到大纲文件: {outline_file}")
                    try:
                        with open(outline_file, "r", encoding="utf-8") as handle:
                            return outline_file, handle.read()
                    except Exception as exc:
                        self.logger.warning(f"读取历史大纲文件失败: {exc}")
                        continue

        self.logger.error("No outline artifact was found in the current workspace/registry or legacy fallback path")
        return None

    def _get_review_checkpoint_file_path(self) -> str:
        if not self.output_dir:
            raise ValueError("输出目录未设置")
        if self.job_workspace is not None and self.project_name:
            return self.job_workspace.checkpoint_path(f"{self.project_name}_review_checkpoint.json")
        if self.project_name:
            return os.path.join(self.output_dir, f"{self.project_name}_review_checkpoint.json")
        return os.path.join(self.output_dir, "review_checkpoint.json")

    def _get_review_word_file_path(self) -> str:
        if not self.output_dir:
            raise ValueError("输出目录未设置")
        if self.job_workspace is not None and self.project_name:
            return self.job_workspace.report_path(f"{self.project_name}_literature_review.docx")
        if self.project_name:
            return os.path.join(self.output_dir, f"{self.project_name}_literature_review.docx")
        return os.path.join(self.output_dir, "literature_review.docx")

    def _get_failed_review_sections_file_path(self) -> str:
        if not self.output_dir:
            raise ValueError("输出目录未设置")
        if self.job_workspace is not None and self.project_name:
            return self.job_workspace.report_path(f"{self.project_name}_failed_review_sections.json")
        if self.project_name:
            return os.path.join(self.output_dir, f"{self.project_name}_failed_review_sections.json")
        return os.path.join(self.output_dir, "failed_review_sections.json")

    @staticmethod
    def _metadata_match_variants(value: Any) -> Set[str]:
        text = str(value or "").strip()
        variants: Set[str] = set()
        if not text:
            return variants

        variants.add(text)
        for encoding in ("gb18030", "gbk"):
            try:
                repaired = text.encode(encoding).decode("utf-8").strip()
            except Exception:
                continue
            if repaired:
                variants.add(repaired)
        return variants

    @staticmethod
    def _is_placeholder_metadata_value(value: Any) -> bool:
        variants = LiteratureReviewGenerator._metadata_match_variants(value)
        if not variants:
            return True

        placeholder_texts = {"未知", "未知年份", "未知期刊"}
        placeholder_keywords = {"unknown", "n/a", "na", "none", "null"}
        return any(
            candidate.casefold() in placeholder_keywords or candidate in placeholder_texts
            for candidate in variants
        )

    @staticmethod
    def _quality_reason_parts(reason: str) -> List[str]:
        return [part.strip() for part in re.split(r"[;；\n]+", str(reason or "")) if part.strip()]

    @staticmethod
    def _is_metadata_quality_issue(reason_part: str) -> bool:
        lowered = str(reason_part or "").casefold()
        metadata_markers = (
            "author",
            "authors",
            "year",
            "journal",
            "metadata",
            "paper_metadata",
            "paper_info",
            "作者",
            "年份",
            "期刊",
            "元数据",
        )
        return any(marker.casefold() in lowered for marker in metadata_markers)

    def _is_metadata_only_quality_failure(self, reason: str) -> bool:
        parts = self._quality_reason_parts(reason)
        return bool(parts) and all(self._is_metadata_quality_issue(part) for part in parts)

    def _apply_filename_metadata_backfill(self, paper: PaperInfo) -> List[str]:
        pdf_path = str(paper.get("pdf_path") or "")
        if not pdf_path:
            return []
        stem = os.path.splitext(os.path.basename(pdf_path))[0].strip()
        if not stem:
            return []

        updated_fields: List[str] = []
        if not str(paper.get("title") or "").strip():
            paper["title"] = stem
            updated_fields.append("title")

        if self._is_placeholder_metadata_value(paper.get("year")):
            match = re.search(r"\b(?:19|20)\d{2}\b", stem)
            if match:
                paper["year"] = match.group(0)
                updated_fields.append("year")
        return updated_fields

    def _apply_paper_metadata_to_ai_summary(self, paper: PaperInfo, ai_summary: Any) -> List[str]:
        if not isinstance(ai_summary, dict):
            return []
        paper_metadata = ai_summary.setdefault("paper_metadata", {})
        if not isinstance(paper_metadata, dict):
            paper_metadata = {}
            ai_summary["paper_metadata"] = paper_metadata

        updated_fields: List[str] = []

        def set_if_missing(field: str, value: Any) -> None:
            current = paper_metadata.get(field)
            current_missing = (
                not current
                or (field in {"year", "journal"} and self._is_placeholder_metadata_value(current))
                or (field == "authors" and not current)
            )
            if value and current_missing:
                paper_metadata[field] = value
                updated_fields.append(f"paper_metadata.{field}")

        title = str(paper.get("title") or "").strip()
        authors = paper.get("authors") or []
        year = str(paper.get("year") or "").strip()
        journal = str(paper.get("journal") or "").strip()
        doi = str(paper.get("doi") or "").strip()

        set_if_missing("title", title)
        if authors:
            set_if_missing("authors", authors)
        if year and not self._is_placeholder_metadata_value(year):
            set_if_missing("year", year)
        if journal and not self._is_placeholder_metadata_value(journal):
            set_if_missing("journal", journal)
        set_if_missing("doi", doi)

        if updated_fields:
            quality_audit = ai_summary.setdefault("quality_audit", {})
            inferred_fields = quality_audit.setdefault("inferred_fields", [])
            if isinstance(inferred_fields, list):
                for field in updated_fields:
                    note = f"{field} from paper_info"
                    if note not in inferred_fields:
                        inferred_fields.append(note)
        return updated_fields

    def _resolve_stage1_metadata_for_quality(self, paper: PaperInfo, ai_summary: Any, stage1_text: str) -> List[str]:
        updated_fields: List[str] = []
        try:
            updated_fields.extend(self._apply_filename_metadata_backfill(paper))
        except Exception as exc:
            self.logger.warning(f"Filename metadata backfill failed: {exc}")
        try:
            updated_fields.extend(self._apply_ai_metadata_backfill(paper, ai_summary))
        except Exception as exc:
            self.logger.warning(f"AI metadata backfill failed: {exc}")
        try:
            updated_fields.extend(self._apply_stage1_text_metadata_backfill(paper, stage1_text))
        except Exception as exc:
            self.logger.warning(f"Stage-1 text metadata backfill failed: {exc}")
        try:
            updated_fields.extend(self._apply_paper_metadata_to_ai_summary(paper, ai_summary))
        except Exception as exc:
            self.logger.warning(f"Summary metadata sync failed: {exc}")
        return updated_fields

    def _missing_metadata_fields_for_summary(self, paper: PaperInfo, ai_summary: Any) -> List[str]:
        metadata = sanitize_metadata_fields(get_paper_metadata(ai_summary or {}))
        authors = paper.get("authors") or metadata.get("authors") or []
        year = paper.get("year") or metadata.get("year") or ""
        journal = paper.get("journal") or metadata.get("journal") or ""
        missing: List[str] = []
        if not authors:
            missing.append("authors")
        if self._is_placeholder_metadata_value(year):
            missing.append("year")
        if self._is_placeholder_metadata_value(journal):
            missing.append("journal")
        return missing

    def _mark_summary_metadata_manual_review(self, paper: PaperInfo, ai_summary: Any, quality_reason: str) -> List[str]:
        if not isinstance(ai_summary, dict):
            return []
        missing_fields = self._missing_metadata_fields_for_summary(paper, ai_summary)
        quality_audit = ai_summary.setdefault("quality_audit", {})
        if not isinstance(quality_audit, dict):
            quality_audit = {}
            ai_summary["quality_audit"] = quality_audit
        quality_audit["needs_manual_review"] = True

        missing_critical = quality_audit.setdefault("missing_critical_fields", [])
        if isinstance(missing_critical, list):
            for field in missing_fields:
                marker = f"paper_metadata.{field}"
                if marker not in missing_critical:
                    missing_critical.append(marker)

        inferred_fields = quality_audit.setdefault("inferred_fields", [])
        if isinstance(inferred_fields, list):
            warning = f"metadata unresolved after local resolution: {quality_reason}"
            if warning not in inferred_fields:
                inferred_fields.append(warning)

        conflict_flags = quality_audit.setdefault("conflict_flags", [])
        if isinstance(conflict_flags, list) and "metadata_needs_manual_review" not in conflict_flags:
            conflict_flags.append("metadata_needs_manual_review")
        return missing_fields

    def _mark_metadata_manual_review_if_missing(self, paper: PaperInfo, ai_summary: Any, reason: str) -> List[str]:
        missing_fields = self._missing_metadata_fields_for_summary(paper, ai_summary)
        if not missing_fields:
            return []
        return self._mark_summary_metadata_manual_review(paper, ai_summary, reason)

    def _apply_ai_metadata_backfill(self, paper: PaperInfo, ai_summary: Any) -> List[str]:
        metadata = sanitize_metadata_fields(get_paper_metadata(ai_summary or {}))
        updated_fields: List[str] = []

        extracted_title = str(metadata.get("title") or "").strip()
        extracted_authors = list(metadata.get("authors") or [])
        extracted_year = str(metadata.get("year") or "").strip()
        extracted_journal = str(metadata.get("journal") or "").strip()
        extracted_doi = str(metadata.get("doi") or "").strip()

        if extracted_title and extracted_title != str(paper.get("title") or "").strip():
            paper["title"] = extracted_title
            updated_fields.append("标题")

        if extracted_authors and not paper.get("authors"):
            paper["authors"] = extracted_authors
            updated_fields.append("作者")

        if extracted_year and (
            self._is_placeholder_metadata_value(paper.get("year"))
            or not str(paper.get("year") or "").strip()
        ):
            paper["year"] = extracted_year
            updated_fields.append("年份")

        if extracted_journal and (
            self._is_placeholder_metadata_value(paper.get("journal"))
            or not str(paper.get("journal") or "").strip()
        ):
            paper["journal"] = extracted_journal
            updated_fields.append("期刊")

        if extracted_doi and not str(paper.get("doi") or "").strip():
            paper["doi"] = extracted_doi
            updated_fields.append("DOI")

        return updated_fields


    def _save_failed_review_sections(self, failed_sections: List[Dict[str, Any]]) -> None:
        failed_sections_file = self._get_failed_review_sections_file_path()
        if failed_sections:
            payload = {
                "failed_sections": failed_sections,
                "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            }
            with open(failed_sections_file, "w", encoding="utf-8") as handle:
                json.dump(payload, handle, ensure_ascii=False, indent=2)
            return

        if os.path.exists(failed_sections_file):
            os.remove(failed_sections_file)

    def _load_failed_review_sections(self) -> List[Dict[str, Any]]:
        failed_sections_file = self._get_failed_review_sections_file_path()
        if not os.path.exists(failed_sections_file):
            return []
        try:
            with open(failed_sections_file, "r", encoding="utf-8") as handle:
                payload = json.load(handle)
            failed_sections = payload.get("failed_sections", []) if isinstance(payload, dict) else []
            if isinstance(failed_sections, list):
                return [item for item in failed_sections if isinstance(item, dict)]
        except Exception as exc:
            self.logger.warning(f"读取失败章节记录失败: {exc}")
        return []

    def _clear_failed_review_section(self, section_number: int) -> None:
        failed_sections = [
            item
            for item in self._load_failed_review_sections()
            if int(item.get("section_number", 0) or 0) != section_number
        ]
        self._save_failed_review_sections(failed_sections)

    def _remove_paragraph(self, paragraph: Any) -> None:
        element = paragraph._element  # type: ignore[attr-defined]
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    def _trim_review_document_from_section(self, word_file: str, section_number: int) -> bool:
        if not os.path.exists(word_file):
            return True
        if not DOCX_AVAILABLE or Document is None:
            self.logger.warning("python-docx 不可用，将退回整篇重生而不是裁剪已有文档。")
            return False

        try:
            document = Document(word_file)  # type: ignore[operator]
            current_section_number = 0
            removing = False

            for paragraph in list(document.paragraphs):
                text = paragraph.text.strip()
                heading_match = re.match(r"^第\s*(\d+)\s*章", text)
                if heading_match:
                    current_section_number = int(heading_match.group(1))
                    if current_section_number >= section_number:
                        removing = True
                elif text == "参考文献" and current_section_number >= section_number:
                    removing = True

                if removing:
                    self._remove_paragraph(paragraph)

            document.save(word_file)
            return True
        except Exception as exc:
            self.logger.warning(f"裁剪已有综述文档失败，将退回整篇重生: {exc}")
            return False

    def generate_specific_review_section(self, section_number: int) -> bool:
        try:
            if not self.config and not self.load_configuration():
                return False
            if not self.output_dir and not self.setup_output_directory():
                return False
            if not self.load_existing_summaries():
                self.logger.error("无法加载摘要文件，请先运行阶段一")
                return False
            if not self.summaries:
                self.logger.error("没有可用的摘要数据，请先运行阶段一")
                return False

            outline_artifact = self._load_outline_artifact()
            outline_file = self._get_legacy_outline_file_path()
            if outline_artifact is None:
                self.logger.error(f"大纲文件不存在: {outline_file}，请先运行 --generate-outline")
                return False

            _outline_file, outline_content = outline_artifact

            section_title = self.extract_section_title_from_outline(outline_content, section_number)
            if not section_title:
                self.logger.error(f"未在大纲中找到第 {section_number} 章")
                return False

            return self.create_literature_review_section(section_number, section_title, outline_content)
        except Exception as exc:
            self.logger.error(f"补写指定章节失败: {exc}")
            return False

    def retry_failed_review_sections(self) -> bool:
        try:
            self._check_cancelled()
            if not self.config and not self.load_configuration():
                return False
            if not self.output_dir and not self.setup_output_directory():
                return False
            if not self.load_existing_summaries():
                self.logger.error("无法加载摘要文件，请先运行阶段一")
                return False

            failed_sections = self._load_failed_review_sections()
            if not failed_sections:
                self.logger.info("未找到失败章节记录，将直接重跑阶段二综述生成。")
                return self.generate_full_review_from_outline()

            section_numbers = sorted(
                {
                    int(item.get("section_number", 0) or 0)
                    for item in failed_sections
                    if int(item.get("section_number", 0) or 0) > 0
                }
            )
            if not section_numbers:
                self.logger.info("失败章节记录为空，将直接重跑阶段二综述生成。")
                return self.generate_full_review_from_outline()

            first_failed_section = section_numbers[0]
            checkpoint_file = self._get_review_checkpoint_file_path()
            word_file = self._get_review_word_file_path()
            resume_from_section = 0

            if os.path.exists(word_file) and self._trim_review_document_from_section(word_file, first_failed_section):
                resume_from_section = first_failed_section - 1

            checkpoint_payload = {
                "last_completed_section": resume_from_section,
                "last_section_title": "",
                "update_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            }
            atomic_write_json(checkpoint_file, checkpoint_payload)

            self.logger.info(
                f"将从第 {resume_from_section + 1} 章开始重新生成，以补齐之前失败的章节。"
            )
            return self.generate_full_review_from_outline()
        except Exception as exc:
            self.logger.error(f"重试失败章节时出错: {exc}")
            return False
    
    def process_paper(self, paper: PaperInfo, paper_index: int, file_index: Optional[FileIndex], total_papers: int) -> Optional[ProcessingResult]:
        """处理单篇论文"""
        try:
            self._check_cancelled()
            paper_key = LiteratureReviewGenerator.get_paper_key(paper)  # type: ignore
            
            # 检查是否已在断点中处理过
            if paper_key in self._checkpoint_processed_papers:
                self.logger.info(f"跳过已处理论文: {paper.get('title', '未知标题')}")
                # 从现有摘要中找到对应的条目
                for summary in self.summaries:
                    if summary.get('status') == 'success' and LiteratureReviewGenerator.get_paper_key(summary.get('paper_info', {})) == paper_key:
                        return summary
                return None
            
            if paper_key in self._checkpoint_failed_papers:
                self.logger.info(f"重新尝试上次失败论文: {paper.get('title', '未知标题')}")
                self._checkpoint_failed_papers.discard(paper_key)
            
            self.logger.info(f"[{paper_index+1}/{total_papers}] 正在处理: {paper.get('title', '未知标题')}")
            
            # 获取PDF文件路径
            paper_label = self._paper_progress_label(paper)
            self._emit_progress(stage="analyze", item_label=paper_label, message=f"正在处理: {paper_label}")
            pdf_path = paper.get('pdf_path')
            if not pdf_path and self.mode == "zotero":
                # Zotero模式下查找PDF文件
                file_title = paper.get('title', '')
                _file_authors = paper.get('authors', [])
                paths_config: Dict[str, str] = self.config.get('Paths', {}) if self.config else {}
                library_path: str = self.library_path or paths_config.get('library_path', '')
                
                if not library_path:
                    failure_reason = "配置文件中缺少library_path路径"
                    self.logger.error(failure_reason)
                    return {
                        'paper_info': paper,
                        'status': 'failed',
                        'failure_reason': failure_reason
                    }
                
                # 创建文件索引（如果还没有）
                if not file_index:
                    file_index = create_file_index(library_path)
                
                # 使用 file_finder.py 中强大的 find_pdf 函数
                find_result = find_pdf(dict(paper), library_path, file_index)
                
                if find_result:
                    pdf_path = find_result
                    self.logger.info(f"智能查找到PDF: {os.path.basename(pdf_path)}")
                else:
                    failure_reason = "未找到PDF文件"
                    self.logger.error(f"未找到PDF文件: {file_title} - 原因: {failure_reason}")
                    return {
                        'paper_info': paper,
                        'status': 'failed',
                        'failure_reason': failure_reason
                    }
            elif not pdf_path and self.mode == "direct":
                # 直接模式下PDF路径应该已经存在
                pdf_path = paper.get('pdf_path', '')
            
            if not pdf_path or not os.path.exists(pdf_path):
                failure_reason = f"PDF文件不存在: {pdf_path}"
                self.logger.error(failure_reason)
                return {
                    'paper_info': paper,
                    'status': 'failed',
                    'failure_reason': failure_reason
                }
            
            # Keep the existing parser retry route: hybrid -> forced Docling -> MinerU remote if configured -> legacy.
            mineru_configured = bool(
                (
                    getattr(self.preprocess_manager, 'mineru_api_token', '')
                    if self.preprocess_manager is not None
                    else os.getenv('MINERU_API_TOKEN', '')
                )
                or ''
            )
            preprocess_strategies = ['hybrid', 'docling']
            if mineru_configured:
                preprocess_strategies.append('mineru')
            preprocess_strategies.append('legacy')
            attempt_history = []
            
            # 初始化变量，确保在所有路径中都有定义
            preprocess_metadata = {}
            stage1_input_snapshot = {}
            model_used = 'primary'
            ai_result = None
            strategy_succeeded = False

            def record_attempt_failure(
                strategy_name: str,
                current_metadata: Dict[str, Any],
                *,
                model_name: str,
                quality_reason: str,
                extractor_name: Optional[str] = None,
            ) -> None:
                route_snapshot = self._stage1_route_snapshot(strategy_name, current_metadata)
                attempt_history.append({
                    'preprocess_strategy': strategy_name,
                    'preprocess_profile': str(route_snapshot.get('preprocess_profile') or strategy_name),
                    'parser_mode': str(route_snapshot.get('parser_mode') or ''),
                    'extractor_used': extractor_name or str(route_snapshot.get('extractor_used') or 'unknown'),
                    'selected_text_source': str(route_snapshot.get('selected_text_source') or ''),
                    'stage1_quality_level': str(route_snapshot.get('stage1_quality_level') or ''),
                    'stage1_quality_reasons': list(route_snapshot.get('stage1_quality_reasons') or []),
                    'selected_text_length': int(route_snapshot.get('selected_text_length') or 0),
                    'stage1_page_count': int(route_snapshot.get('stage1_page_count') or 0),
                    'mineru_remote_requested': bool(route_snapshot.get('mineru_remote_requested')),
                    'mineru_remote_enabled': bool(route_snapshot.get('mineru_remote_enabled')),
                    'mineru_attempted': bool(route_snapshot.get('mineru_attempted')),
                    'mineru_succeeded': bool(route_snapshot.get('mineru_succeeded')),
                    'mineru_route': str(route_snapshot.get('mineru_route') or ''),
                    'stage1_route': route_snapshot,
                    'model_used': model_name,
                    'quality_reason': quality_reason,
                    'success': False,
                })
            
            for strategy in preprocess_strategies:
                self._check_cancelled()
                ai_result = None
                
                # 准备阶段一输入
                self.logger.info(f"正在准备阶段一输入 (策略: {strategy}): {os.path.basename(pdf_path)}")
                self._emit_progress(stage="analyze", item_label=paper_label, message=f"正在准备阶段一输入 (策略: {strategy}): {os.path.basename(pdf_path)}")
                
                pdf_text, preprocess_metadata = self._prepare_stage1_input(pdf_path, strategy)
                stage1_route_snapshot = self._stage1_route_snapshot(strategy, preprocess_metadata)
                preprocess_metadata['stage1_route'] = stage1_route_snapshot
                self.logger.info(self._format_stage1_route_snapshot(stage1_route_snapshot))
                route_human_message = self._stage1_route_human_message(stage1_route_snapshot)
                if route_human_message:
                    self.logger.info(route_human_message)

                if not pdf_text or len(pdf_text.strip()) < 500:  # type: ignore
                    failure_reason = f"阶段一输入准备失败或内容过少({len(pdf_text) if pdf_text else 0}字符)"  # type: ignore
                    stage1_reasons = [
                        str(reason)
                        for reason in (preprocess_metadata.get('stage1_quality_reasons') or [])
                        if str(reason).strip()
                    ]
                    if stage1_reasons:
                        failure_reason = (
                            f"阶段一输入质量闸阻止: {', '.join(stage1_reasons)} "
                            f"({len(pdf_text) if pdf_text else 0}字符)"
                        )
                    self.logger.warning(f"策略 {strategy} 失败: {failure_reason}")
                    record_attempt_failure(
                        strategy,
                        preprocess_metadata,
                        model_name='N/A',
                        quality_reason=failure_reason,
                    )
                    continue

                input_kind = preprocess_metadata.get('analysis_input_kind', 'text')
                extractor_used = preprocess_metadata.get('extractor_used', 'unknown')
                self.logger.success(
                    f"阶段一输入准备成功 (策略: {strategy}): {len(pdf_text)}字符 "
                    f"({input_kind} / {extractor_used})"
                )
                self._check_cancelled()

                try:
                    updated_fields = self._apply_stage1_text_metadata_backfill(paper, pdf_text)
                    if updated_fields:
                        self.logger.info(f"已从阶段一输入回填元数据字段: {', '.join(updated_fields)}")
                except Exception as e:
                    self.logger.warning(f"阶段一输入元数据回填失败: {e}")
                
                # 调用AI API生成摘要
                self.logger.info("正在调用AI生成摘要...")
                
                # 提取分析引擎API配置
                self._emit_progress(stage="analyze", item_label=paper_label, message=f"正在调用AI生成摘要: {paper_label}")
                reader_api_config: APIConfig = get_reader_api_config(self.config)
                backup_api_config: APIConfig = get_backup_reader_api_config(self.config)
                
                visual_bundle = self._build_stage1_visual_bundle(
                    paper=paper,
                    pdf_path=pdf_path,
                    preprocess_metadata=preprocess_metadata,
                )

                # 构建显式的阶段一输入：文本始终是主输入，视觉证据只作为受控补充
                try:
                    stage1_input = self._build_stage1_model_input(
                        pdf_text=pdf_text,
                        reader_api_config=reader_api_config,
                        visual_bundle=visual_bundle,
                        paper=paper,
                    )
                    analysis_prompt = str(stage1_input.get("prompt_text") or pdf_text)
                except Exception as e:
                    self.logger.warning(f"无法构建显式阶段一输入，回退到文本提示词: {e}")
                    stage1_input = {
                        "input_mode": "text_only",
                        "prompt_text": f"请分析以下论文内容，生成结构化摘要：\n\n{pdf_text}",
                        "user_message_content": None,
                        "selected_visual_refs": [],
                        "visual_manifest_path": "",
                        "visual_bundle_path": "",
                        "visual_selection_policy_snapshot": {},
                        "multimodal_capability": {},
                        "fallback_reason": "stage1_input_builder_error",
                    }
                    analysis_prompt = str(stage1_input.get("prompt_text") or pdf_text)

                stage1_user_content = stage1_input.get("user_message_content")
                stage1_input_snapshot = {
                    key: value
                    for key, value in stage1_input.items()
                    if key not in {"prompt_text", "user_message_content"}
                }
                preprocess_metadata["visual_artifact_manifest_path"] = str(stage1_input.get("visual_manifest_path") or "")
                preprocess_metadata["visual_bundle_path"] = str(stage1_input.get("visual_bundle_path") or "")
                preprocess_metadata["selected_visual_count"] = len(stage1_input.get("selected_visual_refs") or [])
                preprocess_metadata["stage1_input_mode"] = str(stage1_input.get("input_mode") or "text_only")
                preprocess_metadata["stage1_input_fallback_reason"] = str(stage1_input.get("fallback_reason") or "")

                # 调用AI接口生成摘要（自动处理引擎切换）
                reader_result = self._call_stage1_reader_with_scheduler(
                    analysis_prompt,
                    reader_api_config,
                    backup_api_config,
                    user_content=stage1_user_content,
                )
                ai_result = reader_result.get("content")
                model_used = str(reader_result.get("engine_type") or "primary")
                
                if not ai_result:
                    api_message = str(reader_result.get("message") or reader_result.get("error_kind") or "")
                    failure_reason = "AI摘要生成失败"
                    if api_message:
                        failure_reason = f"{failure_reason}: {api_message}"
                    self.logger.warning(f"策略 {strategy} 失败: {failure_reason}")
                    record_attempt_failure(
                        strategy,
                        preprocess_metadata,
                        model_name=model_used,
                        quality_reason=failure_reason,
                        extractor_name=extractor_used,
                    )
                    continue
                
                self.logger.success("AI摘要生成成功")
                
                # =================== CONTENT QUALITY CHECK ===================
                # 使用新的上下文管理模块进行质量检查，如果质量不达标则标记为失败
                
                # 构建模拟的ProcessingResult对象用于质量检查
                try:
                    updated_fields = self._apply_ai_metadata_backfill(paper, ai_result)
                    if updated_fields:
                        self.logger.info(f"已从 AI 结果回填元数据字段: {', '.join(updated_fields)}")
                    else:
                        self.logger.info("未检测到可回填的论文元数据，继续保留现有 paper_info。")
                except Exception as e:
                    self.logger.warning(f"元数据回填失败: {e}")

                try:
                    synced_fields = self._apply_paper_metadata_to_ai_summary(paper, ai_result)
                    if synced_fields:
                        self.logger.info(f"Synced paper_info metadata into AI summary: {', '.join(synced_fields)}")
                except Exception as e:
                    self.logger.warning(f"AI summary metadata sync failed: {e}")

                temp_result: Dict[str, Any] = {
                    'paper_info': paper,
                    'status': 'success',
                    'ai_summary': ai_result,
                    'source_mode': self.mode,
                }
                
                # 使用context_manager的质量检查功能
                is_quality_ok, quality_reason = validate_summary_quality(temp_result)
                
                if not is_quality_ok:
                    if self._is_metadata_only_quality_failure(quality_reason):
                        resolved_fields = self._resolve_stage1_metadata_for_quality(paper, ai_result, pdf_text)
                        if resolved_fields:
                            self.logger.info(f"Metadata-only quality issue resolved fields: {', '.join(resolved_fields)}")
                        temp_result = {
                            'paper_info': paper,
                            'status': 'success',
                            'ai_summary': ai_result,
                            'source_mode': self.mode,
                        }
                        is_quality_ok, quality_reason = validate_summary_quality(temp_result)
                        if is_quality_ok:
                            self._mark_metadata_manual_review_if_missing(
                                paper,
                                ai_result,
                                "metadata unresolved after local resolution",
                            )
                            self.logger.info("Metadata resolution fixed Stage 1 quality check")
                            strategy_succeeded = True
                            break
                        if self._is_metadata_only_quality_failure(quality_reason):
                            missing_fields = self._mark_summary_metadata_manual_review(paper, ai_result, quality_reason)
                            self.logger.warning(
                                "Stage 1 summary body is usable but metadata remains incomplete; "
                                f"saved with manual-review flag: {', '.join(missing_fields) or quality_reason}"
                            )
                            strategy_succeeded = True
                            break
                    # 🚨 内容质量检查失败，尝试备用引擎
                    failure_reason = f"AI生成内容为空或不完整: {quality_reason}"
                    self.logger.warning(f"策略 {strategy} 主引擎质量检查失败: {failure_reason}")
                    
                    # 检查是否配置了备用引擎
                    backup_api_key = backup_api_config.get('api_key', '')
                    if backup_api_key and backup_api_key.strip():
                        self.logger.info("主引擎内容质量检查失败，尝试备用引擎...")
                        
                        # 使用备用引擎直接调用（绕过主引擎）
                        backup_reader_result = self._call_stage1_reader_with_scheduler(
                            analysis_prompt,
                            reader_api_config,
                            backup_api_config,
                            user_content=stage1_user_content,
                            skip_engines={model_used},
                        )
                        backup_result = backup_reader_result.get("content")
                        backup_model_used = str(backup_reader_result.get("engine_type") or "backup")
                        
                        if backup_result:
                            self.logger.success("备用引擎AI摘要生成成功")
                            
                            # 检查备用引擎结果的质量
                            try:
                                updated_fields = self._apply_ai_metadata_backfill(paper, backup_result)
                                if updated_fields:
                                    self.logger.info(f"已从备用 AI 结果回填元数据字段: {', '.join(updated_fields)}")
                            except Exception as e:
                                self.logger.warning(f"备用 AI 元数据回填失败: {e}")

                            temp_result_backup: Dict[str, Any] = {
                                'paper_info': paper,
                                'status': 'success',
                                'ai_summary': backup_result,
                                'source_mode': self.mode,
                            }
                            
                            is_quality_ok_backup, quality_reason_backup = validate_summary_quality(temp_result_backup)
                            
                            if is_quality_ok_backup:
                                self.logger.info("备用引擎内容质量检查通过")
                                ai_result = backup_result  # 使用备用引擎的结果
                                # 继续后续处理
                                model_used = backup_model_used
                                self._mark_metadata_manual_review_if_missing(
                                    paper,
                                    ai_result,
                                    "metadata unresolved after local resolution",
                                )
                                strategy_succeeded = True
                                break
                            else:
                                self.logger.warning(f"备用引擎内容质量检查也失败: {quality_reason_backup}")
                                # 备用引擎也失败，记录尝试并继续下一个策略
                                record_attempt_failure(
                                    strategy,
                                    preprocess_metadata,
                                    model_name=backup_model_used,
                                    quality_reason=f"主引擎: {quality_reason}; 备用引擎: {quality_reason_backup}",
                                    extractor_name=extractor_used,
                                )
                                continue
                        else:
                            self.logger.error("备用引擎AI摘要生成失败")
                            # 记录尝试并继续下一个策略
                            record_attempt_failure(
                                strategy,
                                preprocess_metadata,
                                model_name=backup_model_used,
                                quality_reason=f"主引擎: {quality_reason}; 备用引擎调用失败",
                                extractor_name=extractor_used,
                            )
                            continue
                    else:
                        # 没有配置备用引擎，记录尝试并继续下一个策略
                        self.logger.info("未配置备用引擎，尝试下一个预处理策略")
                        record_attempt_failure(
                            strategy,
                            preprocess_metadata,
                            model_name='primary',
                            quality_reason=quality_reason,
                            extractor_name=extractor_used,
                        )
                        continue
                else:
                    # 质量检查通过
                    unresolved_metadata = self._mark_metadata_manual_review_if_missing(
                        paper,
                        ai_result,
                        "metadata unresolved after local resolution",
                    )
                    if unresolved_metadata:
                        self.logger.warning(
                            "Stage 1 summary body passed but metadata remains incomplete; "
                            f"saved with manual-review flag: {', '.join(unresolved_metadata)}"
                        )
                    strategy_succeeded = True
                    break
            
            # 检查是否所有策略都失败了
            if not strategy_succeeded:
                # 所有策略都失败，返回详细的失败报告
                failure_reason = "所有预处理策略都失败"
                self.logger.error(failure_reason)
                
                # 构建详细的失败报告
                detailed_reason = f"所有预处理策略都失败: \n"
                for i, attempt in enumerate(attempt_history):
                    detailed_reason += (
                        f"  尝试 {i+1} (策略: {attempt['preprocess_strategy']}, "
                        f"Profile: {attempt.get('preprocess_profile', '')}, "
                        f"Parser: {attempt.get('parser_mode', '')}, "
                        f"提取器: {attempt['extractor_used']}, "
                        f"Source: {attempt.get('selected_text_source', '')}, "
                        f"Quality: {attempt.get('stage1_quality_level', '')}, "
                        f"MinerU: requested={attempt.get('mineru_remote_requested')}, "
                        f"enabled={attempt.get('mineru_remote_enabled')}, "
                        f"attempted={attempt.get('mineru_attempted')}, "
                        f"succeeded={attempt.get('mineru_succeeded')}, "
                        f"route={attempt.get('mineru_route', '')}, "
                        f"模型: {attempt['model_used']}): {attempt['quality_reason']}\n"
                    )
                
                failed_result: ProcessingResult = {
                    'paper_info': paper,
                    'status': 'failed',
                    'failure_reason': detailed_reason,
                    'attempt_history': attempt_history
                }
                return failed_result
            
            self.logger.info("内容质量检查通过")
            # ================================================================
            
            # 概念增强分析（如果启用）
            if self.concept_mode and self.concept_profile and ai_result:
                self.logger.info(f"正在对 '{paper.get('title', '未知标题')}' 进行概念增强分析...")
                
                # 读取概念分析提示词模板
                try:
                    with open('prompts/prompt_concept_analysis.txt', 'r', encoding='utf-8') as f:
                        concept_prompt_template = f.read()
                    self.logger.success(f"加载概念分析提示词模板: {len(concept_prompt_template)}字符")
                except Exception as e:
                    self.logger.warning(f"无法加载概念分析提示词模板，使用默认提示词: {e}")
                    concept_prompt_template = "基于提供的背景概念信息和论文摘要，分析该论文在背景概念发展中的作用。\n\n【背景概念】\n{{CONCEPT_PROFILE}}\n\n【论文摘要】\n{{PAPER_SUMMARY}}"
                
                # 准备概念分析的提示词
                concept_prompt = concept_prompt_template.replace(
                    '{{CONCEPT_PROFILE}}', json.dumps(self.concept_profile, ensure_ascii=False)
                ).replace(
                    '{{PAPER_SUMMARY}}', json.dumps(ai_result, ensure_ascii=False)
                )
                
                # 获取写作引擎的 API 配置
                writer_api_config: APIConfig = get_writer_api_config(self.config)

                # 调用概念分析接口
                concept_analysis_result = get_concept_analysis(concept_prompt, writer_api_config, logger=self.logger, config=self.config)
                
                if concept_analysis_result:
                    # 将概念分析结果合并到最终的摘要中
                    ai_result['concept_analysis'] = concept_analysis_result
                    self.logger.success("概念增强分析成功。")
                else:
                    self.logger.warning("概念增强分析失败。")

            # =================== METADATA BACKFILL ===================
            # AI提取的元数据回填到paper_info中，解决Direct PDF Mode下的元数据显示问题
            try:
                updated_fields = self._apply_ai_metadata_backfill(paper, ai_result)
                if updated_fields:
                    self.logger.info(f"已从 AI 结果回填元数据字段: {', '.join(updated_fields)}")
                else:
                    self.logger.info("未检测到可回填的论文元数据，继续保留现有 paper_info。")
                if False:
                    common_core = ai_result['common_core']
                    
                    # 提取AI分析出的元数据
                    extracted_title = common_core.get('title', '').strip()
                    extracted_authors = common_core.get('authors', [])
                    extracted_year = common_core.get('year', '').strip()
                    extracted_journal = common_core.get('journal', '').strip()
                    extracted_doi = common_core.get('doi', '').strip()
                    
                    # 验证提取的元数据是否有效（非空且不是"未知"等占位符）
                    valid_title = extracted_title and extracted_title not in ['', '未知', 'N/A', '无标题']
                    valid_year = extracted_year and extracted_year not in ['', '未知', 'N/A', '未知年份']
                    valid_journal = extracted_journal and extracted_journal not in ['', '未知', 'N/A', '未知期刊']
                    
                    # 更新paper_info中的元数据字段
                    if valid_title:
                        paper['title'] = extracted_title
                    
                    # 处理authors字段：可能是字符串或列表
                    if extracted_authors:
                        if isinstance(extracted_authors, list):
                            # 如果是列表，直接使用
                            if extracted_authors:  # 确保列表不为空
                                paper['authors'] = extracted_authors
                        elif isinstance(extracted_authors, str):
                            # 如果是字符串，尝试分割为列表
                            authors_str = extracted_authors.strip()
                            if authors_str and authors_str not in ['', '未知', 'N/A']:
                                # 尝试按常见分隔符分割
                                import re
                                authors_list = re.split(r'[,，、;；和and]\s*', authors_str)
                                authors_list = [author.strip() for author in authors_list if author.strip()]
                                if authors_list:
                                    paper['authors'] = authors_list
                    
                    # 更新年份和期刊信息
                    if valid_year:
                        paper['year'] = extracted_year
                    
                    if valid_journal:
                        paper['journal'] = extracted_journal
                    
                    # 更新DOI（如果有的话）
                    if extracted_doi:
                        paper['doi'] = extracted_doi
                    
                    # 记录元数据更新情况
                    updated_fields: List[str] = []
                    if valid_title:
                        updated_fields.append('标题')
                    if extracted_authors:
                        updated_fields.append('作者')
                    if valid_year:
                        updated_fields.append('年份')
                    if valid_journal:
                        updated_fields.append('期刊')
                    if extracted_doi:
                        updated_fields.append('DOI')
                    
                    if updated_fields:
                        self.logger.info(f"✅ 元数据回填成功，更新字段: {', '.join(updated_fields)}")
                    else:
                        self.logger.info("ℹ️  未发现有效的AI提取元数据，使用默认值")
                        
            except Exception as e:
                self.logger.warning(f"元数据回填失败: {e}")
            # =============================================================

            # 构造结果
            result: ProcessingResult = {
                'paper_info': paper,
                'status': 'success',
                'ai_summary': ai_result,  # type: ignore
                'processing_time': datetime.now().isoformat(),
                'text_length': len(pdf_text) if pdf_text else 0,  # type: ignore
                'preprocess': preprocess_metadata,
                'stage1_input': stage1_input_snapshot,
                'source_mode': self.mode,
                'model_used': model_used,
                'attempt_history': attempt_history
            }

            if not self._persist_paper_artifact(result):
                return {
                    'paper_info': paper,
                    'status': 'failed',
                    'failure_reason': 'paper_artifact_v1 persistence failed',
                    'attempt_history': attempt_history
                }
            
            return result
            
        except Exception as e:
            failure_reason = f"处理论文时发生异常: {str(e)}"
            self.logger.error(failure_reason)
            traceback.print_exc()
            failed_result: ProcessingResult = {
                'paper_info': paper,
                'status': 'failed',
                'failure_reason': failure_reason
            }
            return failed_result
    
    def save_summaries(self) -> bool:
        """保存摘要到JSON文件（线程安全版本）"""
        try:
            if not self.output_dir or not self.summary_file:
                self.logger.error("输出目录或摘要文件路径未设置")
                return False
            
            # 确保输出目录存在
            if not ensure_dir(self.output_dir):
                self.logger.error(f"无法创建输出目录: {self.output_dir}")
                return False
            
            # 创建备份文件（如果原文件存在）
            if os.path.exists(self.summary_file):
                backup_file = f"{self.summary_file}.backup"
                try:
                    import shutil
                    shutil.copy2(self.summary_file, backup_file)
                    self.logger.debug(f"已更新摘要文件备份: {backup_file}")
                except Exception as e:
                    self.logger.debug(f"无法更新摘要文件备份: {e}")
            
            # 使用线程锁确保线程安全（如果存在）
            if hasattr(self, 'save_lock'):
                with self.save_lock:
                    atomic_write_json(self.summary_file, self.summaries)
            else:
                # 无锁版本（向后兼容）
                atomic_write_json(self.summary_file, self.summaries)

            self._register_workspace_artifact(
                artifact_role="summary",
                artifact_type="summary_file",
                artifact_version="v1",
                path=self.summary_file,
                producer="main.LiteratureReviewGenerator.save_summaries",
            )
            self._write_stage1_progress_snapshot()
            
            self.logger.debug(f"[保存] 摘要文件已更新: {len(self.summaries)}条记录")
            return True
            
        except Exception as e:
            self.logger.error(f"保存摘要文件失败: {e}")
            self.logger.error(f"摘要列表类型: {type(self.summaries)}")
            self.logger.error(f"摘要列表长度: {len(self.summaries)}")
            
            # 尝试保存错误报告
            try:
                error_file = f"{self.summary_file}.error" if self.summary_file else None
                if error_file:
                    with open(error_file, 'w', encoding='utf-8') as f:
                        json.dump({
                            'error': str(e),
                            'timestamp': datetime.now().isoformat(),
                            'summaries_count': len(self.summaries),
                            'summaries_type': str(type(self.summaries))
                        }, f, ensure_ascii=False, indent=2)
            except:
                pass
            return False
    
    def generate_excel_report(self) -> bool:
        """生成Excel报告"""
        try:
            if not self.output_dir or not self.project_name:
                return False
            
            excel_file = self._get_report_file_path('_analyzed_papers.xlsx')
            
            # 生成Excel报告
            success = generate_excel_report(self)
            
            if success:
                self.logger.success(f"Excel报告已生成: {excel_file}")
                return True
            else:
                self.logger.error("Excel报告生成失败")
                return False
                
        except Exception as e:
            self.logger.error(f"生成Excel报告失败: {e}")
            return False
    
    def generate_failure_report(self) -> bool:
        """生成失败报告"""
        try:
            if not self.output_dir or not self.project_name:
                return False
            
            failure_report_file = self._get_report_file_path('_failed_papers_report.txt')
            
            # 生成失败报告
            success = generate_failure_report(self)
            
            if success:
                self.logger.success(f"失败报告已生成: {failure_report_file}")
                return True
            else:
                self.logger.error("失败报告生成失败")
                return False
                
        except Exception as e:
            self.logger.error(f"生成失败报告失败: {e}")
            return False
    
    def generate_retry_zotero_report(self) -> bool:
        """生成Zotero重跑报告（仅Zotero模式）"""
        try:
            if self.mode != "zotero":
                return True  # 直接模式下不需要生成重跑报告
            
            if not self.output_dir or not self.project_name:
                return False
            
            # 类型守卫：确保output_dir和project_name不是None
            assert self.output_dir is not None and self.project_name is not None
            
            retry_report_file = self._get_report_file_path('_zotero_report_for_retry.txt')
            
            # 生成重跑报告
            success = generate_retry_zotero_report(self)
            
            if success:
                self.logger.success(f"重跑报告已生成: {retry_report_file}")
                return True
            else:
                self.logger.error("重跑报告生成失败")
                return False
                
        except Exception as e:
            self.logger.error(f"生成重跑报告失败: {e}")
            return False
    
    def _paper_reuse_entry(self, paper: Mapping[str, Any], *, paper_key: str) -> Dict[str, Any]:
        return {
            "paper_key": paper_key,
            "title": str(paper.get("title") or ""),
            "doi": normalize_doi(paper.get("doi")),
            "canonical_paper_key": str(paper.get("canonical_paper_key") or paper_key),
        }

    def _build_stage1_reuse_report(self, sources: List[SummarySource], rejected_candidates: List[Dict[str, Any]]) -> Dict[str, Any]:
        source_items = [
            {
                "path": source.path,
                "source_type": source.source_type,
                "label": source.label,
                "priority": source.priority,
            }
            for source in sources
        ]
        return {
            "artifact_type": "summary_reuse_report",
            "artifact_version": "v2",
            "created_at": datetime.utcnow().replace(microsecond=0).isoformat() + "Z",
            "project_name": self.project_name or "",
            "mode": self.mode,
            "summary_file": self.summary_file or "",
            "configured_reuse_sources": source_items,
            "source_items": source_items,
            "rejected_candidates": list(rejected_candidates),
            "reused_papers": [],
            "not_reused": [],
            "skipped_papers": [],
            "newly_analyzed_papers": [],
            "failed_papers": [],
            "degraded_artifacts": [],
            "preview": {
                "source_count": len(source_items),
                "reusable_record_count": 0,
                "matched_count": 0,
                "ambiguous_count": 0,
                "needs_analysis_count": 0,
            },
        }

    def _log_stage1_reuse_preview(self, report: Mapping[str, Any]) -> None:
        preview = report.get("preview", {}) if isinstance(report, Mapping) else {}
        source_count = int(preview.get("source_count", 0) or 0)
        reusable_record_count = int(preview.get("reusable_record_count", 0) or 0)
        matched_count = int(preview.get("matched_count", 0) or 0)
        ambiguous_count = int(preview.get("ambiguous_count", 0) or 0)
        needs_analysis_count = int(preview.get("needs_analysis_count", 0) or 0)
        self.logger.info(
            f"[stage1 reuse] Preview: {source_count} sources, {reusable_record_count} reusable summaries, "
            f"{matched_count} direct matches, {ambiguous_count} ambiguous, {needs_analysis_count} still need analysis"
        )

    def _apply_stage1_cross_run_reuse(self) -> bool:
        self._stage1_reuse_report = None
        self._stage1_reused_paper_keys = set()

        if not self.reuse_stage1:
            return True

        paths_config: Dict[str, str] = self.config.get('Paths', {}) if self.config else {}
        output_root = paths_config.get('output_path', './output')
        sources = collect_summary_sources(
            explicit_paths=self.reuse_summary_files,
            output_root=output_root,
            current_workspace_root=self.output_dir,
            current_summary_file=self.summary_file,
        )
        self.logger.info(f"[stage1 reuse] Found {len(sources)} candidate summary sources")

        try:
            catalog = SummaryCatalog.from_sources(sources, logger=self.logger)
        except SummarySourceError as exc:
            self.logger.error(f"[stage1 reuse] Failed to load an explicit reuse source: {exc}")
            return False

        report = self._build_stage1_reuse_report(sources, catalog.rejected_candidates)
        preview = report["preview"]
        preview["reusable_record_count"] = len(catalog.records)
        reused_any = False

        for paper in self.papers:
            paper_key = LiteratureReviewGenerator.get_paper_key(paper)  # type: ignore[arg-type]
            paper_entry = self._paper_reuse_entry(paper, paper_key=paper_key)
            if paper_key in self._checkpoint_processed_papers or paper_key in self._checkpoint_failed_papers:
                report["skipped_papers"].append(
                    {
                        **paper_entry,
                        "reason": "already_processed" if paper_key in self._checkpoint_processed_papers else "already_failed",
                    }
                )
                continue

            match: Optional[SummaryMatch] = catalog.resolve_for_paper(paper)
            if match is None:
                report["not_reused"].append({**paper_entry, "reason": "no_matching_summary"})
                preview["needs_analysis_count"] += 1
                continue

            if match.is_ambiguous:
                report["not_reused"].append(
                    {
                        **paper_entry,
                        "reason": "ambiguous_match",
                        "match_type": match.match_type,
                        "ambiguous_candidates": [
                            {
                                **describe_summary_candidate(candidate.summary),
                                "source_path": candidate.source.path,
                                "source_type": candidate.source.source_type,
                            }
                            for candidate in match.ambiguous_candidates
                        ],
                    }
                )
                preview["ambiguous_count"] += 1
                preview["needs_analysis_count"] += 1
                continue

            if match.winner is None:
                report["not_reused"].append({**paper_entry, "reason": "no_matching_summary"})
                preview["needs_analysis_count"] += 1
                continue

            canonical_doi = normalize_doi(paper.get("doi"))
            reused_summary = build_reused_summary(
                current_paper=paper,
                matched_summary=match.winner.summary,
                reuse_source=match.winner.source,
                match_type=match.match_type,
                canonical_doi=canonical_doi,
            )

            replaced = False
            for index, existing_summary in enumerate(self.summaries):
                existing_paper_info = existing_summary.get('paper_info', {})
                if LiteratureReviewGenerator.get_paper_key(existing_paper_info) == paper_key:
                    self.summaries[index] = reused_summary  # type: ignore[index]
                    replaced = True
                    break
            if not replaced:
                self.summaries.append(reused_summary)  # type: ignore[arg-type]

            self._checkpoint_processed_papers.add(paper_key)
            self._checkpoint_failed_papers.discard(paper_key)
            self._stage1_reused_paper_keys.add(paper_key)
            reused_any = True

            artifact_ready = self._persist_paper_artifact(reused_summary)
            report_entry = {
                **paper_entry,
                "reason": match.match_type,
                "match_type": match.match_type,
                "winner_source": {
                    "path": match.winner.source.path,
                    "source_type": match.winner.source.source_type,
                    "label": match.winner.source.label,
                },
                "source_path": match.winner.source.path,
                "source_type": match.winner.source.source_type,
                "ambiguous_candidates": [],
                "artifact_ready": artifact_ready,
            }
            report["reused_papers"].append(report_entry)
            if not artifact_ready:
                report["degraded_artifacts"].append(report_entry)
            preview["matched_count"] += 1

        preview["needs_analysis_count"] = len(report["not_reused"])

        self._stage1_reuse_report = report
        self._log_stage1_reuse_preview(report)

        if reused_any:
            self.save_summaries()
            self.save_checkpoint()
            self.logger.info(f"[stage1 reuse] Reused {len(report['reused_papers'])} papers via global summary catalog")
        else:
            self.logger.info("[stage1 reuse] No reusable summaries matched the current paper set")

        return True

    def _persist_stage1_reuse_report(self) -> bool:
        if not self._stage1_reuse_report:
            return True

        report = dict(self._stage1_reuse_report)
        current_paper_keys = {
            LiteratureReviewGenerator.get_paper_key(paper): paper
            for paper in self.papers
        }

        newly_analyzed: List[Dict[str, Any]] = []
        for summary in self.summaries:
            if str(summary.get('status') or '').strip().lower() != 'success':
                continue
            paper_info = summary.get('paper_info', {})
            if not isinstance(paper_info, Mapping):
                continue
            paper_key = LiteratureReviewGenerator.get_paper_key(dict(paper_info))
            if paper_key not in current_paper_keys or paper_key in self._stage1_reused_paper_keys:
                continue
            newly_analyzed.append(
                {
                    "paper_key": paper_key,
                    "title": str(paper_info.get('title') or ''),
                    "doi": normalize_doi(paper_info.get('doi')),
                }
            )

        failed_papers: List[Dict[str, Any]] = []
        for failed in self.failed_papers:
            paper_info = failed.get('paper_info', {})
            if not isinstance(paper_info, Mapping):
                continue
            paper_key = LiteratureReviewGenerator.get_paper_key(dict(paper_info))
            if paper_key not in current_paper_keys or paper_key in self._stage1_reused_paper_keys:
                continue
            failed_papers.append(
                {
                    "paper_key": paper_key,
                    "title": str(paper_info.get('title') or ''),
                    "doi": normalize_doi(paper_info.get('doi')),
                    "failure_reason": str(failed.get('failure_reason') or ''),
                }
            )

        report["newly_analyzed_papers"] = newly_analyzed
        report["failed_papers"] = failed_papers
        report["counts"] = {
            "configured_reuse_sources": len(report.get("configured_reuse_sources", [])),
            "reused": len(report.get("reused_papers", [])),
            "not_reused": len(report.get("not_reused", [])),
            "skipped": len(report.get("skipped_papers", [])),
            "newly_analyzed": len(newly_analyzed),
            "failed": len(failed_papers),
        }

        report_path = self._get_summary_reuse_report_path()
        atomic_write_json(report_path, report)
        self._register_workspace_artifact(
            artifact_role="summary_reuse",
            artifact_type="summary_reuse_report",
            artifact_version="v1",
            path=report_path,
            producer="main.LiteratureReviewGenerator.run_stage_one",
            depends_on=[
                ArtifactDependencyRef(artifact_type="summary_file", path=self.summary_file or ""),
            ] if self.summary_file else [],
        )
        self._stage1_reuse_report = report
        return True

    def process_all_papers(self) -> bool:
        """处理所有论文（并发处理版本）"""
        try:
            self._check_cancelled()
            if not self.papers:
                self.logger.error("没有论文需要处理")
                return False
            
            total_papers = len(self.papers)
            self.logger.info(f"开始并发处理 {total_papers} 篇论文")
            
            # 确定最大工作线程数
            performance_config: Dict[str, str] = self.config.get('Performance', {}) if self.config else {}
            max_workers = int(performance_config.get('max_workers', 3))
            self.logger.info(f"使用 {max_workers} 个工作线程")
            
            # 创建文件索引（Zotero模式）
            file_index: Optional[FileIndex] = None
            if self.mode == "zotero":
                paths_config: Dict[str, str] = self.config.get('Paths', {}) if self.config else {}
                library_path: str = paths_config.get('library_path', '')
                if library_path:
                    self.logger.info("正在创建文件索引...")
                    file_index = create_file_index(library_path)
                    self.logger.success(f"文件索引创建完成，包含 {len(file_index)} 个文件")
            
            # 确定需要处理的论文（跳过已处理的）
            papers_to_process: List[Tuple[int, 'PaperInfo']] = []
            skipped_count = 0
            
            for i, paper in enumerate(self.papers):
                self._check_cancelled()
                paper_key = LiteratureReviewGenerator.get_paper_key(paper)  # type: ignore
                if paper_key in self._checkpoint_processed_papers:
                    skipped_count += 1
                    continue
                papers_to_process.append((i, paper))
            
            self.logger.info(f"需要处理: {len(papers_to_process)}篇论文，跳过: {skipped_count}篇论文")
            
            if not papers_to_process:
                self.logger.success("所有论文都已处理完成")
                self._emit_stage1_progress(
                    total=0,
                    current=0,
                    success_count=0,
                    failure_count=0,
                    message="所有论文都已处理完成",
                )
                return True
            
            # 重置计数器
            self.reset_counters()
            self.reset_stage1_reader_engine_round_state()
            run_success_count = 0
            run_failure_count = 0
            tracked_total = len(papers_to_process)
            self._emit_stage1_progress(
                total=tracked_total,
                current=0,
                success_count=0,
                failure_count=0,
                message=f"开始并发处理 {tracked_total} 篇论文",
            )
            
            # 创建进度条
            progress_bar = tqdm(total=len(papers_to_process), desc="[阶段一] 正在分析文献")
            
            # 创建线程池并提交任务
            with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
                # 提交所有任务
                future_to_paper: Dict[concurrent.futures.Future['ProcessingResult | None'], Tuple[int, 'PaperInfo']] = {
                    executor.submit(self.process_paper, paper, i, file_index, total_papers): (i, paper)
                    for i, paper in papers_to_process
                }
                
                # 处理完成的任务
                for future in concurrent.futures.as_completed(future_to_paper):
                    self._check_cancelled()
                    _, paper = future_to_paper[future]
                    paper_key = LiteratureReviewGenerator.get_paper_key(paper)  # type: ignore
                    
                    try:
                        result = future.result()
                        
                        if result and result.get('status') == 'success':
                            # 处理成功
                            with self.save_lock:
                                self.summaries.append(result)
                                self._checkpoint_processed_papers.add(paper_key)
                            
                            # 线程安全地增加计数器
                            with self.save_lock:
                                self.processed_count.increment()
                            run_success_count += 1
                            
                            # 更新进度条
                            progress_bar.update(1)
                            completed_label = self._paper_progress_label(result.get('paper_info', paper))
                            self._emit_stage1_progress(
                                total=tracked_total,
                                current=run_success_count + run_failure_count,
                                success_count=run_success_count,
                                failure_count=run_failure_count,
                                item_label=completed_label,
                                message=f"已完成 {run_success_count + run_failure_count}/{tracked_total}: {completed_label}",
                            )
                            # 更新进度条的后缀信息
                            progress_bar.set_postfix_str(f"成功: {self.processed_count.value}, 失败: {self.failed_count.value}")
                        else:
                            # 处理失败
                            failure_reason = result.get('failure_reason') or '未知错误' if result else '处理返回空结果'
                            if not isinstance(failure_reason, str):  # type: ignore
                                failure_reason = '未知错误'
                            failed_paper = result.get('paper_info', paper) if result else paper
                            
                            self.failed_papers.append({  # type: ignore
                                    'paper_info': failed_paper,
                                    'failure_reason': failure_reason
                                })
                                # 更新身份基断点跟踪
                            self._checkpoint_failed_papers.add(paper_key)
                            
                            # 线程安全地增加计数器
                            with self.save_lock:
                                self.failed_count.increment()
                            run_failure_count += 1
                            failed_label = self._paper_progress_label(failed_paper)
                            self._emit_stage1_progress(
                                total=tracked_total,
                                current=run_success_count + run_failure_count,
                                success_count=run_success_count,
                                failure_count=run_failure_count,
                                item_label=failed_label,
                                message=f"处理失败 {run_success_count + run_failure_count}/{tracked_total}: {failed_label} | {self._short_progress_message(failure_reason)}",
                            )
                            
                            # 更新进度条
                            progress_bar.update(1)
                            # 更新进度条的后缀信息
                            progress_bar.set_postfix_str(f"成功: {self.processed_count.value}, 失败: {self.failed_count.value}")
                        
                        # 每完成一个任务就立即保存数据，确保数据不丢失
                        if result and result.get('status') == 'success':
                            save_result = self.save_summaries()
                            if not save_result:
                                self.logger.error("⚠️ 警告: 数据保存失败，请检查磁盘空间和权限")
                        else:
                            # 失败的情况下定期保存（每3个失败保存一次）
                            if (self.processed_count.get_value() + self.failed_count.get_value()) % 3 == 0:
                                self.save_summaries()
                                self.save_checkpoint()
                        
                    except Exception as e:
                        # 任务执行异常
                        failure_reason = f"处理过程发生异常: {str(e)}"
                        
                        with self.save_lock:
                            self.failed_papers.append({  # type: ignore
                                'paper_info': paper,
                                'failure_reason': failure_reason
                            })
                            # 更新身份基断点跟踪
                            self._checkpoint_failed_papers.add(paper_key)
                        
                        # 线程安全地增加计数器
                        with self.save_lock:
                            self.failed_count.increment()
                        run_failure_count += 1
                        failed_label = self._paper_progress_label(paper)
                        self._emit_stage1_progress(
                            total=tracked_total,
                            current=run_success_count + run_failure_count,
                            success_count=run_success_count,
                            failure_count=run_failure_count,
                            item_label=failed_label,
                            message=f"处理异常 {run_success_count + run_failure_count}/{tracked_total}: {failed_label} | {self._short_progress_message(failure_reason)}",
                        )
                        
                        self.logger.error(f"任务执行异常: {e}")
                        self.logger.error(f"失败: {self.processed_count.value}成功, {self.failed_count.value}失败 - {failure_reason}")
                        
                        # 异常情况下立即保存，确保数据不丢失
                        save_result = self.save_summaries()
                        if not save_result:
                            self.logger.error("⚠️ 警告: 异常情况下数据保存失败")
                        self.save_checkpoint()
            
            # 最终保存所有数据
            self.save_summaries()
            self.save_checkpoint()
            self._emit_stage1_progress(
                total=tracked_total,
                current=run_success_count + run_failure_count,
                success_count=run_success_count,
                failure_count=run_failure_count,
                item_label=self._paper_progress_label(papers_to_process[-1][1]) if papers_to_process else "",
                message=f"阶段一完成：成功 {run_success_count}，失败 {run_failure_count}",
            )
            
            self.logger.success("\n并发处理完成！")
            self.logger.info(f"总文献数: {total_papers}")
            self.logger.info(f"本次处理: {len(papers_to_process)}篇")
            self.logger.info(f"跳过已处理: {skipped_count}篇")
            self.logger.info(f"成功处理: {self.processed_count.value}")
            self.logger.info(f"失败: {self.failed_count.value}")
            self.logger.info(f"摘要文件: {self.summary_file}")
            
            # 自动重试循环 - 第一阶段末尾
            if self.failed_papers:
                self.logger.warning(f"有{len(self.failed_papers)}篇论文处理失败，启动自动重试循环...")
                
                # 读取重试配置
                retry_config: Dict[str, str] = self.config.get('Retry_Settings', {}) if self.config else {}
                max_retry_rounds: int = int(retry_config.get('max_retry_rounds', 2))
                base_retry_delay: int = int(retry_config.get('base_retry_delay', 30))
                max_retry_delay: int = int(retry_config.get('max_retry_delay', 120))
                self._emit_stage1_progress(
                    total=tracked_total,
                    current=run_success_count + run_failure_count,
                    success_count=run_success_count,
                    failure_count=run_failure_count,
                    message=f"进入自动重试，共 {max_retry_rounds} 轮",
                    retry_round=0,
                    retry_total_rounds=max_retry_rounds,
                )
                
                self.logger.info(f"🔄 重试配置: 最大重试轮数={max_retry_rounds}, 基础间隔={base_retry_delay}秒, 最大间隔={max_retry_delay}秒")
                
                # 定义可重试的失败类型关键词
                retriable_keywords = ['api', 'network', 'http', 'timeout', '500', '502', '503', '504', '429', '连接', '超时', '错误', '失败']
                
                # 分离可重试和永久失败的论文
                retriable_failures: List['FailedPaper'] = []
                permanent_failures: List['FailedPaper'] = []
                
                for failed_item in self.failed_papers:
                    failure_reason: str = failed_item.get('failure_reason', '').lower()
                    paper_info: Dict[str, Any] = failed_item.get('paper_info', {})  # type: ignore
                    
                    # 检查失败原因是否包含可重试关键词
                    is_retriable = any(keyword in failure_reason for keyword in retriable_keywords)
                    
                    if is_retriable:
                        retriable_failures.append(failed_item)
                    else:
                        permanent_failures.append(failed_item)
                
                self.logger.info(f"可重试失败论文: {len(retriable_failures)}篇")
                self.logger.info(f"永久失败论文: {len(permanent_failures)}篇")
                
                # 执行自动重试（使用配置中的参数）
                for retry_round in range(1, max_retry_rounds + 1):
                    self._check_cancelled()
                    self.reset_stage1_reader_engine_round_state()
                    self._emit_stage1_progress(
                        total=tracked_total,
                        current=run_success_count + run_failure_count,
                        success_count=run_success_count,
                        failure_count=run_failure_count,
                        message=f"开始第 {retry_round} 轮自动重试，待重试 {len(retriable_failures)} 篇",
                        retry_round=retry_round,
                        retry_total_rounds=max_retry_rounds,
                    )
                    if not retriable_failures:
                        self.logger.info("没有可重试的失败论文，结束重试循环")
                        break
                    
                    # 如果不是第一轮重试，添加延迟等待API限制恢复
                    if retry_round > 1:
                        # 使用配置的重试间隔，支持上限控制
                        calculated_delay = retry_round * base_retry_delay
                        retry_delay = min(calculated_delay, max_retry_delay)
                        self._emit_stage1_progress(
                            total=tracked_total,
                            current=run_success_count + run_failure_count,
                            success_count=run_success_count,
                            failure_count=run_failure_count,
                            message=f"第 {retry_round - 1} 轮重试后等待 {retry_delay} 秒",
                            retry_round=retry_round,
                            retry_total_rounds=max_retry_rounds,
                        )
                        self.logger.info(f"第 {retry_round-1} 轮重试失败，等待 {retry_delay} 秒让API限制恢复...")
                        self.logger.info(f"⏰ 间隔计算: {retry_round} × {base_retry_delay} = {calculated_delay}秒，已限制上限为 {max_retry_delay}秒")
                        self.logger.info("⏳ 等待中... 这有助于避免API频率限制，提高重试成功率")
                        
                        # 显示倒计时
                        for i in range(retry_delay, 0, -5):
                            if i > 5:
                                self.logger.info(f"⏰ 剩余等待时间: {i} 秒...")
                                time.sleep(5)
                            else:
                                break
                        time.sleep(retry_delay % 5)  # 完成剩余等待时间
                        self.logger.info("✅ 等待完成，开始重试...")
                    
                    self.logger.info(f"正在对 {len(retriable_failures)} 篇失败文献进行第 {retry_round} 轮自动重试...")
                    
                    # 准备重试论文数据
                    retry_papers: List[Tuple[int, Dict[str, Any]]] = []
                    retry_indices: List[int] = []
                    for failed_item in retriable_failures:
                        paper_info: Dict[str, Any] = failed_item.get('paper_info', {})  # type: ignore
                        # 找到原始论文索引

                        for i, original_paper in enumerate(self.papers):

                            if LiteratureReviewGenerator.get_paper_key(original_paper) == LiteratureReviewGenerator.get_paper_key(paper_info):
                                # 计算论文key
                                paper_key = LiteratureReviewGenerator.get_paper_key(original_paper)
                                # 关键修复：从失败集合中移除，避免被process_paper跳过
                                if paper_key in self._checkpoint_failed_papers:
                                    self._checkpoint_failed_papers.discard(paper_key)
                                    self.logger.info(f"已从失败集合中移除论文以便重试: {original_paper.get('title', '未知标题')}")
                                
                                retry_papers.append((i, original_paper))  # type: ignore  # 使用original_paper而不是paper_info
                                retry_indices.append(i)
                                break
                    
                    if not retry_papers:
                        self.logger.warning("无法找到重试论文的原始索引，结束重试")
                        break
                    
                    # 重置当前轮次的失败列表
                    current_round_failures: List[Dict[str, Any]] = []
                    
                    # 创建线程池进行重试处理
                    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as retry_executor:
                        retry_futures: Dict[concurrent.futures.Future['ProcessingResult | None'], Tuple[Dict[str, Any], int]] = {}
                        for original_index, paper in retry_papers:
                            future = retry_executor.submit(self.process_paper, paper, original_index, file_index, total_papers)  # type: ignore
                            retry_futures[future] = (paper, original_index)
                        
                        # 处理重试结果
                        retry_progress_bar = tqdm(concurrent.futures.as_completed(retry_futures), 
                                                total=len(retry_papers), desc=f"[重试第{retry_round}轮] 正在重试文献")
                        
                        for future in retry_progress_bar:
                            paper, original_index = retry_futures[future]
                            paper_key = LiteratureReviewGenerator.get_paper_key(paper)  # type: ignore
                            
                            try:
                                result = future.result()
                                if result and result.get('status') == 'success':
                                    # 重试成功，添加到结果列表
                                    with self.save_lock:
                                        self.summaries.append(result)
                                        self._checkpoint_processed_papers.add(paper_key)
                                        # 从失败集合中移除，保持状态一致性
                                        self._checkpoint_failed_papers.discard(paper_key)
                                        # 从失败列表中移除
                                        self.failed_papers = [fp for fp in self.failed_papers  # type: ignore
                                                          if LiteratureReviewGenerator.get_paper_key(fp.get('paper_info', {})) != paper_key]
                                    
                                    with self.save_lock:
                                        self.processed_count.increment()
                                        self.failed_count.decrement()
                                    run_success_count += 1
                                    run_failure_count = max(run_failure_count - 1, 0)
                                    
                                    retry_progress_bar.update(1)
                                    completed_label = self._paper_progress_label(result.get('paper_info', paper))
                                    self._emit_stage1_progress(
                                        total=tracked_total,
                                        current=run_success_count + run_failure_count,
                                        success_count=run_success_count,
                                        failure_count=run_failure_count,
                                        item_label=completed_label,
                                        message=f"重试成功 {run_success_count}/{tracked_total}: {completed_label}",
                                        retry_round=retry_round,
                                        retry_total_rounds=max_retry_rounds,
                                    )
                                    retry_progress_bar.set_postfix_str(f"成功: {self.processed_count.value}, 失败: {self.failed_count.value}")
                                else:
                                    # 重试仍然失败
                                    failure_reason = result.get('failure_reason', '重试失败') if result else '重试返回空结果'
                                    current_round_failures.append({
                                        'paper_info': paper,
                                        'failure_reason': failure_reason
                                    })
                                    failed_label = self._paper_progress_label(paper)
                                    self._emit_stage1_progress(
                                        total=tracked_total,
                                        current=run_success_count + run_failure_count,
                                        success_count=run_success_count,
                                        failure_count=run_failure_count,
                                        item_label=failed_label,
                                        message=f"重试未通过 {run_success_count + run_failure_count}/{tracked_total}: {failed_label} | {self._short_progress_message(failure_reason)}",
                                        retry_round=retry_round,
                                        retry_total_rounds=max_retry_rounds,
                                    )
                                    
                                    retry_progress_bar.update(1)
                                    retry_progress_bar.set_postfix_str(f"成功: {self.processed_count.value}, 失败: {self.failed_count.value}")
                                
                                # 重试成功时立即保存，确保数据不丢失
                                if result and result.get('status') == 'success':
                                    save_result = self.save_summaries()
                                    if not save_result:
                                        self.logger.error("⚠️ 警告: 重试成功数据保存失败")
                                else:
                                    # 失败情况下定期保存
                                    if (self.processed_count.get_value() + self.failed_count.get_value()) % 3 == 0:
                                        self.save_summaries()
                                        self.save_checkpoint()
                                
                            except Exception as e:
                                # 重试异常
                                failure_reason = f"重试过程发生异常: {str(e)}"
                                current_round_failures.append({
                                    'paper_info': paper,
                                    'failure_reason': failure_reason
                                })
                                failed_label = self._paper_progress_label(paper)
                                self._emit_stage1_progress(
                                    total=tracked_total,
                                    current=run_success_count + run_failure_count,
                                    success_count=run_success_count,
                                    failure_count=run_failure_count,
                                    item_label=failed_label,
                                    message=f"重试异常 {run_success_count + run_failure_count}/{tracked_total}: {failed_label} | {self._short_progress_message(failure_reason)}",
                                    retry_round=retry_round,
                                    retry_total_rounds=max_retry_rounds,
                                )
                                
                                self.logger.error(f"重试任务执行异常: {e}")
                                # 重试异常时立即保存，确保数据不丢失
                                save_result = self.save_summaries()
                                if not save_result:
                                    self.logger.error("⚠️ 警告: 重试异常时数据保存失败")
                                self.save_checkpoint()
                    
                    # 更新重试失败列表
                    retriable_failures = current_round_failures  # type: ignore
                    
                    if current_round_failures:
                        self.logger.warning(f"第 {retry_round} 轮重试后，仍有 {len(current_round_failures)} 篇论文失败")
                    else:
                        self.logger.success(f"第 {retry_round} 轮重试成功，所有论文处理完成！")
                        break
                
                # 合并最终失败列表
                final_failed_papers = permanent_failures + retriable_failures
                self.failed_papers = final_failed_papers  # type: ignore
                
                # 更新失败计数
                self.failed_count.set(len(self.failed_papers))
                run_failure_count = len(self.failed_papers)
                run_success_count = max(tracked_total - run_failure_count, 0)
                self._emit_stage1_progress(
                    total=tracked_total,
                    current=run_success_count + run_failure_count,
                    success_count=run_success_count,
                    failure_count=run_failure_count,
                    message=f"自动重试完成：成功 {run_success_count}，失败 {run_failure_count}",
                    retry_round=max_retry_rounds if max_retry_rounds > 0 else 0,
                    retry_total_rounds=max_retry_rounds,
                )
                
                self.logger.info(f"🔄 自动重试循环完成！")
                self.logger.info(f"📊 使用配置: {max_retry_rounds}轮重试，基础间隔{base_retry_delay}秒，上限{max_retry_delay}秒")
                self.logger.info(f"📈 最终失败论文数: {len(self.failed_papers)}篇")
            
            # 生成失败报告
            if self.failed_papers:
                self.logger.warning(f"有{len(self.failed_papers)}篇论文处理失败，将生成失败报告")
            
            # 最终保存所有数据
            self.save_summaries()
            self.save_checkpoint()

            # Stage 1 strict success semantics
            stage1_cfg = self._get_stage1_config()
            failed_count = len(self.failed_papers)
            total_papers = len(self.papers)
            success_count = max(total_papers - failed_count, 0)
            success_ratio = success_count / total_papers if total_papers > 0 else 1.0
            allow_partial = stage1_cfg["allow_partial_success"]
            min_ratio = stage1_cfg["min_success_ratio"]

            partial_success = failed_count > 0
            stage1_ok = not partial_success or (
                allow_partial and success_ratio >= min_ratio
            )

            self.stage1_result_summary = {
                "partial_success": partial_success,
                "failed_count": failed_count,
                "success_count": success_count,
                "total": total_papers,
                "success_ratio": round(success_ratio, 4),
                "allow_partial_success": allow_partial,
                "min_success_ratio": min_ratio,
            }

            if not stage1_ok:
                self.logger.warning(
                    f"Stage 1 success check failed: {failed_count} failed papers, "
                    f"ratio={success_ratio:.2f}, allow_partial={allow_partial}, "
                    f"min_ratio={min_ratio}"
                )
                return False

            return True

        except KeyboardInterrupt:
            self.logger.error("\n\n用户中断处理")
            self.logger.info(f"已处理: {self.processed_count.value}篇文献，失败: {self.failed_count.value}篇")
            self.save_summaries()
            self.save_checkpoint()
            return False
        except Exception as e:
            self.logger.error(f"并发处理过程中出错: {e}")
            self.logger.info(f"已处理: {self.processed_count.value}篇文献，失败: {self.failed_count.value}篇")
            self.save_summaries()
            self.save_checkpoint()
            return False

    def save_checkpoint(self) -> bool:
        """保存基于身份的断点文件 - 委托给CheckpointManager"""
        return self.checkpoint_manager.save_checkpoint(self)

    def load_checkpoint(self) -> bool:
        """加载基于身份的断点文件 - 委托给CheckpointManager"""
        return self.checkpoint_manager.load_checkpoint(self)

    def run_stage_one(self, override_zotero_report_path: Optional[str] = None) -> bool:
        """阶段一：文献解析与AI摘要生成（基于身份的断点续传版本）"""
        self.logger.info("=" * 60 + "\n文献综述自动生成器 - 阶段一（身份基断点续传）\n" + "=" * 60)
        try:
            # 加载配置文件
            if not self.load_configuration(): 
                return False
            # 确保配置已正确加载到实例变量
            if not self.config:
                self.logger.error("配置未正确加载")
                return False
            
            # 如果提供了重写的Zotero报告路径，在此处应用
            if override_zotero_report_path:
                self.logger.info(f"[重跑模式] 已将文献来源强制指定为 -> {override_zotero_report_path}")
            
            if not self.setup_output_directory(): 
                return False
            
            # 加载基于身份的断点文件
            # 先尝试加载旧 checkpoint
            checkpoint_loaded = self.load_checkpoint()
            
            # 检查是否有新的 progress snapshot
            progress_snapshot_path: Optional[str] = None
            if hasattr(self, 'job_workspace') and self.job_workspace:
                snapshot_path = self.job_workspace.artifact_path("stage1_progress_snapshot.json")
                if os.path.exists(snapshot_path):
                    progress_snapshot_path = snapshot_path

            # 加载现有摘要（兼容旧版本）
            self.load_existing_summaries()
            
            if checkpoint_loaded:
                self._merge_stage1_progress_from_loaded_summaries()
                self.logger.info("[断点恢复] 成功加载旧 checkpoint，将从上次中断处继续处理")
                self.logger.info("[断点续传] 已加载处理进度，将跳过已处理的论文")
            elif progress_snapshot_path and self._restore_stage1_progress_from_snapshot(progress_snapshot_path):
                self._merge_stage1_progress_from_loaded_summaries()
                self.logger.info("[进度恢复] 找到新的 progress snapshot，将从上次中断处继续处理")
                self.logger.info("[断点续传] 已加载处理进度，将跳过已处理的论文")
            elif self._rebuild_stage1_progress_from_loaded_summaries():
                self.logger.info("[进度恢复] 未找到有效 checkpoint，但已根据现有摘要恢复处理进度")
                self.logger.info("[断点续传] 已根据摘要重建跳过集合，将跳过已处理的论文")
            else:
                self.logger.info("[全新开始] 未找到有效断点或 progress snapshot，将开始全新处理")
                
                # 确保重置断点数据
                self._checkpoint_processed_papers = set()
                self._checkpoint_failed_papers = set()
                self.processed_count.set(0)
                self.failed_count.set(0)
            
            # 逻辑分叉：根据运行模式选择数据源
            if self.mode == "zotero":
                # Zotero模式：解析Zotero报告，传递覆盖路径
                if not self.parse_zotero_report(override_zotero_report_path): 
                    return False
            else:
                # 直接模式：扫描PDF文件夹
                if not self.scan_pdf_folder(): 
                    return False
            
            # 验证论文数据完整性
            if not self.papers:
                self.logger.error("未找到任何论文数据")
                return False
            
            self.logger.info(f"Paper metadata loaded: {len(self.papers)} papers")

            if not self._apply_stage1_cross_run_reuse():
                return False
            
            # Process all papers after same-project resume and optional cross-run DOI reuse
            success = self.process_all_papers()
            if self.reuse_stage1 and not self._persist_stage1_reuse_report():
                self.logger.warning("[stage1 reuse] Failed to persist the reuse report")
            
            # If processing succeeded, continue with report generation and cleanup
            if success:
                # 清除断点文件（表示全部完成）
                if self.output_dir and self.project_name:
                    # 类型守卫：确保output_dir和project_name不是None
                    assert self.output_dir is not None and self.project_name is not None
                    checkpoint_file = self._get_stage1_checkpoint_file_path()
                    # 检查是否应保留检查点文件
                    keep_checkpoints = self._keep_checkpoints_after_completion()
                    if os.path.exists(checkpoint_file) and not keep_checkpoints:
                        try:
                            os.remove(checkpoint_file)
                            self.logger.info("已清除断点文件，所有论文处理完成")
                        except Exception as e:
                            self.logger.warning(f"无法清除断点文件: {e}")
                    elif keep_checkpoints:
                        self.logger.info("配置要求保留检查点文件，已跳过清理")
                
                # 调用统一的报告生成方法
                self.generate_all_reports()
            
            return success
            
        except Exception as e:
            self.logger.error(f"阶段一运行失败: {e}")
            # 即使失败也要保存断点
            self.save_checkpoint()
            return False

    
    



    def generate_all_reports(self) -> None:
        """生成所有分析阶段的报告 - 委托给ReportingService"""
        self.reporting_service.generate_all_reports(self)
    
    def extract_section_title_from_outline(self, outline_content: str, section_number: int) -> Optional[str]:
        """从大纲内容中提取指定章节的标题"""
        try:
            lines = outline_content.split('\n')
            current_section = 0
            
            for line in lines:
                # 查找二级标题（##）
                if line.startswith('## '):
                    current_section += 1
                    if current_section == section_number:
                        return line[3:].strip()
            
            return None
        except Exception as e:
            self.logger.error(f"提取章节标题失败: {e}")
            return None

    def create_literature_review_section(self, section_number: int, section_title: str, outline_content: str) -> bool:
        """创建文献综述的指定章节内容"""
        try:
            section_content = self.generate_review_section_content(section_title, outline_content)
            if not section_content:
                self.logger.error(f"第{section_number}章内容生成失败")
                return False
            
            # section_content应该是纯文本字符串
            if not isinstance(section_content, str):  # type: ignore
                self.logger.warning("预期收到纯文本，但收到其他格式，正在转换...")
                section_text = str(section_content)
            else:
                section_text = section_content
            
            # 生成Word文档路径（添加项目名称前缀）
            if not self.output_dir:
                self.logger.error("输出目录未设置")
                return False

            word_file = self._get_review_word_file_path()
            
            # 将章节内容追加到Word文档
            success = self.append_section_to_word_document(section_number, section_title, section_text, word_file)
            
            if success:
                self._clear_failed_review_section(section_number)
                self.logger.success(f"第{section_number}章已追加到文献综述: {word_file}")
                return True
            else:
                return False
                
        except Exception as e:
            self.logger.error(f"创建文献综述章节失败: {e}")
            return False

    def generate_review_section_content(self, section_title: str, outline_content: str) -> Optional[str]:
        """生成指定章节的内容（带智能续写循环和上下文优化）"""
        try:
            # 🆕 使用context_manager优化上下文数据
            self.logger.info("正在优化综述生成上下文...")
            
            # 优化上下文并智能截断
            # Gemini 3 Pro有1M token上下文，使用950000作为安全阈值（仅在最极端情况下触发截断）
            optimized_context: str = optimize_context_for_synthesis(
                self.summaries, 
                outline_content, 
                max_tokens=950000
            )
            
            self.logger.info(f"上下文优化完成：原始数据 -> 优化后格式")
            
            # 直接使用优化的prompt文件
            with open('prompts/optimized_prompt_synthesize_section.txt', 'r', encoding='utf-8') as f:
                prompt_template = f.read()
            
            # 替换占位符
            section_prompt: str = prompt_template.replace('{{SUMMARIES_JSON_ARRAY}}', optimized_context)
            section_prompt = section_prompt.replace('{{SECTION_TITLE}}', section_title)
            section_prompt = section_prompt.replace('{{REVIEW_OUTLINE}}', outline_content)
            section_prompt = self._inject_free_mode_context(section_prompt)

            self.logger.info(f"生成综述提示词: {len(section_prompt)}字符")

            # 提取写作引擎API配置
            writer_api_config: APIConfig = get_writer_api_config(self.config)

            self.logger.info(f"正在调用写作引擎生成章节内容: {section_title}")

            # 智能续写循环实现
            partial_section_content = ""  # 存储已生成的章节内容
            continuation_attempts = 0  # 续写计数器
            max_continuation_attempts = 5  # 最大续写次数（安全熔断）

            while continuation_attempts <= max_continuation_attempts:
                if continuation_attempts == 0:
                    # 首次调用，使用优化后的提示词
                    self.logger.info(f"[章节生成] 首次调用生成章节: {section_title}")
                    result = self._call_section_api_optimized(
                        section_prompt,
                        writer_api_config, 
                        is_continuation=False
                    )
                else:
                    # 续写调用，使用续写提示词
                    self.logger.info(f"[章节续写] 第{continuation_attempts}次续写: {section_title}")
                    result = self._call_section_api_optimized(
                        section_prompt,
                        writer_api_config, 
                        is_continuation=True,
                        partial_content=partial_section_content
                    )

                if not result:
                    self.logger.error(f"[章节生成] API调用失败，章节生成中断")
                    return None

                # 解析返回结果

                section_content = result.get('content', '')  # type: ignore

                finish_reason = result.get('finish_reason', 'stop')  # type: ignore

                # 截取 === 正文 === 后的内容
                if '=== 正文 ===' in section_content:
                    section_content = section_content.split('=== 正文 ===')[1].strip()
                # 移除其他中间文本标记
                section_content = section_content.replace('=== 逻辑规划 ===', '').replace('=== 文献矩阵 ===', '').replace('=== 核查计划 ===', '').strip()

                if not section_content or len(section_content.strip()) < 100:
                    self.logger.warning(f"[章节生成] 返回内容过短({len(section_content)}字符)，重试...")
                    continuation_attempts += 1
                    continue

                # 将新内容追加到已生成内容中
                if continuation_attempts == 0:
                    partial_section_content = section_content
                else:
                    partial_section_content += section_content

                self.logger.success(f"[章节生成] 本次生成 {len(section_content)} 字符，累计 {len(partial_section_content)} 字符")

                # 检查是否需要继续续写
                if finish_reason == 'stop':
                    self.logger.success(f"[章节生成] 章节生成完成，无需续写")
                    return partial_section_content
                elif finish_reason == 'length':
                    self.logger.info(f"[章节生成] 内容被截断，准备续写...")
                    continuation_attempts += 1
                    if continuation_attempts > max_continuation_attempts:
                        self.logger.warning(f"[章节生成] 达到最大续写次数({max_continuation_attempts})，返回部分生成的内容")
                        return partial_section_content
                else:
                    self.logger.warning(f"[章节生成] 未知的finish_reason: {finish_reason}，假设完成")
                    return partial_section_content

            # 达到最大续写次数
            self.logger.warning(f"[章节生成] 达到最大续写次数({max_continuation_attempts})，返回部分生成的内容")
            return partial_section_content

        except Exception as e:
            self.logger.error(f"生成章节内容失败: {e}")
            return None

    def _call_section_api(self, section_title: str, summaries_string: str, outline_string: str, 
                         writer_api_config: 'APIConfig', is_continuation: bool = False, 
                         partial_content: str = "") -> Optional[Dict[str, Any]]:
        """调用章节生成API的私有方法"""
        try:
            # Determine system prompt
            try:
                with open('prompts/prompt_system_section.txt', 'r', encoding='utf-8') as f:
                    system_prompt = f.read()
                self.logger.success(f"加载章节系统提示词模板: {len(system_prompt)}字符")
            except Exception as e:
                self.logger.warning(f"无法加载章节系统提示词模板，使用默认提示词: {e}")
                system_prompt = """你是一个学术文献综述专家。请基于提供的文献分析结果和完整大纲，撰写指定章节的正文内容。

要求：
1. 直接输出纯文本格式的章节正文内容
2. 不要包含章节标题
3. 内容需要专业、客观、全面
4. 适当引用具体文献以支持论点，使用结构化的 citation refs 或 token，格式为 [cite:paper_key]，其中 paper_key 是论文的唯一标识符
5. 语言风格需专业、学术
6. 只撰写指定章节的内容，不要包含其他章节
7. 不要使用传统的 (作者, 年份) 格式，只使用 [cite:paper_key] 格式的引用标记"""

            # Determine final prompt
            if is_continuation:
                try:
                    with open('prompts/prompt_continue_section.txt', 'r', encoding='utf-8') as f:
                        section_prompt_template = f.read()
                    self.logger.success(f"加载章节续写提示词模板: {len(section_prompt_template)}字符")
                except Exception as e:
                    self.logger.warning(f"无法加载章节续写提示词模板，使用默认提示词: {e}")
                    section_prompt_template = "【角色】你是一位正在撰写综述特定章节的学者，刚才思路被打断了。\n【任务】请你继续完成一份未写完的章节正文。\n\n【全部论文分析数据】\n{{SUMMARIES_JSON_ARRAY}}\n\n【综述完整大纲】\n{{REVIEW_OUTLINE}}\n\n【当前需要撰写的章节标题】\n{{SECTION_TITLE}}\n\n【已完成的章节草稿】\n{{PARTIAL_SECTION_CONTENT}}"

                final_prompt = section_prompt_template.replace('{{SUMMARIES_JSON_ARRAY}}', summaries_string)
                final_prompt = final_prompt.replace('{{REVIEW_OUTLINE}}', outline_string)
                final_prompt = final_prompt.replace('{{SECTION_TITLE}}', section_title)
                final_prompt = final_prompt.replace('{{PARTIAL_SECTION_CONTENT}}', partial_content)
            else:
                try:
                    with open('prompts/prompt_synthesize_section.txt', 'r', encoding='utf-8') as f:
                        section_prompt_template = f.read()
                    self.logger.success(f"加载章节提示词模板: {len(section_prompt_template)}字符")
                except Exception as e:
                    self.logger.warning(f"无法加载章节提示词模板，使用默认提示词: {e}")
                    section_prompt_template = "基于以下文献摘要信息和大纲，请撰写指定章节的内容。\n\n【全部论文分析数据】\n{{SUMMARIES_JSON_ARRAY}}\n\n【综述完整大纲】\n{{REVIEW_OUTLINE}}\n\n【当前需要撰写的章节标题】\n{{SECTION_TITLE}}"

                final_prompt = section_prompt_template.replace('{{SUMMARIES_JSON_ARRAY}}', summaries_string)
                final_prompt = final_prompt.replace('{{REVIEW_OUTLINE}}', outline_string)
                final_prompt = final_prompt.replace('{{SECTION_TITLE}}', section_title)

            self.logger.success(f"生成最终章节提示词: {len(final_prompt)}字符")

            # Call unified AI API function
            ai_response = _call_ai_api(
                prompt=final_prompt,
                api_config=writer_api_config,
                system_prompt=system_prompt,
                max_tokens=6000,
                temperature=0.7,
                response_format="text" # Expecting plain text
            )

            if ai_response:
                # _call_ai_api returns content directly for text format
                return {
                    'content': ai_response,
                    'finish_reason': 'stop' # _call_ai_api doesn't return finish_reason for text, assume stop
                }
            else:
                self.logger.error(f"章节内容生成失败: _call_ai_api 返回空值")
                return None

        except Exception as e:
            self.logger.error(f"调用章节API失败: {e}")
            return None

    def _call_section_api_optimized(self, section_prompt: str, writer_api_config: 'APIConfig', 
                                   is_continuation: bool = False, partial_content: str = "") -> Optional[Dict[str, Any]]:
        """🆕 优化的章节生成API调用（使用预处理的提示词）"""
        try:
            # Determine system prompt
            try:
                with open('prompts/prompt_system_section.txt', 'r', encoding='utf-8') as f:
                    system_prompt = f.read()
                self.logger.success(f"加载章节系统提示词模板: {len(system_prompt)}字符")
            except Exception as e:
                self.logger.warning(f"无法加载章节系统提示词模板，使用默认提示词: {e}")
                system_prompt = """你是一个学术文献综述专家。请基于提供的文献分析结果和完整大纲，撰写指定章节的正文内容。

要求：
1. 深度综合不同学者的观点，对比异同
2. 每个论点必须引用至少1-2篇文献，使用结构化的 citation refs 或 token，格式为 [cite:paper_key]，其中 paper_key 是论文的唯一标识符
3. 逻辑连贯，段落间有过渡
4. 避免流水账式写法，按主题组织内容
5. 不要使用传统的 (作者, 年份) 格式，只使用 [cite:paper_key] 格式的引用标记"""

            # 对于续写调用，添加续写标记
            if is_continuation and partial_content:
                continuation_prompt = f"""请继续撰写上文的章节内容。上文内容：
{partial_content}

请继续上文的内容，保持逻辑连贯，确保：
1. 与上文风格一致
2. 内容自然衔接
3. 继续深化主题分析

继续内容："""
                final_prompt = f"{section_prompt}\n\n{continuation_prompt}"
            else:
                final_prompt = section_prompt

            self.logger.success(f"生成最终章节提示词: {len(final_prompt)}字符")

            max_section_api_attempts = 3
            last_message = "empty response"
            for attempt in range(1, max_section_api_attempts + 1):
                # Call detailed API to get real finish_reason for continuation logic
                result = _call_ai_api_text_detailed(
                    prompt=final_prompt,
                    api_config=writer_api_config,
                    system_prompt=system_prompt,
                    max_tokens=6000,
                    temperature=0.7,
                    logger=self.logger,
                )

                content = result.get("content")
                if isinstance(content, str) and content.strip():
                    return {
                        'content': content,
                        'finish_reason': result.get('finish_reason', 'stop'),
                    }

                error_kind = str(result.get("error_kind") or "").strip()
                last_message = (
                    str(result.get("message") or "").strip()
                    or error_kind
                    or "empty response from Writer_API"
                )
                if error_kind in {"quota_exhausted", "fatal_config_or_auth"}:
                    break
                if attempt < max_section_api_attempts:
                    wait_seconds = 2 * (2 ** (attempt - 1))
                    self.logger.warning(
                        f"章节内容生成返回空内容或失败: {last_message}；"
                        f"{wait_seconds:.1f}秒后重试 ({attempt + 1}/{max_section_api_attempts})"
                    )
                    time.sleep(wait_seconds)

            self.logger.error(f"章节内容生成失败: {last_message}")
            return None

        except Exception as e:
            self.logger.error(f"调用优化章节API失败: {e}")
            return None

    def append_section_to_word_document(self, section_number: int, section_title: str, section_text: str, word_file: str) -> bool:
        """将章节内容追加到Word文档（带样式配置）"""
        return append_section_to_word_document(self, section_number, section_title, section_text, word_file)

    def generate_full_review_from_outline(self) -> bool:
        """从大纲生成完整文献综述"""
        self.logger.info("=" * 60 + "\n文献综述自动生成器 - 阶段二：综述生成\n" + "=" * 60)
        try:
            self._check_cancelled()
            if not self.load_configuration(): 
                return False
            if not self.setup_output_directory(): 
                return False
            if not self.load_existing_summaries():
                self.logger.error("无法加载摘要文件，请先运行阶段一")
                return False
            if not self.summaries:
                self.logger.error("没有找到任何摘要，请先运行阶段一")
                return False
            
            writer_config: Dict[str, Any] = (self.config or {}).get('Writer_API', {})  # type: ignore
            if 'dummy' in (writer_config.get('api_key') or ''):  # type: ignore
                if not self.output_dir:
                    self.logger.error("输出目录未设置")
                    return False

                word_file = self._get_review_word_file_path()
                doc = Document()  # type: ignore
                doc.add_heading('Dummy Literature Review', 0)
                doc.add_paragraph('This is a dummy literature review.')
                doc.save(word_file)
                if not self._persist_review_draft(
                    outline_file="",
                    review_sections=[
                        {
                            "section_number": 1,
                            "section_title": "Dummy Literature Review",
                            "content": "This is a dummy literature review.",
                        }
                    ],
                    references=[],
                    word_file=word_file,
                    generation_mode="dummy_full_review",
                ):
                    return False
                self.logger.success(f"Dummy review saved to {word_file}")
                return True

            writer_api_config: APIConfig = get_writer_api_config(self.config)
            
            # 加载大纲文件
            if not self.output_dir:
                self.logger.error("输出目录未设置")
                return False

            outline_artifact = self._load_outline_artifact()
            review_checkpoint_file = self._get_review_checkpoint_file_path()
            word_file = self._get_review_word_file_path()
            outline_file = self._get_legacy_outline_file_path()
            
            if outline_artifact is None:
                self.logger.error(f"大纲文件不存在: {outline_file}，请先运行 --generate-outline 生成大纲")
                return False
            
            outline_file, outline_content = outline_artifact
            
            # 解析大纲中的所有章节
            import re
            section_matches = re.findall(r"^##\s*(\d+)\.\s*(.*)", outline_content, re.MULTILINE)
            
            if not section_matches:
                self.logger.error("大纲中没有找到任何章节（格式：## 数字. 标题）")
                return False
            
            self.logger.info(f"从大纲中解析到 {len(section_matches)} 个章节")
            
            # 验证章节编号连续性
            section_numbers = [int(match[0]) for match in section_matches]
            section_numbers.sort()
            for i in range(1, len(section_numbers)):
                if section_numbers[i] != section_numbers[i-1] + 1:
                    self.logger.error(f"大纲章节编号不连续：发现第{section_numbers[i-1]}章后直接是第{section_numbers[i]}章")
                    self.logger.error("请检查大纲文件，确保章节编号连续（如1, 2, 3...）")
                    return False
            self.logger.success("大纲章节编号验证通过：编号连续")
            section_titles_by_number = {
                int(section_num): section_title
                for section_num, section_title in section_matches
            }
            
            # 检查断点续传文件
            last_completed_section = 0
            checkpoint_loaded = False
            if os.path.exists(review_checkpoint_file):
                try:
                    with open(review_checkpoint_file, 'r', encoding='utf-8') as f:
                        checkpoint = json.load(f)
                        last_completed_section = checkpoint.get('last_completed_section', 0)
                        checkpoint_loaded = True
                except Exception as e:
                    self.logger.warning(f"读取断点文件失败，将从头开始: {e}")
                    last_completed_section = 0
            else:
                self.logger.info("[全新开始] 未发现断点文件，将从第1章开始生成")

            if last_completed_section > 0 and not os.path.exists(word_file):
                self.logger.warning(
                    f"[断点续传] 发现综述断点 last_completed_section={last_completed_section}，"
                    "但未找到可恢复的综述文档；本次实际起点为第1章，将重新生成，"
                    "避免跳过已丢失的章节正文。"
                )
                last_completed_section = 0
            elif last_completed_section > 0:
                self.logger.info(
                    f"[断点续传] 已确认可恢复综述文档，将从第 {last_completed_section + 1} 章继续。"
                )
            elif checkpoint_loaded:
                self.logger.info("[全新开始] 未发现有效断点，将从第1章开始生成")
            
            # 洁净启动机制：全新任务时删除旧文件
            if last_completed_section == 0 and os.path.exists(word_file):
                self.logger.info(f"检测到已存在的旧综述文件，将创建全新版本: {word_file}")
                try:
                    os.remove(word_file)
                except Exception as e:
                    self.logger.error(f"无法删除旧的综述文件，请检查文件权限: {e}")
                    return False
            
            # 创建或加载Word文档
            doc = None
            if os.path.exists(word_file) and last_completed_section > 0:
                # 断点续传：加载现有文档
                try:
                    doc = Document(word_file)  # type: ignore
                    self.logger.info(f"[断点续传] 已加载现有文档: {word_file}")
                except Exception as e:
                    self.logger.error(f"加载现有文档失败，将创建新文档: {e}")
                    doc = Document()  # type: ignore
            else:
                # 全新开始：创建新文档
                doc = Document()  # type: ignore
                
                # 加载样式配置
                style_config = self.config.get('Styling') if self.config else {}  # type: ignore
                font_name = style_config.get('font_name', 'Times New Roman')  # type: ignore
                font_size_body = int(style_config.get('font_size_body', '12'))  # type: ignore
                font_size_heading1 = int(style_config.get('font_size_heading1', '16'))  # type: ignore
                font_size_heading2 = int(style_config.get('font_size_heading2', '14'))  # type: ignore
                
                # 设置默认字体
                doc.styles['Normal'].font.name = font_name  # type: ignore
                doc.styles['Normal'].font.size = Pt(font_size_body)  # type: ignore
                
                # 设置中文字体
                doc.styles['Normal']._element  # type: ignore.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # type: ignore
                
                # 设置标题样式
                doc.styles['Heading 1'].font.name = font_name  # type: ignore
                doc.styles['Heading 1'].font.size = Pt(font_size_heading1)  # type: ignore
                doc.styles['Heading 1']._element  # type: ignore.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # type: ignore
                
                doc.styles['Heading 2'].font.name = font_name  # type: ignore
                doc.styles['Heading 2'].font.size = Pt(font_size_heading2)  # type: ignore
                doc.styles['Heading 2']._element  # type: ignore.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # type: ignore
                
                title = doc.add_heading('文献综述', level=0)
                if title is not None:  # type: ignore
                    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # type: ignore

                # 应用标题样式
                for run in title.runs:

                    run.font.name = font_name  # type: ignore

                    run.font.size = Pt(font_size_heading1 + 2)  # 主标题稍大  # type: ignore
                
                # 添加生成时间
                date_para = doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y年%m月%d日')}")
                date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # type: ignore
                
                # 应用日期样式

                for run in date_para.runs:

                    run.font.name = font_name  # type: ignore

                    run.font.size = Pt(font_size_body)  # type: ignore

            review_sections_by_number: Dict[int, Dict[str, Any]] = {}
            if last_completed_section > 0 and os.path.exists(word_file):
                review_sections_by_number.update(
                    self._extract_review_sections_from_word_document(
                        word_file,
                        section_titles_by_number=section_titles_by_number,
                    )
                )
            
            # 用tqdm包装章节列表，显示进度条
            failed_sections: List[Dict[str, Any]] = []
            total_sections = len(section_matches)
            progress_bar = tqdm(enumerate(section_matches, 1), total=total_sections, desc="[阶段二] 正在生成综述章节")
            
            # 发送阶段二开始的进度
            self._emit_progress(
                stage="outline",
                total=total_sections,
                current=0,
                message=f"开始生成综述章节，共 {total_sections} 章",
                indeterminate=False
            )
            
            # 逐章生成内容（从断点开始）
            for i, (section_num, section_title) in progress_bar:
                self._check_cancelled()
                # 跳过已完成的章节
                if i <= last_completed_section:
                    self.logger.info(f"[跳过] 第{section_num}章已完成，继续下一章...")
                    continue
                
                # 新增：跳过参考文献和附录章节
                if "参考文献" in section_title or "附录" in section_title:
                    self.logger.info(f"[跳过] 第{section_num}章 '{section_title}' 将在最后由程序自动生成。")
                    continue
                
                # 更新进度条的当前章节信息
                progress_bar.set_postfix_str(f"当前章节: {section_num}. {section_title[:30]}...")
                
                # 发送章节开始的进度
                self._emit_progress(
                    stage="outline",
                    total=total_sections,
                    current=i,
                    message=f"正在生成第{section_num}章: {section_title}",
                    item_label=section_title,
                    indeterminate=False
                )
                
                self.logger.info(f"正在生成第{section_num}章: {section_title}")
                
                # 生成章节内容
                section_content = self.generate_review_section_content(section_title, outline_content)
                if not section_content:
                    self.logger.error(f"第{section_num}章内容生成失败")
                    # 发送章节失败的进度
                    self._emit_progress(
                        stage="outline",
                        total=total_sections,
                        current=i,
                        message=f"第{section_num}章内容生成失败: {section_title}",
                        item_label=section_title,
                        indeterminate=False
                    )
                    failed_sections.append(
                        {
                            "section_number": int(section_num),
                            "section_title": section_title,
                            "failure_reason": "section_content_generation_failed",
                            "update_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        }
                    )
                    break
                
                # 发送章节成功的进度
                self._emit_progress(
                    stage="outline",
                    total=total_sections,
                    current=i,
                    message=f"第{section_num}章内容生成成功: {section_title}",
                    item_label=section_title,
                    indeterminate=False
                )
                
                # 检查章节内容是否包含结构化citation token
                if "[[cite:" not in section_content:
                    self.logger.warning(
                        f"Section {section_num} is missing structured citation tokens; attempting one narrow retry."
                    )
                    # One narrow citation-token insertion retry.
                    # Boundary markers prevent the embedded section content from
                    # being interpreted as prompt instructions.
                    retry_section_escaped = section_content.replace("'''", r"\'\'\'")
                    retry_prompt = (
                        "The following section content is missing [[cite:paper_key]] citation tokens. "
                        "Please rewrite it inserting appropriate [[cite:...]] tokens for every factual claim "
                        "referenced from the provided paper summaries. Preserve all other content and structure.\n\n"
                        f"=== SECTION TO FIX ===\nSection: {section_title}\n\n"
                        "'''CONTENT START'''\n"
                        f"{retry_section_escaped[:4000]}\n"
                        "'''CONTENT END'''\n\n"
                        "Rewrite the above CONTENT section with citation tokens added."
                    )
                    retry_result = self._call_section_api_optimized(
                        retry_prompt,
                        writer_api_config,
                        is_continuation=False,
                    )
                    if retry_result and "[[cite:" in str(retry_result.get("content", "")):
                        section_content = retry_result["content"]
                        self.logger.success(f"Citation-token retry succeeded for section {section_num}")
                    else:
                        self.logger.error(
                            f"Section {section_num} still lacks structured citation tokens after retry; "
                            "canonical review generation is blocked."
                        )
                        failed_sections.append(
                            {
                                "section_number": int(section_num),
                                "section_title": section_title,
                                "failure_reason": "missing_structured_citation_tokens",
                                "update_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            }
                        )
                        break

                review_sections_by_number[int(section_num)] = {
                    "section_number": int(section_num),
                    "section_title": section_title,
                    "content": section_content,
                }
                
                checkpoint_data: Dict[str, Any] = {
                    'last_completed_section': i,
                    'last_section_title': section_title,
                    'update_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                }
                atomic_write_json(review_checkpoint_file, checkpoint_data)
                
                self.logger.success(f"第{section_num}章已处理并更新断点")
            
            # 在所有章节处理完成后，一次性保存文档
            self._save_failed_review_sections(failed_sections)
            
            # Emit final stage-two progress after all sections are processed.
            if not failed_sections:
                self._emit_progress(
                    stage="outline",
                    total=total_sections,
                    current=total_sections,
                    message="All review sections generated successfully.",
                    indeterminate=False
                )
            else:
                self._emit_progress(
                    stage="outline",
                    total=total_sections,
                    current=total_sections - len(failed_sections),
                    message=f"Review generation finished with {len(failed_sections)} failed sections.",
                    indeterminate=False
                )

            if not failed_sections:
                if os.path.exists(review_checkpoint_file):
                    keep_checkpoints = self._keep_checkpoints_after_completion()
                    if not keep_checkpoints:
                        os.remove(review_checkpoint_file)
                        self.logger.info("Removed review checkpoint after section generation completed.")
                    else:
                        self.logger.info("Keeping review checkpoint because configuration requires it.")
            else:
                self.logger.info(
                    f"Preserving review checkpoint because {len(failed_sections)} section(s) failed."
                )

            if failed_sections:
                self.logger.warning("Skipping review_draft_v1 registration because one or more review sections failed")
                return False
            else:
                review_sections = [
                    review_sections_by_number[int(section_num)]
                    for section_num, _section_title in section_matches
                    if int(section_num) in review_sections_by_number
                ]
                references: List[str] = []

                # Persist review_draft_v2 once with empty references so the canonical manifest can be built from block citations.
                if not self._persist_review_draft_v2(
                    outline_file=outline_file,
                    review_sections=review_sections,
                    references=references,
                    word_file=word_file,
                ):
                    return False

                review_draft_path = self._review_draft_v2_path()
                if not self._persist_citation_manifest(
                    review_draft_path=review_draft_path,
                    review_word_path=word_file,
                ):
                    return False

                citation_manifest = self._load_citation_manifest()
                if not citation_manifest:
                    self.logger.error("Canonical citation_manifest_v3 was not persisted.")
                    return False

                from docx_writer import generate_apa_references_from_manifest, rebuild_review_docx_from_structured_artifacts

                references = generate_apa_references_from_manifest(
                    citation_manifest,
                    self,
                    allow_compat_fallback=False,
                )

                # Persist both draft projections again with canonical references attached.
                if not self._persist_review_draft(
                    outline_file=outline_file,
                    review_sections=review_sections,
                    references=references,
                    word_file=word_file,
                ):
                    return False

                if not self._persist_review_draft_v2(
                    outline_file=outline_file,
                    review_sections=review_sections,
                    references=references,
                    word_file=word_file,
                ):
                    return False

                with open(review_draft_path, "r", encoding="utf-8") as handle:
                    review_draft = json.load(handle)
                rebuild_review_docx_from_structured_artifacts(
                    self,
                    review_draft,
                    citation_manifest,
                    word_file,
                    allow_compat_fallback=False,
                )

            self.logger.success(f"完整文献综述已生成: {word_file}")
            
            # 第二阶段验证（根据配置决定是否自动运行）
            try:
                if self._stage2_validation_enabled():
                    self.logger.info("根据配置文件自动启动第二阶段验证...")
                    from validator import run_review_validation
                    validation_result: Dict[str, Any] = run_review_validation(self)  # type: ignore
                    if validation_result.get("success"):
                        self.logger.success("第二阶段验证完成！验证报告已生成。")
                    else:
                        self.logger.warning("第二阶段验证失败，请检查验证报告文件。")
                else:
                    self.logger.info("第二阶段验证未在配置中启用。如需运行验证，请使用: --validate-review")
            except Exception as e:
                self.logger.error(f"第二阶段验证运行时出错: {e}")
                self.logger.info("您可以手动运行验证命令: python main.py --validate-review")
            
            return True
            
        except Exception as e:
            self.logger.error(f"从大纲生成文献综述失败: {e}")
            return False

    def generate_word_table_of_contents(self, doc: Any) -> bool:  # type: ignore
        """为Word文档生成自动目录"""
        return generate_word_table_of_contents(doc)

    def _load_citation_manifest(self, *, allow_legacy: bool = False) -> Optional[Dict[str, Any]]:
        """Load the canonical citation manifest, optionally falling back to legacy artifacts for migration or compatibility only."""
        try:
            v3_path = self._citation_manifest_path()
            if os.path.exists(v3_path):
                with open(v3_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            if not allow_legacy:
                return None
            v2_path = self._citation_manifest_v2_path()
            if os.path.exists(v2_path):
                with open(v2_path, 'r', encoding='utf-8') as f:
                    manifest = json.load(f)
                manifest.setdefault("migration_report", {})
                manifest["migration_report"].setdefault("load_source", "v2")
                return manifest
            v1_path = self._citation_manifest_v1_path()
            if os.path.exists(v1_path):
                with open(v1_path, 'r', encoding='utf-8') as f:
                    manifest = json.load(f)
                manifest.setdefault("migration_report", {})
                manifest["migration_report"].setdefault("load_source", "v1")
                return manifest
        except Exception as e:
            self.logger.warning(f"Failed to load citation manifest: {e}")
        return None

    def generate_apa_references(self) -> List[str]:
        """Generate APA references from the canonical citation manifest."""
        citation_manifest = self._load_citation_manifest()
        return generate_apa_references_from_manifest(citation_manifest, self, allow_compat_fallback=False)


    def generate_literature_review_outline(self) -> bool:
        """生成文献综述大纲（带智能续写循环）"""
        self.logger.info("=" * 60 + "\n文献综述自动生成器 - 阶段二：大纲生成\n" + "=" * 60)
        try:
            if not self.load_configuration(): 
                return False
            if not self.setup_output_directory(): 
                return False
            if not self.load_existing_summaries():
                self.logger.error("无法加载摘要文件，请先运行阶段一")
                return False
            if not self.summaries:
                self.logger.error("没有找到任何摘要，请先运行阶段一")
                return False
            
            writer_config: Dict[str, Any] = (self.config or {}).get('Writer_API', {})  # type: ignore
            if 'dummy' in (writer_config.get('api_key') or ''):  # type: ignore
                outline_content = "# Dummy Outline\n\n## Introduction\n\n## Body Paragraph\n\n## Conclusion"
                outline_file = self._write_outline_artifact(
                    outline_content,
                    producer="main.LiteratureReviewGenerator.generate_literature_review_outline",
                )
                self.logger.success(f"Dummy outline saved to {outline_file}")
                return True
            
            self.logger.info(f"已加载{len(self.summaries)}个文献摘要")
            return self.create_literature_review_outline()
        except Exception as e:
            self.logger.error(f"阶段二运行失败: {e}")
            return False

    def _create_literature_review_outline_v2(self) -> bool:
        """Run Outline Intelligence v2 Stage 2 artifact chain."""
        try:
            compat = self._ensure_compat_config()
            errors = compat.validate_outline_v2_config()
            if errors:
                for error in errors:
                    self.logger.error(f"Outline v2 配置错误: {error}")
                return False

            from outline.pipeline import V2Pipeline

            test_dev_mode = compat.outline_test_dev_fixture_mode()
            pipeline = V2Pipeline(
                job_id=self.job_workspace.job_id if self.job_workspace else "standalone",
                summaries=self.summaries,
                config_view=compat,
                artifact_registry=self.artifact_registry,
                workspace=self.job_workspace,
                output_dir=self.output_dir or "",
                project_name=self.project_name or "review",
                model_caller=self._outline_v2_model_call,
                logger=self.logger,
            )
            result = pipeline.run(
                candidate_count=compat.outline_candidate_count(),
                test_dev_mode=test_dev_mode,
                generator_model=compat.outline_model(),
                structure_critic=compat.structure_critic_model(),
                coverage_critic=compat.coverage_critic_model(),
                arbitrator_model=compat.arbitrator_model(),
                paper_artifacts=self._load_paper_artifacts_for_outline_v2(),
            )
            if not result.ok:
                for error in result.errors:
                    self.logger.error(f"Outline v2 生成失败: {error}")
                return False

            paths = pipeline.persist_artifacts(result)
            if result.coverage_audit and not result.coverage_audit.passed:
                self.logger.warning(
                    "Outline v2 覆盖审计未通过；已写入工件，但采纳会被阻止。"
                )
            self.logger.success(
                "Outline Intelligence v2 工件链已生成: "
                + ", ".join(f"{name}={path}" for name, path in paths.items())
            )
            if compat.outline_require_explicit_adopt():
                self.logger.info(
                    "v2 需要显式采纳。请检查 final_outline/audit 后运行 --adopt-outline-v2；"
                    "generate-review 将只接受 adopted_final_outline.json。"
                )
            return True
        except Exception as exc:
            self.logger.error(f"创建 Outline Intelligence v2 大纲失败: {exc}")
            import traceback
            self.logger.debug(f"详细错误信息: {traceback.format_exc()}")
            return False

    def create_literature_review_outline(self) -> bool:
        """创建文献综述大纲，适配新的纯文本输出格式"""
        try:
            if self._outline_v2_enabled():
                return self._create_literature_review_outline_v2()

            review_data = self.prepare_review_data()
            outline_content = self.generate_review_outline(review_data)
            if not outline_content:
                self.logger.error("文献综述大纲生成失败")
                return False
            
            # outline_content应该是纯文本字符串
            if not isinstance(outline_content, str):  # type: ignore
                self.logger.warning("预期收到纯文本，但收到其他格式，正在转换...")
                outline_text = str(outline_content)
            else:
                outline_text = outline_content
            
            # 生成大纲文件路径（添加项目名称前缀）
            outline_file = self._write_outline_artifact(
                outline_text,
                producer="main.LiteratureReviewGenerator.create_literature_review_outline",
            )
            
            # 保存大纲文件
            
            self.logger.success(f"文献综述大纲已生成: {outline_file}")
            # 根据模式提供不同的命令提示
            if self.mode == "direct":
                self.logger.info("大纲已生成。请检查并修改。然后，运行以下命令生成完整综述：")
                self.logger.info(f"命令: python main.py --pdf-folder \"{self.pdf_folder}\" --generate-review")
            else:
                self.logger.info("大纲已生成。请检查并修改。然后，运行以下命令生成完整综述：")
                self.logger.info(f"命令: python main.py --project-name \"{self.project_name}\" --generate-review")
            return True
                
        except Exception as e:
            self.logger.error(f"创建文献综述大纲失败: {e}")
            import traceback
            traceback.print_exc()
            return False

    def generate_review_outline(self, review_data: Dict[str, Any]) -> Optional[str]:
        """生成综述大纲内容，适配新的两段式JSON输入（智能续写循环版本）"""
        try:
            # 从提示词模板文件读取大纲提示词
            outline_prompt_template = ""
            try:
                with open('prompts/prompt_synthesize_outline.txt', 'r', encoding='utf-8') as f:
                    outline_prompt_template = f.read()
                self.logger.success(f"加载大纲提示词模板: {len(outline_prompt_template)}字符")
            except Exception as e:
                self.logger.warning(f"无法加载大纲提示词模板，使用默认提示词: {e}")
                try:
                    with open('prompts/prompt_default_outline.txt', 'r', encoding='utf-8') as f:
                        outline_prompt_template = f.read()
                    self.logger.success(f"加载默认大纲提示词模板: {len(outline_prompt_template)}字符")
                except Exception as e2:
                    self.logger.error(f"无法加载默认大纲提示词模板: {e2}")
                    outline_prompt_template = "基于以下文献摘要信息，请生成一份详细的文献综述大纲。\n\n{{SUMMARIES_JSON_ARRAY}}"

            # 从提示词模板文件读取续写提示词
            continue_prompt_template = ""
            try:
                with open('prompts/prompt_continue_outline.txt', 'r', encoding='utf-8') as f:
                    continue_prompt_template = f.read()
                self.logger.success(f"加载续写提示词模板: {len(continue_prompt_template)}字符")
            except Exception as e:
                self.logger.warning(f"无法加载续写提示词模板，使用默认提示词: {e}")
                try:
                    with open('prompts/prompt_default_continue_outline.txt', 'r', encoding='utf-8') as f:
                        continue_prompt_template = f.read()
                    self.logger.success(f"加载默认续写提示词模板: {len(continue_prompt_template)}字符")
                except Exception as e2:
                    self.logger.error(f"无法加载默认续写提示词模板: {e2}")
                    continue_prompt_template = "请继续完成这份未写完的文献综述大纲。\n\n【全部论文分析数据】\n{{SUMMARIES_JSON_ARRAY}}\n\n【已完成的大纲草稿】\n{{PARTIAL_OUTLINE}}"

            # 将整个summaries列表转换为格式化的JSON字符串（包含两段式结构）
            summaries_string = json.dumps(self.summaries, ensure_ascii=False, indent=2)
            self.logger.success(f"生成摘要JSON字符串: {len(summaries_string)}字符")

            # 始终使用优化后的高密度格式（去除JSON结构开销），仅在最极端情况下触发截断
            # Gemini 3 Pro有1M token上下文，使用950000作为安全阈值
            estimated_tokens = estimate_tokens(summaries_string)
            max_tokens_for_optimization = 950000  # 优化时最大token数（仅在最极端情况下触发截断）
            
            self.logger.info(f"上下文token数({estimated_tokens})，使用高密度压缩格式...")
            optimized_context = optimize_context_for_outline(self.summaries, max_tokens=max_tokens_for_optimization)
            self.logger.success(f"优化后的上下文长度: {len(optimized_context)}字符 (原长度: {len(summaries_string)}字符)")
            self.logger.info(f"压缩率: {len(optimized_context)/len(summaries_string):.1%}")
            
            # 使用优化后的上下文
            summaries_string = optimized_context

            # 阶段二优先使用 Outline_API，未配置时回退到 Writer_API
            outline_api_config: APIConfig = get_outline_api_config(self.config)

            self.logger.info("正在调用大纲引擎生成文献综述大纲（智能续写循环模式）...")

            # ===== 智能续写循环核心逻辑 =====
            partial_outline = ""  # 存储已生成的大纲内容
            continuation_attempts = 0  # 续写计数器
            max_continuation_attempts = 5  # 最大续写次数（安全熔断机制）
            
            while continuation_attempts <= max_continuation_attempts:
                try:
                    # 根据是否为首次调用选择不同的提示词
                    if continuation_attempts == 0:
                        # 首次调用：使用原始大纲提示词
                        final_prompt = outline_prompt_template.replace('{{SUMMARIES_JSON_ARRAY}}', summaries_string)
                        final_prompt = self._inject_free_mode_context(final_prompt)
                        self.logger.info(f"首次大纲生成，提示词长度: {len(final_prompt)}字符")
                    else:
                        # 续写调用：使用续写提示词
                        final_prompt = continue_prompt_template.replace('{{SUMMARIES_JSON_ARRAY}}', summaries_string)
                        final_prompt = final_prompt.replace('{{PARTIAL_OUTLINE}}', partial_outline)  # type: ignore
                        final_prompt = self._inject_free_mode_context(final_prompt)
                        self.logger.info(f"续写大纲生成(第{continuation_attempts}次)，提示词长度: {len(final_prompt)}字符")  # type: ignore

                    # 调用AI API
                    # 加载系统提示词
                    try:
                        with open('prompts/prompt_system_outline.txt', 'r', encoding='utf-8') as f:
                            system_prompt = f.read()
                        self.logger.success(f"加载大纲系统提示词模板: {len(system_prompt)}字符")
                    except Exception as e:
                        self.logger.warning(f"无法加载大纲系统提示词模板，使用默认提示词: {e}")
                        system_prompt = """你是一个学术文献综述专家。请基于提供的文献分析结果生成一份详细的文献综述大纲。

要求：
1. 直接输出Markdown格式的大纲内容
2. 使用Markdown的标题格式（# 主要标题, ## 章节标题, ### 小节标题）
3. 每个章节标题下，用项目符号（-）列出该章节应包含的核心论点或分析要点
4. 大纲应该结构清晰、逻辑严谨
5. 不要包含任何正文内容，只输出大纲"""

                    ai_response_text = _call_ai_api(
                        prompt=final_prompt,
                        api_config=outline_api_config,
                        system_prompt=system_prompt,
                        max_tokens=8192,
                        temperature=0.7,
                        response_format="text",
                        logger=self.logger  # 添加logger参数以记录详细错误信息
                    )
                    
                    if ai_response_text is None:
                        self.logger.error("API调用失败，无法生成大纲")
                        return None
                    
                    # 模拟旧API的返回结构以适配后续逻辑
                    ai_response = {'choices': [{'message': {'content': ai_response_text}, 'finish_reason': 'stop'}]}  # type: ignore
                    
                    # 提取AI回复内容和完成原因
                    outline_content = ai_response['choices'][0]['message']['content']  # type: ignore
                    finish_reason = ai_response['choices'][0]['finish_reason']  # type: ignore
                    
                    if outline_content and len(outline_content) > 100:  # type: ignore
                        # 将本次生成的内容追加到部分大纲中
                        if continuation_attempts == 0:
                            partial_outline = outline_content  # type: ignore
                        else:
                            partial_outline += "\n\n" + outline_content  # type: ignore
                        
                        self.logger.success(f"大纲片段生成成功，当前总长度: {len(partial_outline)}字符")  # type: ignore
                        self.logger.info(f"完成原因: {finish_reason}")
                        
                        # 检查是否需要继续续写
                        if finish_reason == 'stop':  # type: ignore
                            self.logger.success("大纲生成完成，无需续写")
                            return partial_outline  # type: ignore
                        elif finish_reason == 'length':
                            self.logger.info("大纲被截断，准备续写...")
                            continuation_attempts += 1
                            continue
                        else:
                            self.logger.warning(f"未知的完成原因: {finish_reason}，尝试续写...")
                            continuation_attempts += 1
                            continue
                    else:
                        self.logger.warning(f"大纲内容过短({len(outline_content) if outline_content else 0}字符)，重试...")  # type: ignore
                        continuation_attempts += 1
                        continue

                except Exception as e:
                    self.logger.error(f"大纲生成过程出错: {str(e)}")
                    continuation_attempts += 1
                    if continuation_attempts <= max_continuation_attempts:
                        self.logger.info(f"准备重试第{continuation_attempts}次...")
                        continue
                    else:
                        break
            
            # 安全熔断：达到最大续写次数
            if continuation_attempts > max_continuation_attempts:
                self.logger.error(f"[ERROR] 大纲生成续写次数过多({continuation_attempts}次)，或已陷入死循环。请检查输入数据或Prompt。")
                if partial_outline and len(partial_outline) > 100:  # type: ignore  # 只有部分内容足够长才返回
                    self.logger.warning("返回部分生成的大纲内容")
                    return partial_outline  # type: ignore
                self.logger.error("大纲生成失败，内容过短或为空")
                return None
            
            # 最终检查：只有内容足够长才认为成功
            if partial_outline and len(partial_outline) > 100:  # type: ignore
                return partial_outline  # type: ignore
            else:
                self.logger.error("大纲生成失败，内容过短或为空")
                return None

        except Exception as e:
            self.logger.error(f"生成大纲内容失败: {e}")
            return None

    

    def create_literature_review(self) -> bool:
        """创建文献综述，适配新的纯文本输出格式"""
        try:
            review_data = self.prepare_review_data()
            review_content = self.generate_review_content(review_data)
            if not review_content:
                self.logger.error("文献综述生成失败")
                return False
            
            # review_content现在应该是纯文本字符串
            if not isinstance(review_content, str):  # type: ignore
                self.logger.warning("预期收到纯文本，但收到其他格式，正在转换...")
                review_text = str(review_content)
            else:
                review_text = review_content
            
            # 生成Word文档路径（添加项目名称前缀）
            if not self.output_dir:
                self.logger.error("输出目录未设置")
                return False

            word_file = self._get_review_word_file_path()
            
            # 创建Word文档
            success = self.create_word_document(review_text, word_file)
            
            if success:
                self.logger.success(f"文献综述Word文档已生成: {word_file}")
                return True
            else:
                return False
                
        except Exception as e:
            self.logger.error(f"创建文献综述失败: {e}")
            return False

    def prepare_review_data(self) -> Dict[str, Any]:
        review_data: Dict[str, Any] = {  # type: ignore
            'total_papers': len(self.summaries),
            'successful_papers': len([s for s in self.summaries if s.get('status') == 'success']),
            'failed_papers': len([s for s in self.summaries if s.get('status') != 'success']), 
            'papers': [],
            'research_areas': {}, 
            'methodologies': {}, 
            'key_findings': [], 
            'common_themes': []
        }
        
        for summary in self.summaries:
            if summary.get('status') != 'success': 
                continue
                
            paper_info = summary.get('paper_info', {})
            ai_summary: Union[AISummary, Dict[str, Any], None] = summary.get('ai_summary', {})
            
            # 适配新的两段式结构
            if False:  # legacy branch retained for compatibility cleanup
                # 新的两段式结构
                core_analysis = get_core_analysis(ai_summary or {})  # type: ignore
            else:
                # 兼容旧的单段式结构
                core_analysis = get_core_analysis(ai_summary or {})  # type: ignore
            
            paper_data: Dict[str, Any] = {  # type: ignore
                'title': paper_info.get('title', '未知标题'),
                'authors': paper_info.get('authors', []),
                'year': paper_info.get('year', '未知年份'),
                'journal': paper_info.get('journal', '未知期刊'),
                'summary': core_analysis.get('summary', ''),  # type: ignore
                'key_points': core_analysis.get('key_points', []),  # type: ignore
                'methodology': core_analysis.get('methodology', ''),  # type: ignore
                'findings': core_analysis.get('findings', ''),  # type: ignore
                'conclusions': core_analysis.get('conclusions', ''),  # type: ignore
                'relevance': core_analysis.get('relevance', ''),  # type: ignore
                'limitations': core_analysis.get('limitations', '')  # type: ignore
            }
            
            review_data['papers'].append(paper_data)  # type: ignore
            
            methodology = paper_data['methodology']
            if methodology: 
                review_data['methodologies'][methodology] = review_data['methodologies'].get(methodology, 0) + 1  # type: ignore
                
            findings = paper_data['findings']  # type: ignore
            if findings: 
                review_data['key_findings'].append(findings)  # type: ignore
                
        return review_data

    def generate_review_content(self, review_data: Dict[str, Any]) -> Optional[str]:
        """生成综述内容，适配新的两段式JSON输入"""
        try:
            # 从提示词模板文件读取综述提示词
            synthesize_prompt_template = ""
            try:
                with open('prompts/prompt_synthesize.txt', 'r', encoding='utf-8') as f:
                    synthesize_prompt_template = f.read()
                self.logger.success(f"加载综述提示词模板: {len(synthesize_prompt_template)}字符")
            except Exception as e:
                self.logger.warning(f"无法加载综述提示词模板，使用默认提示词: {e}")
                try:
                    with open('prompts/prompt_default_synthesize.txt', 'r', encoding='utf-8') as f:
                        synthesize_prompt_template = f.read()
                    self.logger.success(f"加载默认综述提示词模板: {len(synthesize_prompt_template)}字符")
                except Exception as e2:
                    self.logger.error(f"无法加载默认综述提示词模板: {e2}")
                    synthesize_prompt_template = "基于以下文献摘要信息，请生成一份完整的文献综述报告。\n\n{{SUMMARIES_JSON_ARRAY}}"

            # 将整个summaries列表转换为格式化的JSON字符串（包含两段式结构）
            summaries_string = json.dumps(self.summaries, ensure_ascii=False, indent=2)
            self.logger.success(f"生成摘要JSON字符串: {len(summaries_string)}字符")

            # 将完整的JSON字符串注入到模板中
            final_prompt = synthesize_prompt_template.replace('{{SUMMARIES_JSON_ARRAY}}', summaries_string)
            final_prompt = self._inject_free_mode_context(final_prompt)
            self.logger.success(f"生成最终综述提示词: {len(final_prompt)}字符")

            writer_api_config: APIConfig = get_writer_api_config(self.config)

            self.logger.info("Calling Writer_API to generate literature review...")

            try:
                with open('prompts/prompt_system_synthesize.txt', 'r', encoding='utf-8') as f:
                    system_prompt = f.read()
                self.logger.success(f"Loaded review system prompt template: {len(system_prompt)} characters")
            except Exception as e:
                self.logger.warning(f"Unable to load review system prompt template, using default prompt: {e}")
                system_prompt = """You are an academic literature review expert. Generate a complete Chinese academic literature review from the provided paper analysis results.

Requirements:
1. Output plain review text, not JSON.
2. Use Markdown headings.
3. Keep the content professional, objective, and comprehensive.
4. Cite specific papers where appropriate.
5. Target 3000-5000 Chinese characters."""

            api_params = (self.config or {}).get("API_Parameters", {}) if self.config else {}
            try:
                max_tokens = int(api_params.get("writer_max_tokens", 8000))
            except (TypeError, ValueError):
                max_tokens = 8000
            try:
                temperature = float(api_params.get("writer_temperature", 0.7))
            except (TypeError, ValueError):
                temperature = 0.7

            result = _call_ai_api_text_detailed(
                prompt=final_prompt,
                api_config=writer_api_config,
                system_prompt=system_prompt,
                max_tokens=max_tokens,
                temperature=temperature,
                logger=self.logger,
            )
            review_content = result.get("content")
            if review_content and len(str(review_content)) > 100:
                self.logger.success("Writer_API returned review text")
                return str(review_content)
            self.logger.error(f"Review generation failed: {result.get('message', 'empty response')}")
            return None

        except Exception as e:
            self.logger.error(f"Failed to generate review content: {e}")
            return None

    @staticmethod
    def build_review_prompt(review_data: Dict[str, Any]) -> str:
        papers_info = []
        for i, paper in enumerate(review_data['papers'], 1):
            paper_text = f"文献 {i}: {paper['title']}\n作者: {', '.join(paper['authors']) if paper['authors'] else '未知'}\n年份: {paper['year']}\n期刊: {paper['journal']}\n\n摘要: {paper['summary']}\n\n研究方法: {paper['methodology']}\n主要发现: {paper['findings']}\n结论: {paper['conclusions']}\n相关性: {paper['relevance']}\n局限性: {paper['limitations']}\n\n关键要点:\n{chr(10).join(['- ' + point for point in paper['key_points']])}"
            papers_info.append(paper_text)  # type: ignore
        all_papers_text = '\n'.join(papers_info)  # type: ignore
        prompt = f"基于以下{review_data['total_papers']}篇学术文献的摘要信息，请生成一份完整的文献综述报告。\n\n文献信息:\n{all_papers_text}\n\n请按照以下结构生成文献综述：\n\n# 文献综述报告\n\n## 1. 引言\n- 研究领域概述\n- 研究背景和意义\n- 文献综述的目的和范围\n\n## 2. 研究现状分析\n- 主要研究主题和趋势\n- 研究方法的分析和比较\n- 关键发现的总结\n\n## 3. 研究热点和前沿\n- 当前研究的热点问题\n- 新兴的研究方向\n- 尚未解决的问题\n\n## 4. 研究方法和质量分析\n- 常用研究方法的评价\n- 研究质量的总体评估\n- 研究的局限性分析\n\n## 5. 综合讨论\n- 主要共识和分歧\n- 研究的理论贡献\n- 实践意义和应用前景\n\n## 6. 未来研究方向\n- 基于现有研究空白的建议\n- 方法学改进的建议\n- 理论和实践的发展方向\n\n## 7. 结论\n- 主要发现总结\n- 对领域的贡献\n- 综述的局限性\n\n## 参考文献\n- 按照学术规范列出所有文献\n\n要求：\n1. 内容要全面、客观、准确\n2. 要有批判性思维和分析\n3. 要指出研究趋势和未来方向\n4. 语言要专业、简洁、清晰\n5. 总字数在3000-5000字之间"
        return prompt

    @staticmethod
    def format_review_content(review_content: Dict[str, Any], review_data: Dict[str, Any]) -> str:
        header = f"# 文献综述报告\n\n**生成时间**: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}\n**文献数量**: {review_data['total_papers']}篇\n**成功处理**: {review_data['successful_papers']}篇\n**失败处理**: {review_data['failed_papers']}篇\n\n---\n\n"
        review_text = review_content if isinstance(review_content, str) else review_content.get('summary', json.dumps(  # type: ignore
            review_content, ensure_ascii=False, indent=2))
        references = "\n\n## 参考文献\n\n"
        for i, paper in enumerate(review_data['papers'], 1):
            authors = ', '.join(paper['authors']) if paper['authors'] else '未知作者';
            year = f" ({paper['year']})" if paper['year'] != '未知年份' else '';
            journal = f". {paper['journal']}" if paper['journal'] != '未知期刊' else ''
            references += f"{i}. {authors}{year}. {paper['title']}{journal}.\n"
        return header + review_text + references

    def create_word_document(self, markdown_text: str, output_path: str) -> bool:
        """将Markdown文本解析并创建Word文档（带样式配置）"""
        return create_word_document(self, markdown_text, output_path)

    def run_priming_phase(self, concept_name: str, seed_folder: str) -> bool:
        """概念学习阶段：分析核心论文以建立概念理解"""
        self.logger.info("=" * 60 + "\n概念学习阶段：建立概念理解\n" + "=" * 60)
        try:
            if not self.load_configuration():
                return False
            if not self.setup_output_directory():
                return False
            
            # 验证种子文件夹
            if not os.path.exists(seed_folder):
                self.logger.error(f"种子文件夹不存在: {seed_folder}")
                return False
            
            # 扫描种子论文
            seed_papers = []
            for root, _, files in os.walk(seed_folder):
                for file in files:
                    if file.lower().endswith('.pdf'):
                        seed_papers.append(os.path.join(root, file))  # type: ignore
            
            if not seed_papers:
                self.logger.error(f"种子文件夹中未找到PDF文件: {seed_folder}")
                return False
            
            self.logger.info(f"找到 {len(seed_papers)} 篇种子论文")  # type: ignore

            # 处理种子论文 - 保持完整信息量，使用并发处理
            concept_papers = []

            # 使用并发处理提高速度，但保持完整信息量
            max_workers = min(2, len(seed_papers))  # type: ignore  # 最多2个并发，避免API限制
            
            def process_seed_paper(pdf_path: str) -> Optional[Dict[str, Any]]:  # type: ignore
                """处理单个种子论文"""
                try:
                    self.logger.info(f"正在分析种子论文: {os.path.basename(pdf_path)}")  # type: ignore
                    
                    # 准备阶段一输入
                    pdf_text, _preprocess_metadata = self._prepare_stage1_input(pdf_path)
                    if not pdf_text or len(pdf_text.strip()) < 500:  # type: ignore
                        self.logger.warning(f"种子论文阶段一输入准备失败: {os.path.basename(pdf_path)}")  # type: ignore
                        return None
                    
                    # 创建论文信息
                    paper_info: PaperInfo = {
                        'title': os.path.splitext(os.path.basename(pdf_path))[0],
                        'authors': [],
                        'year': '未知年份',
                        'journal': '未知期刊',
                        'doi': '',
                        'pdf_path': pdf_path
                    }
                    
                    # 获取API配置
                    reader_api_config: APIConfig = get_reader_api_config(self.config)
                    backup_api_config: APIConfig = get_backup_reader_api_config(self.config)
                    
                    # 构建完整的分析提示词
                    try:
                        prompt_template = self._load_stage1_prompt_template()
                        
                        # 替换占位符
                        prompt_template = self._inject_free_mode_context(prompt_template)
                        analysis_prompt = prompt_template.replace('{{PAPER_FULL_TEXT}}', pdf_text)  # type: ignore
                        
                    except Exception as e:
                        self.logger.warning(f"无法加载分析提示词模板，使用简化提示词: {e}")
                        # 简化提示词
                        analysis_prompt = f"请分析以下论文内容，生成结构化摘要：\n\n{pdf_text}"
                    
                    # 调用AI分析
                    ai_result = get_summary_from_ai_with_fallback(analysis_prompt, reader_api_config, backup_api_config, logger=self.logger, config=self.config)
                    if ai_result:
                        self._apply_ai_metadata_backfill(paper_info, ai_result)
                        self.logger.success(f"种子论文分析成功: {os.path.basename(pdf_path)}")
                        return {
                            'paper_info': paper_info,
                            'ai_summary': ai_result
                        }
                    else:
                        self.logger.warning(f"种子论文分析失败: {os.path.basename(pdf_path)}")
                        return None
                        
                except Exception as e:
                    self.logger.error(f"处理种子论文时出错 {os.path.basename(pdf_path)}: {e}")
                    return None
            
            for future in concurrent.futures.as_completed(future_to_pdf):  # type: ignore
                result: Optional[Dict[str, Any]] = future.result()  # type: ignore
                if result:
                    concept_papers.append(result)  # type: ignore
            
            if not concept_papers:
                self.logger.error("没有成功分析任何种子论文")
                return False
            
            # 生成概念配置
            self.logger.info(f"正在生成概念配置: {concept_name}")
            concept_profile: Dict[str, Any] = self._generate_concept_profile(concept_name, concept_papers)  # type: ignore
            
            if not concept_profile:
                self.logger.error("概念配置生成失败")
                return False
            
            # 保存概念配置
            concept_profile_file: str = self._get_concept_profile_file_path()
            with open(concept_profile_file, 'w', encoding='utf-8') as f:  # type: ignore
                json.dump(concept_profile, f, ensure_ascii=False, indent=2)
            
            self.logger.success(f"概念配置已保存: {concept_profile_file}")
            return True
            
        except Exception as e:
            self.logger.error(f"概念学习阶段失败: {e}")
            return False
    
    
    def _fix_json_string(self, json_str: str) -> str:
        """修复常见的JSON字符串问题"""
        try:
            # 移除可能的注释
            import re
            json_str = re.sub(r'//.*', '', json_str)  # 移除单行注释
            json_str = re.sub(r'/\*.*?\*/', '', json_str, flags=re.DOTALL)  # 移除多行注释
            
            # 修复常见的JSON格式问题
            json_str = json_str.strip()
            
            # 如果字符串以引号开始但不以引号结束，添加结束引号
            if json_str.startswith('"') and not json_str.endswith('"'):
                json_str += '"'
            elif json_str.startswith("'") and not json_str.endswith("'"):
                json_str += "'"
            
            return json_str
        except Exception as e:
            self.logger.error(f"修复JSON字符串失败: {e}")
            return json_str
    
    def _generate_concept_profile(self, concept_name: str, concept_papers: list[Dict[str, Any]]) -> Dict[str, Any]:  # type: ignore
        """根据已分析的种子论文摘要，生成概念配置文件。"""
        try:
            self.logger.info(f"开始生成概念学习笔记: {concept_name}")
            self.logger.info(f"种子论文数量: {len(concept_papers)}")
            
            # 1. 加载概念分析的 Prompt 模板
            try:
                with open('prompts/prompt_prime_concept.txt', 'r', encoding='utf-8') as f:
                    prompt_template = f.read()
                self.logger.success(f"加载概念分析提示词模板: {len(prompt_template)}字符")
            except Exception as e:
                self.logger.error(f"无法加载概念分析提示词模板: {e}")
                return {}  # type: ignore

            # 2. 准备论文数据 (直接从传入的 concept_papers 构建)
            papers_data: list[Dict[str, Any]] = []  # type: ignore
            for paper in concept_papers:  # type: ignore
                # 从 paper['ai_summary'] 提取所需字段，构建 papers_data
                papers_data.append({
                    'file_name': paper.get('file_name', '未知文件'),  # type: ignore
                    'ai_summary': paper.get('ai_summary', {})  # type: ignore
                })
            
            # 3. 构建最终的 Prompt
            papers_json = json.dumps(papers_data, ensure_ascii=False, indent=2)
            final_prompt = prompt_template.replace('{{CONCEPT_NAME}}', concept_name).replace('{{SEED_PAPERS}}', papers_json)
            
            # 调用AI生成概念学习笔记
            writer_api_config: APIConfig = get_writer_api_config(self.config)
            
            # 设置系统提示词
            system_prompt = """你是一位学术研究专家，专门研究概念的历史发展和理论演化。请基于提供的种子论文，生成一个关于指定概念的全面学习笔记，并以JSON格式返回。"""
            
            self.logger.info("正在调用AI生成概念学习笔记...")
            
            # 使用ai_interface.py中的健壮API调用函数
            from ai_interface import _call_ai_api  # type: ignore
            concept_profile = _call_ai_api(
                prompt=final_prompt,
                api_config=writer_api_config,
                system_prompt=system_prompt,
                max_tokens=4000,
                temperature=0.7,
                response_format="json"
            )
            
            if concept_profile:
                self.logger.success(f"概念学习笔记生成成功")
                return concept_profile  # type: ignore
            else:
                self.logger.error("概念学习笔记生成失败")
                return {}  # type: ignore
            
        except Exception as e:
            self.logger.error(f"生成概念配置失败: {e}")
            return {}  # type: ignore
    
    
    
    
    def run_concept_priming(self, seed_papers_folder: str, concept_name: str) -> bool:
        """运行概念学习阶段（保留旧函数名以兼容）"""
        return self.run_priming_phase(concept_name, seed_papers_folder)
        
        
        


def sanitize_path_component(path_component: str) -> str:
    """清理路径组件，移除或替换非法字符"""
    import re
    if not path_component:
        return "unnamed"
    
    # 移除或替换Windows路径中的非法字符
    # Windows不允许的字符: < > : " | ? * 以及控制字符
    sanitized = re.sub(r'[<>:"|?*\x00-\x1f]', '_', path_component)
    
    # 移除开头和结尾的空格和点（Windows不允许）
    sanitized = sanitized.strip(' .')
    
    # 确保名称不为空
    if not sanitized:
        sanitized = "unnamed"
    
    # 限制长度（Windows路径限制）
    if len(sanitized) > 100:
        sanitized = sanitized[:100]
    
    return sanitized

def dispatch_command(args: argparse.Namespace):  # type: ignore
    """命令分派器 - 根据参数调用相应的处理函数"""
    try:
        runtime = detect_runtime_environment()
        print(f"[环境] 当前 Python 环境: {runtime.display_name}")
        if runtime.needs_isolation_recommendation:
            print("[环境] 建议使用独立 conda 环境，避免和现有环境中的包互相冲突。")
            print(f"[环境]   {recommended_conda_create_command()}")
            print(f"[环境]   {recommended_conda_activate_command()}")

        cleanup_requested = bool(getattr(args, "cleanup", False))
        setup_requested = bool(getattr(args, "setup", False))
        prime_with_folder = getattr(args, "prime_with_folder", None)
        concept = getattr(args, "concept", None)
        retry_failed = bool(getattr(args, "retry_failed", False))
        merge_target = getattr(args, "merge", None)
        project_name = getattr(args, "project_name", None)
        pdf_folder = getattr(args, "pdf_folder", None)

        # 清理模式
        if cleanup_requested:
            handle_cleanup_mode(args)
            return
        
        # 检查是否为安装模式
        if setup_requested:
            run_setup_wizard()
            return
        
        # 概念学习模式（Priming Phase）
        if prime_with_folder and concept:
            # 检查是否提供了项目名称
            if not project_name:
                logging.error("概念学习模式需要指定 --project-name 参数")
                sys.exit(1)
            
            generator = LiteratureReviewGenerator(args.config, project_name, None, getattr(args, "queue_file", "output/_queue/queue.json"))
            generator.logger.info("*** 概念学习模式已启动 ***")
            generator.logger.info("=" * 60)
            
            if not generator.load_configuration():
                generator.logger.error("配置加载失败")
                sys.exit(1)
            
            # 设置输出目录
            if not generator.setup_output_directory():
                generator.logger.error("输出目录设置失败")
                sys.exit(1)
            
            # 执行概念学习阶段
            success = generator.run_priming_phase(concept, prime_with_folder)
            if success:
                generator.logger.success("概念学习阶段完成！概念配置文件已生成")
            else:
                generator.logger.error("概念学习阶段失败")
                sys.exit(1)
            return
        
        # 重试模式
        if retry_failed:
            handle_retry_failed(args)
            return
        
        # 合并模式
        if merge_target:
            handle_merge_mode(args)
            return
        
        # 大纲采纳命令
        if hasattr(args, 'adopt_outline_v2') and args.adopt_outline_v2:
            handle_outline_adopt_v2_mode(args)
            return

        if hasattr(args, 'outline_adopt') and args.outline_adopt:
            handle_outline_adopt_mode(args)
            return
        
        # 正常执行模式 - 验证参数
        if not project_name and not pdf_folder:
            logging.error("必须指定--project-name或--pdf-folder参数中的一个")
            sys.exit(1)
        
        # 验证project_name格式
        if project_name:
            # 检查是否可能是完整路径（常见错误）
            if len(project_name) > 100 or '\\' in project_name or '/' in project_name:
                logging.error("❌ --project-name 参数错误")
                logging.error("💡 请不要使用完整路径，应该使用简洁的项目名称")
                logging.error("📝 示例：--project-name \"案例分析\" 而非 --project-name \"C:\\Users\\123\\Desktop\\我的项目\"")
                logging.error("🔄 或者使用 --pdf-folder 指定PDF文件夹路径")
                sys.exit(1)
            
                # project_name length warning
            if len(project_name) > 50:
                logging.warning(f"Project name is quite long ({len(project_name)} chars); a shorter name is recommended.")
            
        from services.job_runner import JobRunner, validate_job_request_options
        from services.workflow_facade import build_job_request

        request = build_job_request(args)
        request_error = validate_job_request_options(request)
        if request_error:
            logging.error(request_error)
            sys.exit(1)
        result = JobRunner().run(request, cancel_token=getattr(args, "_cancel_token", None))
        if not result.success:
            sys.exit(result.exit_code or 1)
            
    except KeyboardInterrupt:
        logging.info("用户中断程序")
        sys.exit(1)
    except Exception as e:
        logging.error(f"程序运行失败: {e}")
        logging.error("=" * 60)
        logging.error("详细错误信息:")
        logging.error(traceback.format_exc())
        logging.error("=" * 60)

        # 检查是否为网络相关异常
        import requests  # type: ignore
        if isinstance(e, (requests.exceptions.ConnectionError, requests.exceptions.Timeout, requests.exceptions.RequestException)):
            logging.error("检测到网络连接中断。")
            logging.error("不用担心，您的进度已被保存。")
            logging.error("请在网络恢复后，重新运行您刚才使用的命令，程序将从中断的地方继续。")
        else:
            logging.error("请检查配置文件、网络连接和文件路径是否正确")

        sys.exit(1)

def parse_failure_report(failure_report_file: str, pdf_folder: Optional[str] = None) -> List[PaperInfo]:  # type: ignore
    """从失败报告文件中解析失败的论文信息"""
    try:
        with open(failure_report_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        papers: List[PaperInfo] = []
        
        # 查找论文标题
        import re
        title_pattern = r'📄 标题:\s*(.+?)(?:\r?\n|$)'
        title_matches: List[str] = re.findall(title_pattern, content)
        
        for title in title_matches:
            title = title.strip()
            if title:
                logging.info(f"从失败报告中提取到论文标题: {title}")
                
                # PDF文件夹路径已经作为参数传入
                logging.info(f"PDF文件夹路径: {pdf_folder}")
                
                # 如果找到了PDF文件夹，在其中搜索
                pdf_path = None
                if pdf_folder and os.path.exists(pdf_folder):
                    import glob
                    # 在文件夹中搜索包含标题的PDF文件
                    pattern = os.path.join(pdf_folder, '**', '*.pdf')
                    all_pdfs = glob.glob(pattern, recursive=True)
                    
                    logging.info(f"在PDF文件夹中找到 {len(all_pdfs)} 个PDF文件")
                    
                    for pdf_file in all_pdfs:
                        pdf_filename = os.path.splitext(os.path.basename(pdf_file))[0]
                        
                        # 通用的匹配逻辑：检查标题和PDF文件名的相似度
                        # 方法1：检查作者姓名（如果标题中有下划线分隔作者）
                        author_match = False
                        if '_' in title:
                            possible_authors = title.split('_')[-1].strip()
                            if possible_authors and possible_authors in pdf_filename:
                                author_match = True
                                logging.info(f"基于作者姓名匹配: {possible_authors}")
                        
                        # 方法2：提取标题中的关键词进行匹配
                        # 去除常见停用词，提取有意义的词汇
                        def extract_keywords(text: str) -> List[str]:
                            """从文本中提取关键词（去除停用词）"""
                            # 常见停用词（中英文）
                            stop_words = {'的', '与', '和', '及', '在', '对', '为', '了', '中', '是', '有', '也', '就', '都',
                                         'the', 'and', 'or', 'in', 'on', 'at', 'for', 'to', 'of', 'a', 'an', 'the',
                                         '研究', '分析', '探讨', '初探', '思考', '基于', '视角'}
                            
                            # 分割成词汇（按非字母数字字符分割）
                            import re
                            words = re.findall(r'[\u4e00-\u9fff]+|[a-zA-Z]+', text)
                            
                            # 过滤停用词，保留长度>=2的词汇
                            keywords = [word for word in words if len(word) >= 2 and word not in stop_words]
                            return keywords
                        
                        title_keywords = extract_keywords(title)
                        filename_keywords = extract_keywords(pdf_filename)
                        
                        # 计算关键词重叠度（使用提取的关键词进行更准确的匹配）
                        keyword_overlap = 0
                        matched_words: List[str] = []
                        for keyword in title_keywords:
                            # 方法1：检查精确匹配（关键词在文件名关键词列表中）
                            if keyword in filename_keywords:
                                keyword_overlap += 1
                                matched_words.append(f"[精确]{keyword}")
                            # 方法2：检查子字符串匹配（关键词在PDF文件名中）
                            elif keyword in pdf_filename:
                                keyword_overlap += 1
                                matched_words.append(f"[包含]{keyword}")
                        
                        # 方法3：计算文本相似度（简单版本）
                        def calculate_similarity(str1: str, str2: str) -> float:
                            """计算两个字符串的相似度（基于重叠字符）"""
                            # 转换为集合（去除重复字符）
                            set1 = set(str1)
                            set2 = set(str2)
                            if not set1 or not set2:
                                return 0.0
                            # Jaccard相似度
                            intersection = len(set1.intersection(set2))
                            union = len(set1.union(set2))
                            return intersection / union if union > 0 else 0.0
                        
                        similarity_score = calculate_similarity(title, pdf_filename)
                        
                        # 匹配条件：作者匹配 或 关键词匹配>=2 或 相似度>0.5
                        if author_match or keyword_overlap >= 2 or similarity_score > 0.5:
                            pdf_path = pdf_file
                            logging.info(f"成功匹配PDF文件: {pdf_file}")
                            if author_match:
                                logging.info("匹配原因: 作者姓名")
                            if keyword_overlap > 0:
                                logging.info(f"匹配到 {keyword_overlap} 个关键词: {matched_words}")
                            if similarity_score > 0.5:
                                logging.info(f"文本相似度: {similarity_score:.2f}")
                            break
                        
                        # 方法4：直接包含检查（如果PDF文件名包含标题的主要部分）
                        else:
                            clean_title = title.replace('——', '').replace('_', '').replace('"', '').replace('（', '').replace('）', '')
                            clean_filename = pdf_filename.replace('_', '').replace('"', '').replace('（', '').replace('）', '')
                            
                            # 如果标题长度>10且被文件名包含，或相反
                            if len(clean_title) > 10 and clean_title in clean_filename:
                                pdf_path = pdf_file
                                logging.info(f"基于整体字符串匹配找到PDF文件: {pdf_file}")
                                break
                            elif len(clean_filename) > 10 and clean_filename in clean_title:
                                pdf_path = pdf_file
                                logging.info(f"基于反向包含匹配找到PDF文件: {pdf_file}")
                                break
                
                # 如果找到了PDF文件，创建论文信息
                if pdf_path and os.path.exists(pdf_path):
                    paper_info: PaperInfo = {
                        'title': title,
                        'authors': [],
                        'year': '未知年份',
                        'journal': '未知期刊',
                        'doi': '',
                        'pdf_path': pdf_path,
                        'file_index': 0
                    }
                    papers.append(paper_info)
                    logging.info(f"成功创建失败论文的重试信息: {title}")
                else:
                    logging.warning(f"未找到论文标题对应的PDF文件: {title}")
                    logging.info(f"PDF文件夹: {pdf_folder}")
                    logging.info(f"PDF文件夹是否存在: {os.path.exists(pdf_folder) if pdf_folder else 'None'}")
        
        # 如果还是没有找到PDF路径，查找PDF文件路径的模式
        if not papers:
            pdf_pattern = r'PDF文件不存在:\s*(.+\.pdf)'
            pdf_matches: List[str] = re.findall(pdf_pattern, content)
            
            for pdf_path in pdf_matches:
                pdf_path = pdf_path.strip()
                if pdf_path and os.path.exists(pdf_path):
                    title = os.path.splitext(os.path.basename(pdf_path))[0]
                    
                    paper_info: PaperInfo = {
                        'title': title,
                        'authors': [],
                        'year': '未知年份',
                        'journal': '未知期刊',
                        'doi': '',
                        'pdf_path': pdf_path,
                        'file_index': 0
                    }
                    papers.append(paper_info)
        
        return papers
        
    except Exception as e:
        logging.error(f"解析失败报告文件出错: {e}")
        return []

def handle_retry_failed(args: argparse.Namespace):  # type: ignore
    """处理重试失败论文模式"""
    if not args.project_name and not args.pdf_folder:
        logging.error("使用--retry-failed命令时必须提供--project-name或--pdf-folder参数中的一个")
        sys.exit(1)
    
    # 验证project_name格式
    if args.project_name:
        if len(args.project_name) > 100 or '\\' in args.project_name or '/' in args.project_name:
            logging.error("❌ --project-name 参数错误")
            logging.error("💡 请不要使用完整路径，应该使用简洁的项目名称")
            logging.error("📝 示例：--project-name \"案例分析\" 而非 --project-name \"C:\\Users\\123\\Desktop\\我的项目\"")
            sys.exit(1)

    generator = LiteratureReviewGenerator(args.config, args.project_name, args.pdf_folder, getattr(args, "queue_file", "output/_queue/queue.json"))
    generator.logger.info("*** 失败论文重试模式已启动 ***")
    
    if not generator.load_configuration() or not generator.setup_output_directory():
        sys.exit(1)

    if not generator.load_existing_summaries():
        generator.logger.error("未找到摘要文件，无法进行重试。请先运行一次完整的分析。")
        sys.exit(1)

    papers_to_retry = []
    retry_report_file = ''  # 初始化变量
    if generator.mode == "zotero":
        retry_report_file = generator._get_report_file_path('_zotero_report_for_retry.txt')
        if not os.path.exists(retry_report_file):  # type: ignore:
            generator.logger.error(f"Zotero模式重试失败：未找到重跑报告文件 '{retry_report_file}'")
            sys.exit(1)
        papers_to_retry = parse_zotero_report(retry_report_file)  # type: ignore
    else:  # direct mode
        generator.logger.info("直接PDF模式：正在从摘要文件和失败报告中识别失败的论文...")
        
        # 首先尝试从summaries.json中查找失败的论文
        failed_summaries = [s for s in generator.summaries if s.get('status') == 'failed']  # type: ignore
        papers_to_retry = [s.get('paper_info') for s in failed_summaries if s.get('paper_info')]  # type: ignore
        
        # 如果没有在summaries.json中找到失败的论文，尝试从失败报告文件中读取
        if not papers_to_retry:
            generator.logger.info("在summaries.json中未找到失败的论文，正在检查失败报告...")
            failure_report_file = generator._get_report_file_path('_failed_papers_report.txt')
            
            if os.path.exists(failure_report_file):
                generator.logger.info(f"找到失败报告文件: {failure_report_file}")
                try:
                    # 解析失败报告文件，传入PDF文件夹路径
                    failed_papers_from_report = parse_failure_report(failure_report_file, generator.pdf_folder)
                    if failed_papers_from_report:
                        papers_to_retry = failed_papers_from_report
                        generator.logger.info(f"从失败报告中提取到 {len(papers_to_retry)} 篇需要重试的论文")
                    else:
                        generator.logger.warning("失败报告文件存在但无法解析")
                except Exception as e:
                    generator.logger.error(f"读取失败报告文件失败: {e}")
            else:
                generator.logger.warning(f"未找到失败报告文件: {failure_report_file}")

    if not papers_to_retry:
        generator.logger.success("没有找到需要重试的失败论文。")
        return

    generator.logger.info(f"识别到 {len(papers_to_retry)} 篇论文需要重试。")
    
    original_summary_count = len(generator.summaries)
    file_index_path: str = generator.config.get('Paths', {}).get('library_path', '') if generator.mode == 'zotero' and generator.config else generator.pdf_folder or ''  # type: ignore
    file_index = create_file_index(file_index_path)  # type: ignore
    performance_config = generator.config.get('Performance') or {}  # type: ignore
    max_workers = int(performance_config.get('max_workers', 3))  # type: ignore

    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:  # type: ignore
        future_to_paper = {executor.submit(generator.process_paper, paper, i, file_index, len(papers_to_retry)): paper for i, paper in enumerate(papers_to_retry)}  # type: ignore
        progress_bar = tqdm(concurrent.futures.as_completed(future_to_paper), total=len(papers_to_retry), desc="[重试模式] 正在处理")  # type: ignore
        for future in progress_bar:
            result: Optional[Dict[str, Any]] = future.result()  # type: ignore
            if result and result.get('status') == 'success':  # type: ignore
                # 在直接PDF模式下，更新原始条目而不是添加新条目
                if generator.mode == "direct":
                    paper_key = LiteratureReviewGenerator.get_paper_key(result.get('paper_info', {}))  # type: ignore
                    # 查找并更新原始条目
                    for i, summary in enumerate(generator.summaries):
                        if LiteratureReviewGenerator.get_paper_key(summary.get('paper_info', {})) == paper_key:  # type: ignore
                            generator.summaries[i] = result  # type: ignore
                            break
                    else:
                        # 如果没有找到原始条目，则添加新条目
                        generator.summaries.append(result)  # type: ignore
                else:
                    # Zotero模式下，直接添加新条目
                    generator.summaries.append(result)  # type: ignore
            else:
                # 处理失败的论文
                failed_paper: Dict[str, Any] = result or {'paper_info': future_to_paper[future], 'failure_reason': '未知重试错误'}  # type: ignore
                if generator.mode == "direct":
                    paper_key = LiteratureReviewGenerator.get_paper_key(failed_paper.get('paper_info', {}))
                    # 查找并更新原始条目
                    for i, summary in enumerate(generator.summaries):
                        if LiteratureReviewGenerator.get_paper_key(summary.get('paper_info', {})) == paper_key:
                            generator.summaries[i] = failed_paper  # type: ignore
                            break
                    else:
                        # 如果没有找到原始条目，则添加新条目
                        generator.summaries.append(failed_paper)  # type: ignore
                    
                    # 确保失败的论文也被添加到failed_papers列表，以便生成失败报告
                    generator.failed_papers.append(failed_paper)  # type: ignore
                else:
                    # Zotero模式下，直接添加到失败列表
                    generator.failed_papers.append(failed_paper)  # type: ignore

    generator.save_summaries()
    
    # 调用统一的报告生成方法
    generator.generate_all_reports()

    # 计算新增成功的论文数量
    success_count = len([s for s in generator.summaries if s.get('status') == 'success'])  # type: ignore
    original_success = len([s for s in generator.summaries[:original_summary_count] if s.get('status') == 'success'])  # type: ignore
    newly_succeeded = success_count - original_success
    failed_count = len([s for s in generator.summaries if s.get('status') == 'failed'])  # type: ignore
    generator.logger.success(f"重试完成！新增成功 {newly_succeeded} 篇，仍然失败 {failed_count} 篇。")  # type: ignore
    
    if not generator.failed_papers and generator.mode == 'zotero' and os.path.exists(retry_report_file):
        try:
            os.remove(retry_report_file)
            generator.logger.info(f"所有失败论文均已成功重试，已自动删除重跑报告文件: {retry_report_file}")
        except Exception as e:
            generator.logger.warning(f"无法自动删除重跑报告文件: {e}")

def handle_merge_mode(args: argparse.Namespace):  # type: ignore
    """处理合并模式"""
    # 验证参数：必须提供project_name或pdf_folder中的一个
    if not args.project_name and not args.pdf_folder:
        logging.error("使用--merge命令时必须提供--project-name或--pdf-folder参数中的一个")
        sys.exit(1)
    
    # 验证project_name格式
    if args.project_name:
        if len(args.project_name) > 100 or '\\' in args.project_name or '/' in args.project_name:
            logging.error("❌ --project-name 参数错误")
            logging.error("💡 请不要使用完整路径，应该使用简洁的项目名称")
            logging.error("📝 示例：--project-name \"案例分析\" 而非 --project-name \"C:\\Users\\123\\Desktop\\我的项目\"")
            sys.exit(1)
    
    generator = LiteratureReviewGenerator(args.config, args.project_name, args.pdf_folder, getattr(args, "queue_file", "output/_queue/queue.json"))
    generator.logger.info("*** 合并模式已启动 ***")
    generator.logger.info("=" * 60)
    
    # 根据模式确定项目名称和文件路径
    try:
        # 加载配置以获取输出路径
        if not generator.load_configuration():
            generator.logger.error("配置加载失败")
            sys.exit(1)
        
        # 设置输出目录以确定项目名称
        if not generator.setup_output_directory():
            generator.logger.error("输出目录设置失败")
            sys.exit(1)
        
        # 确定主文件路径
        main_file = generator.summary_file
        merge_file = args.merge
        
        if not main_file or not os.path.exists(main_file):
            generator.logger.error(f"主文件不存在: {main_file}")
            return
        
        if not os.path.exists(merge_file):
            generator.logger.error(f"合并文件不存在: {merge_file}")
            return
        
        try:
            main_data = load_json_file_with_fallbacks(main_file, logger=generator.logger)  # type: ignore
        except Exception as exc:
            generator.logger.error(f"读取主文件失败: {main_file} - {exc}")
            return
        
        try:
            merge_data = load_json_file_with_fallbacks(merge_file, logger=generator.logger)  # type: ignore
        except Exception as exc:
            generator.logger.error(f"读取合并文件失败: {merge_file} - {exc}")
            return
        
        if not isinstance(main_data, list) or not isinstance(merge_data, list):  # type: ignore
            generator.logger.error("文件格式错误，必须是JSON数组")
            return
        
        # 智能合并：以合并文件中的记录为准
        generator.logger.info(f"主文件包含 {len(main_data)} 篇论文")  # type: ignore
        generator.logger.info(f"合并文件包含 {len(merge_data)} 篇论文")  # type: ignore
        
        # 创建基于DOI的索引（如果没有DOI则使用标题+作者）
        def get_paper_key(paper: 'Dict[str, Any] | PaperInfo'):  # type: ignore
            paper_info = paper.get('paper_info', {})  # type: ignore
            return paper_info.get('doi', f"{paper_info.get('title', '')}_{paper_info.get('authors', [])}")  # type: ignore
        
        # 构建主文件的索引
        main_index = {get_paper_key(paper): i for i, paper in enumerate(main_data)}  # type: ignore
        
        # 合并数据
        merged_count = 0
        added_count = 0
        
        for merge_paper in merge_data:  # type: ignore
            merge_key = get_paper_key(merge_paper)  # type: ignore
            
            if merge_key in main_index:
                # 更新现有记录
                main_index_pos = main_index[merge_key]
                main_data[main_index_pos] = merge_paper  # type: ignore
                merged_count += 1
            else:
                # 添加新记录
                main_data.append(merge_paper)  # type: ignore
                added_count += 1
        
        # 保存合并结果
        backup_file: str = f"{main_file}.backup.{int(time.time())}"  # type: ignore
        os.rename(main_file, backup_file)  # type: ignore
        generator.logger.info(f"已创建备份文件: {backup_file}")  # type: ignore
        
        with open(main_file, 'w', encoding='utf-8') as f:  # type: ignore
            json.dump(main_data, f, ensure_ascii=False, indent=2)  # type: ignore
        
        generator.logger.success("合并完成！")  # type: ignore
        generator.logger.info(f"更新记录: {merged_count} 篇")  # type: ignore
        generator.logger.info(f"新增记录: {added_count} 篇")  # type: ignore
        generator.logger.info(f"总记录数: {len(main_data)} 篇")  # type: ignore
        
    except Exception as e:
        generator.logger.error(f"合并过程中出错: {e}")
        traceback.print_exc()

def handle_run_all_mode(generator: 'LiteratureReviewGenerator'):  # type: ignore
    """处理一键执行模式"""
    generator.logger.info("*** '一键执行'模式已启动 ***")
    generator.logger.info("=" * 60)
    
    # 执行阶段一
    generator.logger.info("开始执行阶段一：文献分析...")
    stage1_success = generator.run_stage_one()
    
    if stage1_success:
        generator.logger.success("\n阶段一执行成功！")
        generator.logger.info("开始执行阶段二：文献综述生成...")
        
        # 执行阶段二：先生成大纲，再生成全文
        generator.logger.info("开始执行阶段二第一步：生成大纲...")
        outline_success = generator.generate_literature_review_outline()
        
        if outline_success:
            generator.logger.success("大纲生成成功！")
            generator.logger.info("开始执行阶段二第二步：从大纲生成全文...")
            stage2_success = generator.generate_full_review_from_outline()
        else:
            stage2_success = False
        
        if stage2_success:
            generator.logger.success("\n一键执行模式完成！所有任务执行成功！")
        else:
            generator.logger.error("\n阶段二执行失败！")
            sys.exit(1)
    else:
        generator.logger.error("\n阶段一执行失败，无法继续执行阶段二！")
        sys.exit(1)

def handle_generate_outline_mode(generator: 'LiteratureReviewGenerator', args: argparse.Namespace):  # type: ignore
    """处理生成大纲模式"""
    success = generator.generate_literature_review_outline()
    if success:
        generator.logger.success("\n大纲生成成功！文献综述大纲已生成完成")
        generator.logger.info(f"您可以编辑大纲文件，然后运行以下命令生成完整综述：")
        if args.project_name:
            # 检查是否是概念模式
            if args.concept:
                generator.logger.info(f"命令: python main.py --project-name \"{args.project_name}\" --concept \"{args.concept}\" --generate-review")
            else:
                generator.logger.info(f"命令: python main.py --project-name \"{args.project_name}\" --generate-review")
        elif args.pdf_folder:
            # 检查是否是概念模式
            if args.concept:
                generator.logger.info(f"命令: python main.py --pdf-folder \"{args.pdf_folder}\" --concept \"{args.concept}\" --generate-review")
            else:
                generator.logger.info(f"命令: python main.py --pdf-folder \"{args.pdf_folder}\" --generate-review")
    else:
        generator.logger.error("\n大纲生成失败！")
        sys.exit(1)

def handle_generate_review_mode(generator: 'LiteratureReviewGenerator'):  # type: ignore
    """处理生成综述模式"""
    success = generator.generate_full_review_from_outline()
    if success:
        generator.logger.success("\n文献综述生成成功！完整综述已生成完成")
    else:
        generator.logger.error("\n文献综述生成失败！")
        sys.exit(1)

def handle_generate_section_mode(generator: 'LiteratureReviewGenerator', args: argparse.Namespace):  # type: ignore
    """Handle generation for a single outline section."""

    section_number = int(getattr(args, "generate_section", 0) or 0)
    if section_number <= 0:
        generator.logger.error("generate-section 参数必须是大于 0 的整数")
        sys.exit(1)

    success = generator.generate_specific_review_section(section_number)
    if success:
        generator.logger.success(f"\n第 {section_number} 章已补写完成")
    else:
        generator.logger.error(f"\n第 {section_number} 章补写失败")
        sys.exit(1)


def handle_retry_review_failed_mode(generator: 'LiteratureReviewGenerator'):  # type: ignore
    """Retry failed stage-two review sections from the saved failure record."""

    success = generator.retry_failed_review_sections()
    if success:
        generator.logger.success("\n失败章节重试完成")
    else:
        generator.logger.error("\n失败章节重试失败")
        sys.exit(1)


def handle_stage_one_mode(generator: 'LiteratureReviewGenerator', args: argparse.Namespace):  # type: ignore
    """处理阶段一模式（默认模式）"""
    generator.logger.info("*** 阶段一模式已启动 ***")
    generator.logger.info("=" * 60)
    
    # 执行阶段一
    generator.logger.info("开始执行阶段一：文献分析...")
    stage1_success = generator.run_stage_one()
    
    if stage1_success:
        generator.logger.success("\n阶段一执行成功！")
        generator.logger.info("您现在可以继续执行以下命令：")
        if args.project_name:
            # 检查是否是概念模式
            if args.concept:
                generator.logger.info(f"生成大纲: python main.py --project-name \"{args.project_name}\" --concept \"{args.concept}\" --generate-outline")
                generator.logger.info(f"一键生成综述: python main.py --project-name \"{args.project_name}\" --concept \"{args.concept}\" --run-all")
            else:
                generator.logger.info(f"生成大纲: python main.py --project-name \"{args.project_name}\" --generate-outline")
                generator.logger.info(f"一键生成综述: python main.py --project-name \"{args.project_name}\" --run-all")
        elif args.pdf_folder:
            # 检查是否是概念模式
            if args.concept:
                generator.logger.info(f"生成大纲: python main.py --pdf-folder \"{args.pdf_folder}\" --concept \"{args.concept}\" --generate-outline")
                generator.logger.info(f"一键生成综述: python main.py --pdf-folder \"{args.pdf_folder}\" --concept \"{args.concept}\" --run-all")
            else:
                generator.logger.info(f"生成大纲: python main.py --pdf-folder \"{args.pdf_folder}\" --generate-outline")
                generator.logger.info(f"一键生成综述: python main.py --pdf-folder \"{args.pdf_folder}\" --run-all")
    else:
        generator.logger.error("\n阶段一执行失败！")
        sys.exit(1)


def main() -> None:  # type: ignore
    """主函数，处理命令行参数和执行相应操作"""
    
    parser = argparse.ArgumentParser(
        description="auto-generate - AI 文献分析与综述写作工作台\n"
                   "支持本地 GUI 和 CLI 两种模式；CLI 直接运行，GUI 工作台自动管理后台队列。",
        formatter_class=argparse.RawTextHelpFormatter,
        epilog="示例:\n"
               "  # 运行交互式设置\n"
               "  python main.py --setup\n"
               "\n"
               "  # 使用 PDF 文件夹分析文献\n"
               "  python main.py --pdf-folder \"D:\\papers\" --analyze-only\n"
               "\n"
               "  # 一键运行所有阶段\n"
               "  python main.py --project-name \"my_review\" --run-all\n"
    )
    
    # 基础选项
    base_group = parser.add_argument_group('基础选项')
    base_group.add_argument(
        '--config',
        type=str,
        default='config.ini',
        help='配置文件路径（默认：config.ini）'
    )
    base_group.add_argument(
        '--project-name', '-p',
        type=str, 
        help='项目名称（用于创建独立的输出文件夹）'
    )
    base_group.add_argument(
        '--pdf-folder', '-f',
        type=str, 
        help='PDF 文件夹路径（直接扫描该文件夹处理文献）'
    )
    
    # 工作流选项
    workflow_group = parser.add_argument_group('工作流选项')
    workflow_group.add_argument(
        '--setup',
        action='store_true', 
        help='运行交互式设置向导'
    )
    workflow_group.add_argument(
        '--run-all', '-a',
        action='store_true', 
        help='一键运行：分析 → 大纲 → 综述'
    )
    workflow_group.add_argument(
        '--analyze-only', '-A',
        action='store_true', 
        help='仅运行阶段一：文献分析'
    )
    workflow_group.add_argument(
        '--generate-outline', '-o', '--outline',
        action='store_true', 
        help='仅运行阶段二：生成大纲'
    )
    workflow_group.add_argument(
        '--generate-review', '-r',
        action='store_true', 
        help='仅运行阶段三：生成综述'
    )
    workflow_group.add_argument(
        '--validate-review', '-v',
        action='store_true',
        help='验证生成的综述'
    )
    workflow_group.add_argument(
        '--outline-adopt',
        action='store_true',
        help='显式采纳 Week5 大纲仲裁结果'
    )
    workflow_group.add_argument(
        '--adopt-outline-v2',
        action='store_true',
        help='显式采纳 Outline Intelligence v2 final_outline.json（写入 adopted_final_outline.json）'
    )
    workflow_group.add_argument(
        '--cleanup',
        action='store_true',
        help='清理旧工作空间，只保留最新的'
    )
    
    # 高级选项
    advanced_group = parser.add_argument_group('高级选项')
    advanced_group.add_argument(
        '--prime-with-folder',
        type=str,
        help='概念预热：指定种子论文文件夹'
    )
    advanced_group.add_argument(
        '--concept',
        type=str,
        help='概念预热：指定概念名称'
    )
    advanced_group.add_argument(
        '--retry-failed',
        action='store_true',
        help='重试失败的论文'
    )
    advanced_group.add_argument(
        '--generate-section',
        type=int,
        help='仅生成指定章节号'
    )
    advanced_group.add_argument(
        '--retry-review-failed',
        action='store_true',
        help='重试失败的综述章节'
    )
    advanced_group.add_argument(
        '--free-mode-profile',
        type=str,
        help='Free mode: load a profile JSON file'
    )
    advanced_group.add_argument(
        '--free-mode-idea',
        type=str,
        help='Free mode: pass an idea as plain text'
    )
    advanced_group.add_argument(
        '--merge',
        type=str,
        help='Merge another summaries.json file into the current summaries file'
    )
    advanced_group.add_argument(
        '--summary-file',
        type=str,
        help='Explicitly select a summaries.json file for outline/review/section/validation'
    )
    advanced_group.add_argument(
        '--summary-source',
        dest='summary_sources',
        action='append',
        help='Add a summaries.json file to the current downstream summary source set (repeatable)'
    )
    advanced_group.add_argument(
        '--reuse-stage1',
        action='store_true',
        help='Enable DOI-only reuse of earlier stage-1 summaries during stage-1 analysis'
    )
    advanced_group.add_argument(
        '--reuse-summary-file',
        action='append',
        help='Add an extra summaries.json file to the stage-1 reuse pool (repeatable)'
    )
    
    # Zotero选项
    zotero_group = parser.add_argument_group('Zotero选项')
    zotero_group.add_argument(
        '--zotero-report',
        type=str,
        help='Zotero报告文件路径'
    )
    zotero_group.add_argument(
        '--library-path',
        type=str,
        help='Zotero库路径'
    )

    args = parser.parse_args()
    dispatch_command(args)

def handle_cleanup_mode(args: argparse.Namespace):  # type: ignore
    """清理旧工作空间模式"""
    try:
        from utils import sanitize_path_component
        
        # 确定项目名称
        project_name = None
        if args.project_name:
            project_name = sanitize_path_component(args.project_name)
        elif args.pdf_folder:
            project_name = sanitize_path_component(os.path.basename(os.path.abspath(args.pdf_folder.rstrip("/\\"))))
        
        if not project_name:
            logging.error("必须指定 --project-name 或 --pdf-folder 参数中的一个")
            sys.exit(1)
        
        # 加载配置以获取输出路径
        temp_generator = LiteratureReviewGenerator(args.config, project_name, args.pdf_folder, "output/_queue/queue.json")
        if not temp_generator.load_configuration():
            logging.error("配置加载失败")
            sys.exit(1)
        
        paths_config = temp_generator.config.get('Paths', {}) if temp_generator.config else {}
        output_base_path = paths_config.get('output_path', './output')
        output_base_path_abs = os.path.abspath(output_base_path)
        
        # 查找所有工作空间
        import glob
        import shutil
        workspace_pattern = os.path.join(output_base_path_abs, f"{project_name}__*")
        workspaces = glob.glob(workspace_pattern)
        
        if not workspaces:
            print(f"未找到项目 '{project_name}' 的任何工作空间")
            sys.exit(0)
        
        if len(workspaces) <= 1:
            print(f"项目 '{project_name}' 只有 1 个工作空间，无需清理")
            sys.exit(0)
        
        # 评估每个工作空间的完整性
        def score_workspace(workspace_path: str) -> int:
            """给工作空间打分，分数越高越完整"""
            score = 0
            artifacts_dir = os.path.join(workspace_path, "artifacts")
            reports_dir = os.path.join(workspace_path, "reports")
            
            # 检查重要文件
            if os.path.exists(os.path.join(artifacts_dir, f"{project_name}_summaries.json")):
                score += 10
            if os.path.exists(os.path.join(artifacts_dir, f"{project_name}_literature_review_outline.md")):
                score += 10
            if os.path.exists(os.path.join(reports_dir, f"{project_name}_literature_review.docx")):
                score += 10
            if os.path.exists(os.path.join(reports_dir, f"{project_name}_analyzed_papers.xlsx")):
                score += 5
            
            # 检查文件数量
            try:
                artifact_files = len([f for f in os.listdir(artifacts_dir) if os.path.isfile(os.path.join(artifacts_dir, f))])
                score += min(artifact_files, 10)
            except:
                pass
            
            # 检查修改时间（越新越好，但不如文件完整重要）
            mtime = os.path.getmtime(workspace_path)
            score += int(mtime / 1000000000)  # 时间戳的年分数部分
            
            return score
        
        # 按完整性评分排序
        workspaces_with_scores = [(score_workspace(ws), ws) for ws in workspaces]
        workspaces_with_scores.sort(key=lambda x: x[0], reverse=True)
        
        best_workspace = workspaces_with_scores[0][1]
        other_workspaces = [ws for (score, ws) in workspaces_with_scores[1:]]
        
        print(f"找到 {len(workspaces)} 个工作空间，将保留文件最完整的 1 个，删除其余的 {len(other_workspaces)} 个:")
        print(f"  保留: {os.path.basename(best_workspace)} (得分: {workspaces_with_scores[0][0]})")
        
        deleted_count = 0
        for workspace_path in other_workspaces:
            try:
                print(f"  删除: {os.path.basename(workspace_path)}")
                shutil.rmtree(workspace_path)
                deleted_count += 1
            except Exception as e:
                print(f"  ❌ 删除失败: {os.path.basename(workspace_path)} - {e}")
        
        print(f"\n✅ 清理完成！共删除了 {deleted_count} 个旧工作空间")
        
    except Exception as e:
        logging.error(f"处理清理模式时出错: {e}")
        import traceback
        logging.debug(f"详细错误信息: {traceback.format_exc()}")
        sys.exit(1)

def handle_outline_adopt_mode(args: argparse.Namespace):  # type: ignore
    """处理大纲采纳模式"""
    try:
        # 验证参数
        if not args.project_name and not args.pdf_folder:
            logging.error("必须指定--project-name或--pdf-folder参数中的一个")
            sys.exit(1)
        
        # 初始化生成器
        generator = LiteratureReviewGenerator(args.config, args.project_name, args.pdf_folder, getattr(args, "queue_file", "output/_queue/queue.json"))
        generator.logger.info("=== 大纲采纳模式已启动 ===")
        generator.logger.info("=" * 60)
        
        # 加载配置
        if not generator.load_configuration():
            generator.logger.error("配置加载失败")
            sys.exit(1)
        
        # 设置输出目录
        if not generator.setup_output_directory():
            generator.logger.error("输出目录设置失败")
            sys.exit(1)
        
        # 执行大纲采纳
        if not generator.adopt_outline():
            generator.logger.error("大纲采纳失败")
            sys.exit(1)
        else:
            generator.logger.success("大纲采纳成功！")
    except Exception as e:
        logging.error(f"处理大纲采纳模式时出错: {e}")
        import traceback
        logging.debug(f"详细错误信息: {traceback.format_exc()}")
        sys.exit(1)


def handle_outline_adopt_v2_mode(args: argparse.Namespace):  # type: ignore
    """处理 Outline Intelligence v2 大纲采纳模式"""
    try:
        if not args.project_name and not args.pdf_folder:
            logging.error("必须指定--project-name或--pdf-folder参数中的一个")
            sys.exit(1)

        generator = LiteratureReviewGenerator(
            args.config,
            args.project_name,
            args.pdf_folder,
            getattr(args, "queue_file", "output/_queue/queue.json"),
        )
        generator.logger.info("=== Outline Intelligence v2 采纳模式已启动 ===")
        generator.logger.info("=" * 60)

        if not generator.load_configuration():
            generator.logger.error("配置加载失败")
            sys.exit(1)

        if not generator.setup_output_directory():
            generator.logger.error("输出目录设置失败")
            sys.exit(1)

        if not generator.adopt_outline_v2():
            generator.logger.error("Outline v2 大纲采纳失败")
            sys.exit(1)
        generator.logger.success("Outline v2 大纲采纳成功！")
    except Exception as e:
        logging.error(f"处理 Outline v2 大纲采纳模式时出错: {e}")
        import traceback
        logging.debug(f"详细错误信息: {traceback.format_exc()}")
        sys.exit(1)

if __name__ == "__main__":
    main()
