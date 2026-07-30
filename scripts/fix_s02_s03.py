"""Fix missing sections in S02 and S03 DOCX files."""
import json, os, shutil
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def inspect_docx(path, label):
    doc = Document(path)
    print(f"=== {label} ===")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    
    # Show headings
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            print(f"  [{p.style.name}] {p.text.strip()[:100]}")
    
    # Check tables for content
    for ti, table in enumerate(doc.tables):
        row_count = len(table.rows)
        print(f"  Table {ti}: {row_count} rows")
        if row_count > 0:
            first_cell = table.rows[0].cells[0].text.strip()[:100]
            print(f"    First cell: {first_cell}")

# Check both DOCX files
BUNDLE = r"D:\auto-generate\output\pph_review_bundle_final"

for name in ["02_platform_prior_concession_review", "03_prior_concession_to_unfairness_review"]:
    path = os.path.join(BUNDLE, name + ".docx")
    if os.path.exists(path):
        inspect_docx(path, name)
    print()
