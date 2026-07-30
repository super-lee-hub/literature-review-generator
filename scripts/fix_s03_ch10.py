"""Manually write S03 chapter 10 (conclusion) into the DOCX."""
import shutil, os
from docx import Document

SRC = r"D:\auto-generate\output\pph_s03_concession_to_unfairness__20260728_063445_57ff8d29\reports\pph_s03_concession_to_unfairness_literature_review.docx"
DST = r"D:\auto-generate\output\pph_review_bundle_final\03_prior_concession_to_unfairness_review.docx"
shutil.copy2(SRC, DST)

doc = Document(DST)
doc.add_heading("第10章 结论与理论整合", level=2)

t1 = ("综合上述文献，平台既往让利对后续人际价格劣势情境下价格不公平感的缓释效应，"
      "其理论逻辑可归纳为一条完整的归因修正链。首先，消费者在遭遇人际价格劣势时，"
      "天然倾向于将价格差异归因为平台的过度逐利、机会主义定价或算法剥削"
      "（Campbell, 1999; Xia et al., 2004; Bolton et al., 2003）。"
      "这种归因倾向根植于双重权利原则——消费者认为企业有权获取合理利润，"
      "但不应利用信息不对称或需求波动牟取暴利（Kahneman et al., 1986）。"
      "然而，当消费者知晓或亲历平台曾通过补贴、折扣、免费试用等方式让利时，"
      "既往让利行为传递了平台"非单向获利"的信号：平台曾为获取用户、培育市场而承担真实成本"
      "（Armstrong, 2006; Rochet & Tirole, 2003; Pauwels & Weiss, 2008）。"
      "在此情境下，后续的价格劣势更可能被理解为平台回收前期投入、恢复正常经营利润或维持商业可持续性，"
      "而非纯粹的剥削行为（Effron & Monin, 2010）。")
doc.add_paragraph(t1)

t2 = ("值得强调的是，既往让利的缓释效应并非无条件的。"
      "从竞争性机制角度看，过去的价格优惠可能形成较低的内部参考价格，"
      "反而强化消费者对当前高价的负面反应（Mazumdar et al., 2005; Kalwani et al., 1992）。"
      "此外，频繁促销可能培养消费者的应得权利感，使后续正常定价被视为剥夺"
      "（Xia et al., 2010）。在方法层面，现有研究主要采用情景实验和问卷调查，"
      "以价格公平感知作为核心结果变量，部分研究扩展到购买意愿和负面口碑。"
      "但直接将平台既往让利作为调节变量、以人际价格劣势为自变量的研究尚付阙如。"
      "这构成了本项目的核心研究缺口，也意味着本文提出的理论链条——"
      "既往让利→获利动机归因修正→价格不公平感缓释——虽然各环节均有桥接证据支持，"
      "但整体路径仍需实证检验。")
doc.add_paragraph(t2)

t3 = ("在实践层面，本综述的发现对平台企业具有管理启示。当平台因经营需要调整定价策略、"
      "或算法定价导致消费者间出现价格差异时，主动沟通前期投入和让利历史，"
      "或有助于将消费者的定价动机归因从"剥削"转向"合理经营"。"
      "然而，并非所有价格差异均可借既往让利正当化——涉及弱势群体剥削、"
      "基于受保护特征的价格歧视、或具备市场支配地位的掠夺性定价，"
      "即便存在前期让利，仍应受到伦理和法律的严格审视。"
      "未来研究可从三个方向推进：一是构建并直接检验既往让利缓释效应的完整理论模型；"
      "二是探索让利类型、幅度和时间距离的调节作用；"
      "三是考察不同商业模式（订阅制、交易平台、双边市场）下该效应的异质性。")
doc.add_paragraph(t3)

doc.save(DST)
print("S03 DOCX updated: " + str(os.path.getsize(DST)) + " bytes")

md_path = DST.replace(".docx", ".md")
with open(md_path, "w", encoding="utf-8") as f:
    f.write("# Prior Concession to Price Unfairness Review\n\n")
    f.write("## 第10章 结论与理论整合\n\n")
    f.write(t1 + "\n\n")
    f.write(t2 + "\n\n")
    f.write(t3 + "\n")
print("S03 MD updated: " + md_path)
