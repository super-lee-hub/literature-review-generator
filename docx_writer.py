"""Canonical DOCX renderer for the current review artifact contract."""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Dict, List, Mapping

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from services.citation_catalog import CitationCatalogEntry, format_in_text_citation, format_reference_entry
from services.citation_ref_catalog import LEGAL_CITE_REF_TOKEN_PATTERN, extract_ref_ids_from_token


def _log(logger: Any, level: str, message: str) -> None:
    method = getattr(logger, level, None) or getattr(logger, "info", None)
    if callable(method):
        method(message)


def _entry_lookup(manifest: Mapping[str, Any]) -> dict[str, CitationCatalogEntry]:
    entries: dict[str, CitationCatalogEntry] = {}
    for raw in manifest.get("paper_entries", []):
        if not isinstance(raw, Mapping):
            continue
        paper_id = str(raw.get("paper_id") or raw.get("paper_key") or "").strip()
        paper_key = str(raw.get("paper_key") or paper_id).strip()
        if not paper_id:
            continue
        entry = CitationCatalogEntry(
            index=len(entries) + 1,
            paper_id=paper_id,
            paper_key=paper_key,
            title=str(raw.get("title") or ""),
            authors=[str(item).strip() for item in raw.get("authors", []) if str(item).strip()],
            year=str(raw.get("year") or ""),
            journal=str(raw.get("journal") or ""),
            doi=str(raw.get("doi") or ""),
            aliases=[str(item) for item in raw.get("aliases", [])],
        )
        entries[paper_id] = entry
        entries[paper_key] = entry
    by_ref: dict[str, CitationCatalogEntry] = {}
    for occurrence in manifest.get("occurrences", []):
        if not isinstance(occurrence, Mapping):
            continue
        ref_id = str(occurrence.get("ref_id") or "").strip()
        paper_id = str(occurrence.get("paper_id") or "").strip()
        if ref_id and paper_id in entries:
            by_ref[ref_id] = entries[paper_id]
    return by_ref


def render_structured_citations(
    text: str,
    generator_instance: Any,
    citation_manifest: Mapping[str, Any],
) -> tuple[str, List[str]]:
    del generator_instance
    lookup = _entry_lookup(citation_manifest)
    unresolved: list[str] = []
    raw = str(text or "")

    # --- Group adjacent citation tokens into one multi-id group -----------
    # Writer emission often produces consecutive single-ref tokens, e.g.
    #   [[cite_ref:R006]][[cite_ref:R009]]
    # which would otherwise render as "(A)(B)".  Merge a maximal run of
    # adjacent tokens into a single token with comma-separated ref ids:
    #   [[cite_ref:R006, R009]] -> "(A; B)"
    def _group_adjacent(run_match: re.Match[str]) -> str:
        run = run_match.group(0)
        ids: list[str] = []
        for token in re.findall(r"\[\[cite_ref:[^\]]+\]\]", run):
            for ref_id in extract_ref_ids_from_token(token):
                if ref_id not in ids:
                    ids.append(ref_id)
        if len(ids) <= 1:
            return run
        return f"[[cite_ref:{', '.join(ids)}]]"

    rendered_text = re.sub(r"(?:\[\[cite_ref:[^\]]+\]\])+", _group_adjacent, raw)

    # --- Normalize missing space before a citation group ------------------
    # "text[[cite_ref:R008]]" -> "text [[cite_ref:R008]]"; the renderer never
    # depends on the model emitting the space itself.
    def _ensure_space_before(match: re.Match[str]) -> str:
        return f"{match.group(1)} {match.group(2)}"

    rendered_text = re.sub(
        r"([^\s(,，.。])(\[\[cite_ref:)", _ensure_space_before, rendered_text
    )

    def replace(match: re.Match[str]) -> str:
        token = match.group(0)
        ref_ids = extract_ref_ids_from_token(token)
        if not ref_ids:
            unresolved.append(token)
            return token
        rendered: list[str] = []
        for ref_id in ref_ids:
            entry = lookup.get(ref_id)
            if entry is None:
                unresolved.append(ref_id)
                continue
            value = format_in_text_citation(entry, mode="parenthetical").strip("()")
            rendered.append(value)
        if len(rendered) != len(ref_ids):
            return token
        return f"({'; '.join(rendered)})"

    def record_legacy(match: re.Match[str]) -> str:
        token = match.group(0)
        unresolved.append(token)
        return token

    rendered = re.sub(r"\[\[cite:(?!ref:)[^\]]+\]\]", record_legacy, rendered_text)
    rendered = re.sub(r"\[\[cite_ref:[^\]]+\]\]", replace, rendered)
    return rendered.replace("`", ""), unresolved


def set_advanced_document_styles(
    doc: Any,
    font_name: str = "Times New Roman",
    font_size_body: int = 12,
    font_size_heading1: int = 16,
    font_size_heading2: int = 14,
) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size_body)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    for name, size in (("Heading 1", font_size_heading1), ("Heading 2", font_size_heading2)):
        style = doc.styles[name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_header_and_footer(doc: Any, title: str = "Literature Review") -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._element.extend((begin, instruction, end))


def append_section_to_word_document(
    generator_instance: Any,
    section_number: int,
    section_title: str,
    section_text: str,
    word_file: str,
    *,
    citation_manifest: Mapping[str, Any],
) -> bool:
    try:
        output = Path(word_file)
        doc = Document(str(output)) if output.is_file() else Document()
        if not output.is_file():
            set_advanced_document_styles(doc)
            add_header_and_footer(doc)
        rendered, unresolved = render_structured_citations(
            section_text,
            generator_instance,
            citation_manifest,
        )
        if unresolved:
            raise ValueError("unresolved citation references: " + ", ".join(sorted(set(unresolved))))
        doc.add_heading(f"{section_number}. {section_title}", level=2)
        for paragraph in rendered.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output))
        return True
    except Exception as exc:
        _log(getattr(generator_instance, "logger", None), "error", str(exc))
        return False


def generate_apa_references_from_manifest(
    citation_manifest: Mapping[str, Any],
    generator_instance: Any,
) -> List[str]:
    del generator_instance
    references: list[str] = []
    for entry in citation_manifest.get("bibliography", []):
        if isinstance(entry, Mapping) and entry.get("is_cited", True):
            text = str(entry.get("citation_text") or "").strip()
            if text:
                references.append(text)
    if not references:
        for entry in citation_manifest.get("paper_entries", []):
            if not isinstance(entry, Mapping):
                continue
            text = format_reference_entry(
                CitationCatalogEntry(
                    index=len(references) + 1,
                    paper_id=str(entry.get("paper_id") or ""),
                    paper_key=str(entry.get("paper_key") or entry.get("paper_id") or ""),
                    title=str(entry.get("title") or ""),
                    authors=[str(item) for item in entry.get("authors", [])],
                    year=str(entry.get("year") or ""),
                    journal=str(entry.get("journal") or ""),
                    doi=str(entry.get("doi") or ""),
                    aliases=[str(item) for item in entry.get("aliases", [])],
                )
            )
            if text:
                references.append(text)
    return list(dict.fromkeys(references))


def scan_docx_for_unresolved_citation_tokens(
    docx_path: str,
    citation_manifest: Mapping[str, Any] | None = None,
) -> Dict[str, Any]:
    doc = Document(docx_path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    raw_tokens = re.findall(r"\[\[(?:cite_ref|cite):[^\]]+\]\]", text)
    legal_tokens = LEGAL_CITE_REF_TOKEN_PATTERN.findall(text)
    known_refs = {
        str(item.get("ref_id") or "")
        for item in (citation_manifest or {}).get("occurrences", [])
        if isinstance(item, Mapping)
    }
    unresolved = [
        token
        for token in raw_tokens
        if not token.startswith("[[cite_ref:")
        or not set(extract_ref_ids_from_token(token)).issubset(known_refs)
    ]
    return {
        "docx_path": docx_path,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "legal_tokens": legal_tokens,
        "unresolved_tokens": sorted(set(unresolved)),
        "bare_ref_ids": [],
        "references_seen": "References" in text,
        "passed": not unresolved,
    }


def rebuild_review_docx_from_structured_artifacts(
    generator_instance: Any,
    review_draft: Mapping[str, Any],
    citation_manifest: Mapping[str, Any],
    output_path: str,
) -> None:
    output = Path(output_path)
    if output.exists():
        output.unlink()
    for section in review_draft.get("content", {}).get("sections", []):
        text = "\n\n".join(
            str(block.get("text") or "").strip()
            for block in section.get("blocks", [])
            if isinstance(block, Mapping) and str(block.get("text") or "").strip()
        )
        if not append_section_to_word_document(
            generator_instance,
            int(section.get("section_number") or 0),
            str(section.get("section_title") or ""),
            text,
            str(output),
            citation_manifest=citation_manifest,
        ):
            raise ValueError("section DOCX rendering failed")
    doc = Document(str(output))
    doc.add_heading("References", level=1)
    for reference in generate_apa_references_from_manifest(citation_manifest, generator_instance):
        doc.add_paragraph(reference)
    doc.save(str(output))
    report = scan_docx_for_unresolved_citation_tokens(str(output), citation_manifest)
    if not report["passed"]:
        raise ValueError("DOCX contains unresolved citation tokens")


def rebuild_final_docx_from_manifest(
    generator_instance: Any,
    review_draft: Mapping[str, Any],
    citation_manifest: Mapping[str, Any],
    output_path: str,
    *,
    scan_report_path: str = "",
) -> Dict[str, Any]:
    rebuild_review_docx_from_structured_artifacts(
        generator_instance,
        review_draft,
        citation_manifest,
        output_path,
    )
    report = scan_docx_for_unresolved_citation_tokens(output_path, citation_manifest)
    if scan_report_path:
        target = Path(scan_report_path)
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def create_word_document(generator_instance: Any, markdown_text: str, output_path: str) -> bool:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_advanced_document_styles(doc)
    add_header_and_footer(doc)
    for line in str(markdown_text or "").splitlines():
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.strip():
            doc.add_paragraph(line.strip())
    doc.save(str(output))
    _log(getattr(generator_instance, "logger", None), "info", f"DOCX written: {output}")
    return output.is_file()
