import json
from pathlib import Path
from typing import cast

from docx import Document

import main
from config_loader import ConfigDict
from services.artifact_registry import ArtifactRegistry
from services.config_compat import CompatConfigView
from services.job_workspace import JobWorkspace
from services.progress_state import ResumeStateReport


class _DummyLogger:
    def info(self, *_args, **_kwargs):
        pass

    def warning(self, *_args, **_kwargs):
        pass

    def error(self, *_args, **_kwargs):
        pass

    def success(self, *_args, **_kwargs):
        pass

    def debug(self, *_args, **_kwargs):
        pass


def _resume_report(workspace: JobWorkspace) -> ResumeStateReport:
    return ResumeStateReport(
        artifact_type="resume_state_report",
        artifact_version="v1",
        created_from_job_id=workspace.job_id,
        created_at="2026-04-03T00:00:00Z",
        project_name=workspace.project_name,
        job_id=workspace.job_id,
        state="non_resumable",
        reason="test bootstrap",
        summary_file=workspace.artifact_path(f"{workspace.project_name}_summaries.json"),
        progress_snapshot_file=None,
        checkpoint_file=workspace.checkpoint_path(f"{workspace.project_name}_checkpoint.json"),
        fingerprint_bundle={"request": "demo"},
    )


def _make_bound_generator(tmp_path: Path, project_name: str = "demo", job_id: str | None = None):
    output_dir = tmp_path / "output"
    workspace = JobWorkspace.create(str(output_dir), project_name, job_id=job_id)
    registry = ArtifactRegistry(workspace.paths.registry_path, workspace.job_id)
    config = ConfigDict(
        {
            "Paths": {"output_path": str(output_dir)},
            "Writer_API": {"api_key": "writer-key", "model": "writer-model", "api_base": "https://example.com/v1"},
            "Validation": {"stage1_enabled": "false", "stage2_enabled": "false"},
            "Styling": {
                "font_name": "Times New Roman",
                "font_size_body": "12",
                "font_size_heading1": "16",
                "font_size_heading2": "14",
            },
        }
    )
    compat_view = CompatConfigView.from_config(config)

    generator = main.LiteratureReviewGenerator(project_name=project_name, pdf_folder=None)
    generator.logger = cast(main.CustomLogger, _DummyLogger())
    generator.config = config
    generator.bind_job_workspace(
        workspace=workspace,
        artifact_registry=registry,
        compat_config=compat_view,
        fingerprint_bundle={"request": "demo"},
        resume_state_report=_resume_report(workspace),
    )
    generator.summaries = [{"status": "success", "paper_info": {"title": "Paper A"}}]
    generator.summary_file = workspace.artifact_path(f"{project_name}_summaries.json")
    Path(generator.summary_file).write_text(json.dumps(generator.summaries), encoding="utf-8")
    return generator, workspace, registry


def _stub_stage2_bootstrap(monkeypatch, generator) -> None:
    monkeypatch.setattr(generator, "load_configuration", lambda: True)
    monkeypatch.setattr(generator, "setup_output_directory", lambda: True)
    monkeypatch.setattr(generator, "load_existing_summaries", lambda: True)
    monkeypatch.setattr(generator, "_stage2_validation_enabled", lambda: False)
    monkeypatch.setattr(generator, "generate_word_table_of_contents", lambda _doc: True)


def test_successful_stage2_generation_creates_registered_review_draft(tmp_path: Path, monkeypatch) -> None:
    generator, workspace, _registry = _make_bound_generator(tmp_path, job_id="job-review-success")
    _stub_stage2_bootstrap(monkeypatch, generator)

    outline_text = "# Demo Outline\n\n## 1. First Section\n\n## 2. Second Section"
    outline_path = Path(workspace.artifact_path("demo_literature_review_outline.md"))
    outline_path.write_text(outline_text, encoding="utf-8")

    monkeypatch.setattr(generator, "_load_outline_artifact", lambda: (str(outline_path), outline_text))
    monkeypatch.setattr(
        generator,
        "generate_review_section_content",
        lambda section_title, _outline: f"{section_title} generated content.",
    )
    monkeypatch.setattr(generator, "generate_apa_references", lambda: ["Author, A. (2024). Demo reference."])

    assert generator.generate_full_review_from_outline() is True

    word_path = Path(workspace.report_path("demo_literature_review.docx"))
    registry_payload = json.loads(Path(workspace.paths.registry_path).read_text(encoding="utf-8"))
    review_records = [item for item in registry_payload["artifacts"] if item["artifact_type"] == "review_draft" and item["artifact_version"] == "v1"]

    assert word_path.exists() is True
    assert len(review_records) == 1

    artifact_path = Path(review_records[0]["path"])
    artifact_payload = json.loads(artifact_path.read_text(encoding="utf-8"))
    document_text = "\n".join(paragraph.text for paragraph in Document(str(word_path)).paragraphs)

    assert artifact_payload["artifact_type"] == "review_draft"
    assert artifact_payload["artifact_version"] == "v1"
    assert artifact_payload["created_from_job_id"] == workspace.job_id
    assert artifact_payload["draft_identity"]["draft_id"] == "review_draft:full_review"
    assert artifact_payload["generation_context"]["outline_artifact_id"] == generator.OUTLINE_ARTIFACT_ID
    assert artifact_payload["generation_context"]["outline_source_path"] == str(outline_path)
    assert len(artifact_payload["content"]["sections"]) == 2
    assert artifact_payload["content"]["sections"][0]["section_title"] == "First Section"
    assert artifact_payload["content"]["sections"][1]["content"] == "Second Section generated content."
    assert artifact_payload["content"]["references"] == ["Author, A. (2024). Demo reference."]
    assert artifact_payload["projections"]["docx_path"] == str(word_path)
    assert "First Section generated content." in document_text
    assert any(dep["artifact_type"] == "literature_review_outline" for dep in review_records[0]["depends_on"])


def test_stage2_with_failed_sections_does_not_register_review_draft(tmp_path: Path, monkeypatch) -> None:
    generator, workspace, _registry = _make_bound_generator(tmp_path, job_id="job-review-failed")
    _stub_stage2_bootstrap(monkeypatch, generator)

    outline_text = "# Demo Outline\n\n## 1. First Section\n\n## 2. Second Section"
    outline_path = Path(workspace.artifact_path("demo_literature_review_outline.md"))
    outline_path.write_text(outline_text, encoding="utf-8")

    def _section_content(section_title: str, _outline: str):
        if section_title == "Second Section":
            return None
        return f"{section_title} generated content."

    monkeypatch.setattr(generator, "_load_outline_artifact", lambda: (str(outline_path), outline_text))
    monkeypatch.setattr(generator, "generate_review_section_content", _section_content)
    monkeypatch.setattr(generator, "generate_apa_references", lambda: [])

    assert generator.generate_full_review_from_outline() is True

    word_path = Path(workspace.report_path("demo_literature_review.docx"))
    draft_path = Path(workspace.artifact_path("review_drafts/demo_review_draft_v1.json"))
    failed_sections_path = Path(workspace.report_path("demo_failed_review_sections.json"))
    registry_path = Path(workspace.paths.registry_path)

    assert word_path.exists() is True
    assert draft_path.exists() is False
    assert failed_sections_path.exists() is True
    if registry_path.exists():
        registry_payload = json.loads(registry_path.read_text(encoding="utf-8"))
        assert not any(item["artifact_type"] == "review_draft" for item in registry_payload["artifacts"])


def test_review_draft_resume_path_keeps_existing_sections_and_registers_on_completion(
    tmp_path: Path,
    monkeypatch,
) -> None:
    generator, workspace, _registry = _make_bound_generator(tmp_path, job_id="job-review-resume")
    _stub_stage2_bootstrap(monkeypatch, generator)

    outline_text = "# Demo Outline\n\n## 1. First Section\n\n## 2. Second Section"
    outline_path = Path(workspace.artifact_path("demo_literature_review_outline.md"))
    outline_path.write_text(outline_text, encoding="utf-8")

    word_path = Path(workspace.report_path("demo_literature_review.docx"))
    existing_doc = Document()
    existing_doc.add_heading("Literature Review", 0)
    existing_doc.add_heading("1. First Section", level=1)
    existing_doc.add_paragraph("Existing section content.")
    existing_doc.save(str(word_path))

    checkpoint_path = Path(workspace.checkpoint_path("demo_review_checkpoint.json"))
    checkpoint_path.write_text(
        json.dumps(
            {
                "last_completed_section": 1,
                "last_section_title": "First Section",
                "update_time": "2026-04-03 00:00:00",
            }
        ),
        encoding="utf-8",
    )

    monkeypatch.setattr(generator, "_load_outline_artifact", lambda: (str(outline_path), outline_text))
    monkeypatch.setattr(
        generator,
        "generate_review_section_content",
        lambda section_title, _outline: f"{section_title} generated content.",
    )
    monkeypatch.setattr(generator, "generate_apa_references", lambda: ["Author, A. (2024). Demo reference."])

    assert generator.generate_full_review_from_outline() is True

    registry_payload = json.loads(Path(workspace.paths.registry_path).read_text(encoding="utf-8"))
    review_records = [item for item in registry_payload["artifacts"] if item["artifact_type"] == "review_draft" and item["artifact_version"] == "v1"]
    artifact_path = Path(review_records[0]["path"])
    artifact_payload = json.loads(artifact_path.read_text(encoding="utf-8"))

    assert len(review_records) == 1
    assert [section["section_number"] for section in artifact_payload["content"]["sections"]] == [1, 2]
    assert artifact_payload["content"]["sections"][0]["content"] == "Existing section content."
    assert artifact_payload["content"]["sections"][1]["content"] == "Second Section generated content."


def test_retry_failed_review_sections_still_reenters_full_review_generation(
    tmp_path: Path,
    monkeypatch,
) -> None:
    generator, workspace, _registry = _make_bound_generator(tmp_path, job_id="job-review-retry")
    _stub_stage2_bootstrap(monkeypatch, generator)

    failed_sections_path = Path(workspace.report_path("demo_failed_review_sections.json"))
    failed_sections_path.write_text(
        json.dumps(
            {
                "failed_sections": [
                    {
                        "section_number": 2,
                        "section_title": "Second Section",
                        "failure_reason": "section_content_generation_failed",
                        "update_time": "2026-04-03 00:00:00",
                    }
                ],
                "updated_at": "2026-04-03 00:00:00",
            }
        ),
        encoding="utf-8",
    )

    word_path = Path(workspace.report_path("demo_literature_review.docx"))
    word_path.write_bytes(b"placeholder")

    captured = {}
    monkeypatch.setattr(generator, "_trim_review_document_from_section", lambda _word_file, _section_number: True)

    def _fake_full_review() -> bool:
        checkpoint_payload = json.loads(
            Path(workspace.checkpoint_path("demo_review_checkpoint.json")).read_text(encoding="utf-8")
        )
        captured["last_completed_section"] = checkpoint_payload["last_completed_section"]
        return True

    monkeypatch.setattr(generator, "generate_full_review_from_outline", _fake_full_review)

    assert generator.retry_failed_review_sections() is True
    assert captured["last_completed_section"] == 1


def test_successful_stage2_generation_creates_registered_review_draft_v2(tmp_path: Path, monkeypatch) -> None:
    """Test that review_draft_v2 is created and registered alongside v1."""
    generator, workspace, _registry = _make_bound_generator(tmp_path, job_id="job-review-v2-success")
    _stub_stage2_bootstrap(monkeypatch, generator)

    outline_text = "# Demo Outline\n\n## 1. First Section\n\n## 2. Second Section"
    outline_path = Path(workspace.artifact_path("demo_literature_review_outline.md"))
    outline_path.write_text(outline_text, encoding="utf-8")

    monkeypatch.setattr(generator, "_load_outline_artifact", lambda: (str(outline_path), outline_text))
    monkeypatch.setattr(
        generator,
        "generate_review_section_content",
        lambda section_title, _outline: f"{section_title} generated content.\n\nSecond paragraph for {section_title}.",
    )
    monkeypatch.setattr(generator, "generate_apa_references", lambda: ["Author, A. (2024). Demo reference."])

    assert generator.generate_full_review_from_outline() is True

    registry_payload = json.loads(Path(workspace.paths.registry_path).read_text(encoding="utf-8"))
    review_v2_records = [item for item in registry_payload["artifacts"] if item.get("artifact_type") == "review_draft" and item.get("artifact_version") == "v2"]
    v2_artifact_path = Path(workspace.artifact_path("review_drafts/demo_review_draft_v2.json"))

    assert len(review_v2_records) == 1
    assert v2_artifact_path.exists() is True

    artifact_payload = json.loads(v2_artifact_path.read_text(encoding="utf-8"))
    assert artifact_payload["artifact_type"] == "review_draft"
    assert artifact_payload["artifact_version"] == "v2"
    assert artifact_payload["created_from_job_id"] == workspace.job_id
    assert artifact_payload["draft_identity"]["draft_id"] == "review_draft_v2:full_review"
    assert len(artifact_payload["content"]["sections"]) == 2

    # Check block structure
    first_section = artifact_payload["content"]["sections"][0]
    assert "blocks" in first_section
    assert len(first_section["blocks"]) >= 1
    assert first_section["blocks"][0]["block_id"] == "s1_b1"
    assert first_section["blocks"][0]["block_kind"] == "paragraph"
    assert first_section["blocks"][0]["block_order"] == 1
    assert "text" in first_section["blocks"][0]
    assert "anchor_text" in first_section["blocks"][0]


def test_stage2_with_failed_sections_does_not_register_review_draft_v2(tmp_path: Path, monkeypatch) -> None:
    """Test that review_draft_v2 is NOT created when sections fail."""
    generator, workspace, _registry = _make_bound_generator(tmp_path, job_id="job-review-v2-failed")
    _stub_stage2_bootstrap(monkeypatch, generator)

    outline_text = "# Demo Outline\n\n## 1. First Section\n\n## 2. Second Section"
    outline_path = Path(workspace.artifact_path("demo_literature_review_outline.md"))
    outline_path.write_text(outline_text, encoding="utf-8")

    def _section_content(section_title: str, _outline: str):
        if section_title == "Second Section":
            return None
        return f"{section_title} generated content."

    monkeypatch.setattr(generator, "_load_outline_artifact", lambda: (str(outline_path), outline_text))
    monkeypatch.setattr(generator, "generate_review_section_content", _section_content)
    monkeypatch.setattr(generator, "generate_apa_references", lambda: [])

    assert generator.generate_full_review_from_outline() is True

    v2_draft_path = Path(workspace.artifact_path("review_drafts/demo_review_draft_v2.json"))
    registry_path = Path(workspace.paths.registry_path)

    assert v2_draft_path.exists() is False
    if registry_path.exists():
        registry_payload = json.loads(registry_path.read_text(encoding="utf-8"))
        assert not any(item.get("artifact_version") == "v2" for item in registry_payload["artifacts"])


def test_review_draft_v2_written_to_job_workspace_not_project_root(tmp_path: Path, monkeypatch) -> None:
    """Test that review_draft_v2 is written to job workspace, not output/<project>/."""
    generator, workspace, _registry = _make_bound_generator(tmp_path, job_id="job-review-v2-location")
    _stub_stage2_bootstrap(monkeypatch, generator)

    outline_text = "# Demo Outline\n\n## 1. Single Section"
    outline_path = Path(workspace.artifact_path("demo_literature_review_outline.md"))
    outline_path.write_text(outline_text, encoding="utf-8")

    monkeypatch.setattr(generator, "_load_outline_artifact", lambda: (str(outline_path), outline_text))
    monkeypatch.setattr(
        generator,
        "generate_review_section_content",
        lambda section_title, _outline: f"{section_title} content.",
    )
    monkeypatch.setattr(generator, "generate_apa_references", lambda: [])

    assert generator.generate_full_review_from_outline() is True

    # Verify v2 is in job workspace artifacts directory
    v2_path = Path(workspace.artifact_path("review_drafts/demo_review_draft_v2.json"))
    assert v2_path.exists() is True
    assert str(v2_path).startswith(str(workspace.paths.root_dir))

    # Verify it's NOT in the legacy output location
    legacy_path = Path(tmp_path / "output" / "demo" / "demo_review_draft_v2.json")
    assert not legacy_path.exists()


def test_review_draft_v1_and_v2_coexist_in_registry(tmp_path: Path, monkeypatch) -> None:
    """Test that both v1 and v2 can coexist in the artifact registry."""
    generator, workspace, _registry = _make_bound_generator(tmp_path, job_id="job-review-both-versions")
    _stub_stage2_bootstrap(monkeypatch, generator)

    outline_text = "# Demo Outline\n\n## 1. First Section"
    outline_path = Path(workspace.artifact_path("demo_literature_review_outline.md"))
    outline_path.write_text(outline_text, encoding="utf-8")

    monkeypatch.setattr(generator, "_load_outline_artifact", lambda: (str(outline_path), outline_text))
    monkeypatch.setattr(
        generator,
        "generate_review_section_content",
        lambda section_title, _outline: f"{section_title} content.",
    )
    monkeypatch.setattr(generator, "generate_apa_references", lambda: [])

    assert generator.generate_full_review_from_outline() is True

    registry_payload = json.loads(Path(workspace.paths.registry_path).read_text(encoding="utf-8"))
    review_records = [item for item in registry_payload["artifacts"] if item["artifact_type"] == "review_draft"]

    v1_records = [r for r in review_records if r["artifact_version"] == "v1"]
    v2_records = [r for r in review_records if r["artifact_version"] == "v2"]

    assert len(v1_records) == 1
    assert len(v2_records) == 1
    assert v1_records[0]["artifact_id"] == "review_draft:full_review"
    assert v2_records[0]["artifact_id"] == "review_draft_v2:full_review"
    assert v1_records[0]["artifact_role"] == "review_draft"
    assert v2_records[0]["artifact_role"] == "review_draft_v2"
